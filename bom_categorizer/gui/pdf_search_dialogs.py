# -*- coding: utf-8 -*-
"""
Диалоги для поиска PDF документации
"""

import os
import sys
import json
import platform
import subprocess
from typing import Optional, Dict, List

from PySide6.QtWidgets import (
    QDialog, QVBoxLayout, QHBoxLayout, QLabel, QPushButton,
    QLineEdit, QTextEdit, QGroupBox, QComboBox, QListWidget,
    QListWidgetItem, QFileDialog, QMessageBox, QTabWidget,
    QWidget, QGridLayout, QTextBrowser, QCheckBox, QFormLayout, QDialogButtonBox,
    QSpinBox, QToolTip, QApplication
)
from PySide6.QtCore import Qt, Signal, QThread, QUrl, QTimer
from PySide6.QtGui import QFont, QTextCursor, QColor, QDesktopServices


class PDFSearchDialog(QDialog):
    """Главный диалог поиска PDF"""
    
    def __init__(self, parent, config: dict, unlocked: bool = True, expert_mode: bool = True):
        super().__init__(parent)
        self.parent_window = parent
        self.config = config
        self.unlocked = unlocked
        self.expert_mode = expert_mode
        
        self.setWindowTitle("🔍 Поиск PDF документации")
        self.setModal(False)
        self.resize(730, 900)  # Увеличена ширина на 30% (900 -> 1170)
        
        self._create_ui()
        
    def closeEvent(self, event):
        """Обработчик закрытия диалога"""
        # Скрываем любые активные подсказки, которые могли "застрять"
        QToolTip.hideText()
        
        # Если есть активный воркер поиска, отключаем его сигналы
        if hasattr(self, 'ai_worker') and self.ai_worker:
             try:
                 if self.ai_worker.isRunning():
                     # Отключаем слот, чтобы результат не пытался обновить закрытое окно
                     try:
                        self.ai_worker.finished.disconnect()
                     except Exception:
                        pass
             except Exception:
                 pass
                 
        event.accept()
        
    def _create_ui(self):
        """Создает интерфейс"""
        layout = QVBoxLayout(self)
        
        # Поле поиска
        search_layout = QHBoxLayout()
        search_label = QLabel("Компонент:")
        search_label.setFixedWidth(100)
        
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("Введите название компонента (например: HMC449, LM358)")
        self.search_input.returnPressed.connect(self.on_search)
        
        search_btn = QPushButton("🔎 Найти")
        search_btn.clicked.connect(self.on_search)
        search_btn.setFixedWidth(100)
        
        search_layout.addWidget(search_label)
        search_layout.addWidget(self.search_input)
        search_layout.addWidget(search_btn)
        layout.addLayout(search_layout)
        
        # Вкладки для разных типов поиска
        self.tabs = QTabWidget()
        
        # Вкладка локального поиска - доступна всегда
        self.local_tab = self._create_local_tab()
        self.tabs.addTab(self.local_tab, "📁 Локальный поиск")
        
        # Вкладка AI поиска - только для разблокированных экспертов
        if self.unlocked and self.expert_mode:
            self.ai_tab = self._create_ai_tab()
            self.tabs.addTab(self.ai_tab, "🤖 AI поиск")
        else:
            # Создаем заглушку для AI вкладки
            self.ai_tab = QWidget()
            ai_layout = QVBoxLayout(self.ai_tab)
            ai_layout.addStretch()
            
            lock_label = QLabel("🔒 AI поиск доступен только в экспертном режиме после разблокировки приложения")
            lock_label.setAlignment(Qt.AlignCenter)
            lock_label.setStyleSheet("color: #f38ba8; font-size: 14pt; font-weight: bold;")
            ai_layout.addWidget(lock_label)
            
            hint_label = QLabel("Дважды кликните на имя разработчика для разблокировки")
            hint_label.setAlignment(Qt.AlignCenter)
            hint_label.setStyleSheet("color: #cdd6f4; font-size: 12pt;")
            ai_layout.addWidget(hint_label)
            
            ai_layout.addStretch()
            self.tabs.addTab(self.ai_tab, "🔒 AI поиск")
            # Отключаем вкладку
            self.tabs.setTabEnabled(1, False)
        
        layout.addWidget(self.tabs)
        
        # Кнопки
        button_layout = QHBoxLayout()
        
        # Кнопка настроек - только для разблокированных экспертов
        if self.unlocked and self.expert_mode:
            settings_btn = QPushButton("⚙️ Настройки")
            settings_btn.clicked.connect(self.open_settings)
            button_layout.addWidget(settings_btn)
        
        button_layout.addStretch()
        
        close_btn = QPushButton("Закрыть")
        close_btn.clicked.connect(self.accept)
        button_layout.addWidget(close_btn)
        
        layout.addLayout(button_layout)
    
    def _create_local_tab(self) -> QWidget:
        """Создает вкладку локального поиска"""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        # Папки для поиска
        path_group = QGroupBox("📂 Папки для поиска (рекурсивно)")
        path_layout = QVBoxLayout()
        
        # Список папок
        self.search_dirs_list = QListWidget()
        self.search_dirs_list.setMaximumHeight(156)  # Увеличено на 30% (120 -> 156)
        self.search_dirs_list.setToolTip("Список папок, в которых будет выполняться поиск PDF файлов\nДвойной клик по пути откроет папку в проводнике")
        # Обработчик двойного клика для открытия папки в проводнике
        self.search_dirs_list.itemDoubleClicked.connect(self.open_search_directory)
        path_layout.addWidget(self.search_dirs_list)
        
        # Кнопки управления путями
        buttons_layout = QHBoxLayout()
        
        add_dir_btn = QPushButton("➕ Добавить папку")
        add_dir_btn.clicked.connect(self.add_search_directory)
        add_dir_btn.setToolTip("Добавить временную папку для поиска")
        buttons_layout.addWidget(add_dir_btn)
        
        remove_dir_btn = QPushButton("➖ Удалить")
        remove_dir_btn.clicked.connect(self.remove_search_directory)
        remove_dir_btn.setToolTip("Удалить выбранную папку из списка")
        buttons_layout.addWidget(remove_dir_btn)
        
        save_to_config_btn = QPushButton("💾 Сохранить в конфиг")
        save_to_config_btn.clicked.connect(self.save_search_dirs_to_config)
        save_to_config_btn.setToolTip("Сохранить текущие папки в config_qt.json как пользовательские")
        buttons_layout.addWidget(save_to_config_btn)
        
        reset_btn = QPushButton("🔄 Сброс")
        reset_btn.clicked.connect(self.reset_search_directories)
        reset_btn.setToolTip("Вернуть список папок по умолчанию")
        buttons_layout.addWidget(reset_btn)
        
        buttons_layout.addStretch()
        path_layout.addLayout(buttons_layout)
        
        path_group.setLayout(path_layout)
        layout.addWidget(path_group)
        
        # Загружаем папки по умолчанию
        self._load_default_search_dirs()
        
        # Результаты
        results_label = QLabel("Найденные файлы:")
        results_label.setProperty("class", "bold")
        layout.addWidget(results_label)
        
        self.local_results_list = QListWidget()
        self.local_results_list.setMinimumHeight(200)  # Увеличена минимальная высота для результатов
        self.local_results_list.itemDoubleClicked.connect(self.open_local_file)
        layout.addWidget(self.local_results_list)
        
        # Кнопки действий
        actions_layout = QHBoxLayout()
        
        open_file_btn = QPushButton("📄 Открыть файл")
        open_file_btn.clicked.connect(self.open_selected_local_file)
        actions_layout.addWidget(open_file_btn)
        
        open_folder_btn = QPushButton("📁 Открыть папку")
        open_folder_btn.clicked.connect(self.open_local_file_folder)
        actions_layout.addWidget(open_folder_btn)
        
        actions_layout.addStretch()
        layout.addLayout(actions_layout)
        
        return widget
    
    def _create_ai_tab(self) -> QWidget:
        """Создает вкладку AI поиска"""
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        # Выбор провайдера
        provider_layout = QHBoxLayout()
        provider_label = QLabel("Провайдер AI:")
        provider_label.setFixedWidth(100)
        
        self.ai_provider_combo = QComboBox()
        self.ai_provider_combo.addItems(["Anthropic Claude", "OpenAI GPT-4o", "Telegram Bot"])
        self.ai_provider_combo.setCurrentIndex(2)  # По умолчанию Telegram Bot
        self.ai_provider_combo.setFixedWidth(200)
        
        provider_layout.addWidget(provider_label)
        provider_layout.addWidget(self.ai_provider_combo)
        provider_layout.addStretch()
        layout.addLayout(provider_layout)
        
        # === ВЫБОР ТИПА ПРОМПТА ===
        prompt_group = QGroupBox("📝 Тип запроса")
        prompt_layout = QVBoxLayout()
        
        # Выпадающий список готовых промптов
        prompt_select_layout = QHBoxLayout()
        prompt_select_label = QLabel("Готовый промпт:")
        prompt_select_label.setFixedWidth(120)
        
        self.prompt_combo = QComboBox()
        self.prompt_combo.addItems([
            "🔧 Стандартный",
            "📋 Краткое описание ИВП",
            "📖 Развёрнутое описание ИВП",
            "🔄 Поиск аналогов",
            "📊 Сравнительный анализ",
            "✍️ Свой промпт"
        ])
        self.prompt_combo.currentIndexChanged.connect(self._on_prompt_type_changed)
        
        prompt_select_layout.addWidget(prompt_select_label)
        prompt_select_layout.addWidget(self.prompt_combo, 1)
        prompt_layout.addLayout(prompt_select_layout)
        
        # Поле для своего промпта
        self.custom_prompt_label = QLabel("Свой промпт:")
        self.custom_prompt_label.setVisible(False)
        prompt_layout.addWidget(self.custom_prompt_label)
        
        self.custom_prompt_edit = QTextEdit()
        self.custom_prompt_edit.setPlaceholderText(
            "Введите свой промпт. Используйте {component} для подстановки названия компонента.\n"
            "Например: Опиши компонент {component} и найди его аналоги..."
        )
        self.custom_prompt_edit.setMaximumHeight(100)
        self.custom_prompt_edit.setVisible(False)
        prompt_layout.addWidget(self.custom_prompt_edit)
        
        # Подсказка о текущем промпте
        self.prompt_preview_label = QLabel()
        self.prompt_preview_label.setWordWrap(True)
        self.prompt_preview_label.setStyleSheet("color: #6c7086; font-style: italic; padding: 5px;")
        self._update_prompt_preview()
        prompt_layout.addWidget(self.prompt_preview_label)
        
        prompt_group.setLayout(prompt_layout)
        layout.addWidget(prompt_group)
        
        # === ДОПОЛНИТЕЛЬНЫЙ КОНТЕКСТ (HINT) ===
        hint_group = QGroupBox("💡 Уточняющая подсказка (опционально)")
        hint_layout = QVBoxLayout()
        
        hint_desc = QLabel("Добавьте контекст для уменьшения ошибок AI.")
        hint_desc.setStyleSheet("color: #a6adc8; font-size: 11px;")
        hint_layout.addWidget(hint_desc)
        
        self.hint_edit = QTextEdit()
        self.hint_edit.setPlaceholderText(
            "Например: This is a frequency divider from Analog Devices"
        )
        self.hint_edit.setMinimumHeight(60)
        self.hint_edit.setMaximumHeight(80)
        # Стили для тёмной темы: светлый текст, удобный скроллбар
        self.hint_edit.setStyleSheet("""
            QTextEdit {
                background-color: #1e1e2e;
                color: #cdd6f4;
                border: 1px solid #45475a;
                border-radius: 6px;
                padding: 8px;
                font-size: 13px;
            }
            QTextEdit::placeholder {
                color: #6c7086;
            }
            QScrollBar:vertical {
                background: #313244;
                width: 14px;
                border-radius: 7px;
                margin: 2px;
            }
            QScrollBar::handle:vertical {
                background: #585b70;
                border-radius: 5px;
                min-height: 30px;
            }
            QScrollBar::handle:vertical:hover {
                background: #7f849c;
            }
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {
                height: 0px;
            }
        """)
        hint_layout.addWidget(self.hint_edit)
        
        hint_group.setLayout(hint_layout)
        layout.addWidget(hint_group)
        
        # Результаты AI поиска
        results_label = QLabel("Результаты поиска:")
        results_label.setProperty("class", "bold")
        layout.addWidget(results_label)
        
        self.ai_results_browser = QTextBrowser()
        self.ai_results_browser.setMinimumHeight(250)  # Увеличенная зона результатов
        # Отключаем внутреннюю навигацию, открываем ссылки во внешнем браузере
        self.ai_results_browser.setOpenExternalLinks(False)
        self.ai_results_browser.setOpenLinks(False)
        self.ai_results_browser.anchorClicked.connect(self._open_external_link)
        layout.addWidget(self.ai_results_browser, 1)  # stretch factor для растяжения
        
        # Кнопка сохранения
        save_layout = QHBoxLayout()
        save_btn = QPushButton("💾 Сохранить результат")
        save_btn.clicked.connect(self.save_ai_results)
        save_layout.addWidget(save_btn)
        save_layout.addStretch()
        layout.addLayout(save_layout)
        
        return widget
    
    def _on_prompt_type_changed(self, index: int):
        """Обработчик смены типа промпта"""
        is_custom = index == 5  # "Свой промпт"
        self.custom_prompt_label.setVisible(is_custom)
        self.custom_prompt_edit.setVisible(is_custom)
        self._update_prompt_preview()
    
    def _update_prompt_preview(self):
        """Обновляет превью промпта"""
        index = self.prompt_combo.currentIndex()
        previews = [
            "Стандартный: веб-поиск + проверка по datasheet, статус, ключевые параметры, ссылки",
            "Краткое описание: структурированная справка 150-300 слов с параметрами по типу компонента и аналогами",
            "Развёрнутое описание: полный обзор 400-700 слов с техническими характеристиками, схемотехникой, источниками",
            "Поиск аналогов: pin-to-pin, функциональные, отечественные аналоги + таблица сравнения",
            "Сравнительный анализ: таблица параметров с конкурентами, ценовой анализ, рекомендации",
            "Введите свой текст запроса выше. Используйте {component} для названия компонента"
        ]
        self.prompt_preview_label.setText(f"ℹ️ {previews[index]}")
    
    def _get_prompt_template(self, component_name: str) -> str:
        """Возвращает промпт на основе выбранного типа с учётом подсказки"""
        index = self.prompt_combo.currentIndex()
        
        # Получаем уточняющую подсказку (hint)
        hint = ""
        if hasattr(self, 'hint_edit'):
            hint = self.hint_edit.toPlainText().strip()
        
        prompts = {
            0: f"""Роль: Ты инженер-электронщик, выполняющий поиск компонентов для разработки.

Задача: Найди достоверную информацию о компоненте: {component_name}

ОБЯЗАТЕЛЬНЫЕ ДЕЙСТВИЯ:
1. Выполни веб-поиск по официальному сайту производителя
2. Если компонент найден — открой страницу продукта и/или datasheet
3. Сверь информацию минимум с двумя источниками (производитель + дистрибьютор)

ТРЕБУЕМАЯ ИНФОРМАЦИЯ:
- Полное название (включая все суффиксы маркировки)
- Производитель (оригинальный, не дистрибьютор)
- Статус: Active / NRND / Obsolete / Not Found
- Тип компонента и его функция (1-2 предложения)
- Ключевые параметры (только из datasheet, не додумывать):
  - Для ИС: Vin, Vout, Iout, частота, корпус, рабочая температура
  - Для транзисторов: тип, Vds/Vce, Id/Ic, Rds(on)/hFE, корпус
  - Для пассивных: номинал, допуск, напряжение, корпус
- Прямая ссылка на PDF datasheet (только рабочая, проверенная)
- Где купить (1-2 крупных дистрибьютора с наличием)

КРИТИЧЕСКИ ВАЖНО:
- Если не уверен в параметре — напиши "требует проверки по datasheet"
- Если компонент не найден или название неоднозначно — уточни у меня
- Не путай схожие парт-номера (например LM7805 vs LM78L05)

Формат: структурированный текст с заголовками (не JSON)""",

            1: f"""Роль: Инженер-разработчик радиоэлектронной аппаратуры с опытом во всех направлениях: силовая электроника, аналоговая и цифровая схемотехника, СВЧ-техника, датчики и измерения.

Задача: Составь краткое техническое описание компонента: {component_name}

ОБЯЗАТЕЛЬНО:
1. Выполни поиск на сайте производителя
2. Открой datasheet (и application note если есть)
3. Определи тип компонента и подбери релевантные параметры

СТРУКТУРА ОПИСАНИЯ:

【Идентификация】
- Полное название с суффиксами маркировки
- Производитель (оригинальный)
- Статус: Active / NRND / Obsolete
- Категория компонента (ИС, транзистор, модуль, пассивный, датчик, СВЧ и т.д.)

【Назначение】
- Функция компонента (1-2 предложения)
- Типовые области применения

【Ключевые параметры】
Выбери параметры в зависимости от типа компонента:

▸ DC-DC / LDO / ИВП:
  Vin, Vout, Iout, КПД, Fsw, Iq, защиты (OVP/OCP/OTP), корпус

▸ Операционные усилители:
  Vcc, GBW, Slew Rate, Vos, Ib, шум, число каналов, Rail-to-Rail, корпус

▸ АЦП/ЦАП:
  Разрядность, скорость (SPS/MSPS), интерфейс, Vref, INL/DNL, корпус

▸ Микроконтроллеры / ПЛИС:
  Ядро, Flash/RAM, частота, периферия, интерфейсы, корпус

▸ Транзисторы (MOSFET/BJT/IGBT):
  Тип (N/P), Vds/Vce, Id/Ic, Rds(on)/hFE, Qg, корпус

▸ СВЧ-компоненты (усилители, смесители, генераторы):
  Диапазон частот, Gain, P1dB, NF, OIP3, Vcc/Icc, корпус

▸ Датчики:
  Измеряемая величина, диапазон, точность, интерфейс, Vcc, корпус

▸ Пассивные (если требуется описание):
  Номинал, допуск, напряжение/мощность, температурный коэфф., корпус

【Особенности и преимущества】
- 2-3 ключевых отличия от конкурентов (если известны)

【Аналоги】
- Pin-to-pin совместимые (других производителей)
- Отечественные аналоги: конкретные парт-номера или указание, какие параметры не покрываются
- Вывод: оригинал предпочтителен / есть адекватная замена

【Источники】
- Прямая ссылка на datasheet (PDF)
- Ссылка на страницу продукта

Формат: структурированный текст, 150-300 слов (в зависимости от сложности)
Язык: русский, технические термины допустимы на английском""",

            2: f"""Роль: Инженер-разработчик РЭА, готовящий техническую документацию для проекта.

Задача: Подготовь развёрнутое описание компонента: {component_name}

ОБЯЗАТЕЛЬНЫЕ ДЕЙСТВИЯ:
1. Найди официальный datasheet на сайте производителя
2. Открой и изучи datasheet, application notes, reference designs (если есть)
3. Проверь актуальность: статус производства, последняя ревизия документа
4. При необходимости сверь данные с крупными дистрибьюторами

СТРУКТУРА ОПИСАНИЯ:

═══════════════════════════════════════
1. ИДЕНТИФИКАЦИЯ И СТАТУС
═══════════════════════════════════════
- Полное название (включая все суффиксы маркировки)
- Производитель
- Статус: Active / NRND / Obsolete (с датой если известна)
- Категория: DC-DC / LDO / ОУ / АЦП / МК / транзистор / СВЧ / датчик / и т.д.
- Краткое назначение (1-2 предложения)

═══════════════════════════════════════
2. ТЕХНИЧЕСКИЕ ХАРАКТЕРИСТИКИ
═══════════════════════════════════════
Параметры выбирай в зависимости от типа компонента:

▸ Источники питания (DC-DC, LDO, POL, AC-DC модули):
  - Входное напряжение: диапазон, абсолютный максимум
  - Выходное напряжение: номинал или диапазон регулировки
  - Выходной ток: номинальный, пиковый
  - КПД: при разных нагрузках (указать условия измерения)
  - Частота преобразования: фиксированная/регулируемая
  - Ток покоя (Iq) и ток в shutdown
  - Пульсации выходного напряжения
  - Переходная характеристика (load transient)
  - Защиты: OVP, OCP, SCP, OTP, UVLO (пороги если указаны)
  - Температурный диапазон: рабочий и хранения

▸ Усилители (ОУ, инструментальные, буферы):
  - Напряжение питания, ток потребления
  - GBW, Slew Rate, время установления
  - Входные параметры: Vos, Ib, CMRR, входной диапазон
  - Выходные параметры: размах, ток нагрузки
  - Шум: напряжения и тока
  - Rail-to-Rail: вход/выход

▸ Преобразователи данных (АЦП/ЦАП):
  - Разрядность, архитектура
  - Скорость преобразования
  - Входной диапазон, опорное напряжение
  - Точность: INL, DNL, ENOB, SNR, THD
  - Интерфейс, формат данных

▸ Микроконтроллеры / Процессоры:
  - Ядро, архитектура, разрядность
  - Частота, память (Flash/RAM/EEPROM)
  - Периферия: таймеры, АЦП, интерфейсы
  - Напряжение питания, потребление (активный/sleep)

▸ Силовые ключи (MOSFET/IGBT/GaN/SiC):
  - Тип (N/P), структура
  - Vds/Vce max, Id/Ic (при разных температурах)
  - Rds(on) / Vce(sat) — указать условия
  - Заряд затвора Qg, ёмкости
  - Диод: Vf, trr (для MOSFET)
  - SOA, avalanche rating

▸ СВЧ-компоненты:
  - Диапазон рабочих частот
  - Коэффициент усиления, неравномерность
  - P1dB, OIP3, Psat
  - Коэффициент шума
  - Согласование входа/выхода
  - Питание: напряжение, ток

▸ Датчики:
  - Измеряемая величина, диапазон
  - Точность, разрешение, линейность
  - Время отклика
  - Интерфейс выхода
  - Условия эксплуатации

═══════════════════════════════════════
3. КОНСТРУКТИВ И ТЕПЛОВЫЕ ПАРАМЕТРЫ
═══════════════════════════════════════
- Корпус: тип, размеры (мм), шаг выводов
- Термосопротивление: Rth j-a, Rth j-c
- Максимальная температура кристалла (Tj max)
- Требования к монтажу: пайка, теплоотвод
- Влагостойкость (MSL) если указана

═══════════════════════════════════════
4. ВНЕШНИЕ КОМПОНЕНТЫ И СХЕМОТЕХНИКА
═══════════════════════════════════════
- Минимальная обвязка для работы
- Рекомендуемые номиналы (из datasheet)
- Критичные компоненты и требования к ним
- Особенности разводки PCB (если важны)

═══════════════════════════════════════
5. ОСОБЕННОСТИ И ПРЕИМУЩЕСТВА
═══════════════════════════════════════
- Ключевые отличия от конкурентов
- Уникальные функции или технологии
- Для чего оптимизирован (размер / КПД / цена / точность)

6. АНАЛОГИ И СОВМЕСТИМОСТЬ
═══════════════════════════════════════
- Pin-to-pin совместимые (других производителей)
- Функциональные аналоги (с отличиями)
- Отечественные аналоги: конкретные парт-номера или анализ, 
  какие параметры не покрываются российской базой
- Рекомендация по замене

═══════════════════════════════════════
7. ПРИМЕНЕНИЕ
═══════════════════════════════════════
- Типовые области использования
- Референсные схемы (если есть в документации)
- Ограничения и предостережения

8. ИСТОЧНИКИ
═══════════════════════════════════════
- Ссылка на datasheet (PDF) — обязательно
- Ссылка на application notes (если использовались)
- Ссылка на страницу продукта
- Ревизия документа и дата

Формат: структурированный текст с заголовками
Объём: 400-700 слов (в зависимости от сложности компонента)
Язык: русский, технические термины на английском допустимы

ВАЖНО:
- Указывай только данные из datasheet, не додумывай
- Если параметр не указан в документации — напиши "н/д" или "не указано"
- При расхождении данных между источниками — отметь это явно""",

            3: f"""Роль: Инженер-комплектовщик, подбирающий компонентную базу для серийного производства.

Задача: Найди аналоги для компонента: {component_name}

══════════════════════════════════════════════════
ЭТАП 1: ИДЕНТИФИКАЦИЯ ИСХОДНОГО КОМПОНЕНТА
══════════════════════════════════════════════════

ОБЯЗАТЕЛЬНО сначала:
1. Найди datasheet исходного компонента
2. Определи и зафиксируй:
   • Производитель, полное название
   • Тип и функция
   • Корпус, распиновка
   • Ключевые параметры (5-8 главных для этого типа компонента)
   • Статус: Active / NRND / Obsolete

Это нужно для корректного поиска — без понимания оригинала аналоги будут неточными.

══════════════════════════════════════════════════
ЭТАП 2: ПОИСК АНАЛОГОВ
══════════════════════════════════════════════════

Выполни поиск по:
- Сайтам производителей (cross-reference tools)
- Крупным дистрибьюторам (DigiKey, Mouser, LCSC — у них есть подбор аналогов)
- Специализированным базам (FindChips, Octopart)

══════════════════════════════════════════════════
ЭТАП 3: КАТЕГОРИЗАЦИЯ РЕЗУЛЬТАТОВ
══════════════════════════════════════════════════

【A】PIN-TO-PIN СОВМЕСТИМЫЕ (drop-in replacement)
Критерии: идентичный корпус, распиновка, совместимые параметры

Для каждого (до 5 шт):
- Название, производитель
- Статус производства
- Отличия от оригинала (если есть)
- Совместимость: 100% / требует проверки [указать что]
- Ссылка на datasheet

【Б】ФУНКЦИОНАЛЬНЫЕ АНАЛОГИ (требуют адаптации)
Критерии: та же функция, но другой корпус/распиновка/незначительные отличия параметров

Для каждого (до 5 шт):
- Название, производитель
- Статус производства
- Ключевые отличия от оригинала
- Что потребуется изменить (схема/PCB/обвязка)
- Ссылка на datasheet

【В】ОТЕЧЕСТВЕННЫЕ АНАЛОГИ
- Российские производители: Микрон, Миландр, Элвис, Ангстрем, НИИЭТ и др.
- Для каждого найденного:
  - Название, производитель
  - Параметры vs оригинал (таблица сравнения)
  - Ограничения или компромиссы
  - Ссылка на документацию

- Если аналогов нет — указать:
  - Какие именно параметры не покрываются
  - Ближайший отечественный компонент (если есть частичное совпадение)

【Г】АНАЛОГИ С УЛУЧШЕННЫМИ ПАРАМЕТРАМИ
(если пользователь ищет замену с апгрейдом)
- Новые поколения того же производителя
- Конкуренты с лучшими характеристиками
- Указать: что именно лучше, и есть ли компромиссы

══════════════════════════════════════════════════
ЭТАП 4: СВОДНАЯ ТАБЛИЦА СРАВНЕНИЯ
══════════════════════════════════════════════════

Составь таблицу с ключевыми параметрами:

| Параметр | Оригинал | Аналог 1 | Аналог 2 | Аналог 3 | ... |
|----------|----------|----------|----------|----------|-----|
| Производитель | | | | | |
| Статус | | | | | |
| [Парам 1] | | | | | |
| [Парам 2] | | | | | |
| ... | | | | | |
| Корпус | | | | | |
| Совместимость | — | pin-to-pin / функц. | | | |

Параметры выбирай релевантные типу компонента (как в предыдущих промптах).

══════════════════════════════════════════════════
ЭТАП 5: ДОСТУПНОСТЬ И РЕКОМЕНДАЦИЯ
══════════════════════════════════════════════════

Для топ-3 аналогов проверь:
- Наличие у дистрибьюторов (DigiKey, Mouser, LCSC, российские)
- Минимальный заказ (MOQ) если есть ограничения
- Lead time если указан

Итоговая рекомендация:
- Лучший drop-in replacement: [название] — почему
- Лучший отечественный (если есть): [название] — с какими оговорками
- Альтернатива если оригинал EOL: [название]

══════════════════════════════════════════════════
ИСТОЧНИКИ
══════════════════════════════════════════════════
- Ссылки на все упомянутые datasheet
- Ссылки на страницы продуктов
- Ссылки на cross-reference инструменты (если использовались)

Формат: структурированный текст + таблица сравнения
Язык: русский

ВАЖНО:
- Не выдумывай аналоги — только реально существующие компоненты
- Проверяй статус каждого аналога (obsolete не предлагать как основной вариант)
- Если pin-to-pin аналогов не существует — так и напиши
- При неуверенности в совместимости — отмечай "требует проверки по datasheet" """,

            4: f"""Роль: Инженер-аналитик, выполняющий сравнительный обзор компонентной базы для выбора оптимального решения.

Задача: Проведи сравнительный анализ компонента {component_name} с конкурентами.

══════════════════════════════════════════════════
ЭТАП 1: ИДЕНТИФИКАЦИЯ И КЛАССИФИКАЦИЯ
══════════════════════════════════════════════════

ОБЯЗАТЕЛЬНО сначала:
1. Найди datasheet анализируемого компонента
2. Определи:
   • Полное название, производитель
   • Категория (DC-DC / LDO / ОУ / АЦП / MOSFET / СВЧ / датчик / и т.д.)
   • Ключевая функция (1 предложение)
   • Статус: Active / NRND / Obsolete
   • Позиционирование: бюджетный / mainstream / high-performance

══════════════════════════════════════════════════
ЭТАП 2: ПОДБОР КОНКУРЕНТОВ
══════════════════════════════════════════════════

Выполни поиск конкурентов по:
- Сайтам производителей (parametric search)
- DigiKey/Mouser (фильтры по параметрам)
- Обзорам и application notes

Критерии отбора:
- Та же категория и близкие ключевые параметры
- Разные производители (для объективности)
- Только Active или NRND (obsolete — только для справки)
- 3-6 конкурентов (оптимально 4-5)

Обязательно включи (если существуют):
- Прямого конкурента от TI, Analog Devices, Infineon, ST, ON Semi, NXP, Microchip, Renesas, ROHM (в зависимости от категории)
- Отечественный аналог (Миландр, НИИЭТ, Микрон, Элвис и др.)

══════════════════════════════════════════════════
ЭТАП 3: СРАВНИТЕЛЬНАЯ ТАБЛИЦА
══════════════════════════════════════════════════

Параметры таблицы выбирай в зависимости от типа компонента:

▸ DC-DC / LDO / ИВП:
| Параметр | Компонент | Конкурент 1 | Конкурент 2 | ... |
|----------|-----------|-------------|-------------|-----|
| Производитель | | | | |
| Статус | | | | |
| Топология | | | | |
| Vin (мин-макс) | | | | |
| Vout (диапазон) | | | | |
| Iout max | | | | |
| КПД (при усл.) | | | | |
| Iq / Shutdown | | | | |
| Fsw | | | | |
| Защиты | | | | |
| Корпус | | | | |
| Особенности | | | | |

▸ Операционные усилители:
| Параметр | Компонент | Конкурент 1 | Конкурент 2 | ... |
|----------|-----------|-------------|-------------|-----|
| Производитель | | | | |
| Статус | | | | |
| Каналов | | | | |
| Vcc (диапазон) | | | | |
| GBW | | | | |
| Slew Rate | | | | |
| Vos | | | | |
| Ib | | | | |
| Шум (nV/√Hz) | | | | |
| Rail-to-Rail | | | | |
| Icc | | | | |
| Корпус | | | | |

▸ АЦП / ЦАП:
| Параметр | Компонент | Конкурент 1 | Конкурент 2 | ... |
|----------|-----------|-------------|-------------|-----|
| Производитель | | | | |
| Статус | | | | |
| Разрядность | | | | |
| Скорость (SPS) | | | | |
| Архитектура | | | | |
| INL / DNL | | | | |
| SNR / ENOB | | | | |
| Vref | | | | |
| Интерфейс | | | | |
| Каналов | | | | |
| Потребление | | | | |
| Корпус | | | | |

▸ Силовые транзисторы (MOSFET/IGBT):
| Параметр | Компонент | Конкурент 1 | Конкурент 2 | ... |
|----------|-----------|-------------|-------------|-----|
| Производитель | | | | |
| Статус | | | | |
| Тип (N/P) | | | | |
| Vds/Vce max | | | | |
| Id/Ic (@ 25°C) | | | | |
| Rds(on) (@ Vgs) | | | | |
| Qg | | | | |
| Vth | | | | |
| Диод (Vf, trr) | | | | |
| Корпус | | | | |
| Rth j-c | | | | |

▸ СВЧ-компоненты (LNA, PA, смесители):
| Параметр | Компонент | Конкурент 1 | Конкурент 2 | ... |
|----------|-----------|-------------|-------------|-----|
| Производитель | | | | |
| Статус | | | | |
| Диапазон частот | | | | |
| Gain | | | | |
| NF | | | | |
| P1dB / Psat | | | | |
| OIP3 | | | | |
| Vcc / Icc | | | | |
| Корпус | | | | |

▸ Датчики:
| Параметр | Компонент | Конкурент 1 | Конкурент 2 | ... |
|----------|-----------|-------------|-------------|-----|
| Производитель | | | | |
| Статус | | | | |
| Измер. величина | | | | |
| Диапазон | | | | |
| Точность | | | | |
| Разрешение | | | | |
| Интерфейс | | | | |
| Vcc / Icc | | | | |
| Время отклика | | | | |
| Корпус | | | | |

══════════════════════════════════════════════════
ЭТАП 4: ЦЕНОВОЙ АНАЛИЗ (опционально)
══════════════════════════════════════════════════

Проверь актуальные цены у дистрибьюторов (DigiKey, Mouser, LCSC):
- Цена за 1 шт
- Цена за 100 шт (или 1000 если релевантно)
- Наличие на складе
- Lead time если нет в наличии

Примечание: цены на дату запроса, могут измениться.

══════════════════════════════════════════════════
ЭТАП 5: АНАЛИЗ И ВЫВОДЫ
══════════════════════════════════════════════════

【Сильные стороны {component_name}】
- В чём превосходит конкурентов (конкретные параметры)
- Уникальные функции

【Слабые стороны {component_name}】
- Где уступает конкурентам (конкретные параметры)
- Чего не хватает

【Лучший по параметрам】
- Самый высокий КПД / точность / скорость: [название]
- Самый низкий шум / потребление: [название]
- Лучшее соотношение параметры/цена: [название]

【Рекомендации по выбору】

Выбрать {component_name} если:
- [условие 1]
- [условие 2]

Выбрать [конкурент X] если:
- [условие 1]
- [условие 2]

Выбрать отечественный [название] если:
- [условие — например, требования локализации]
- Ограничения: [какие параметры хуже]

══════════════════════════════════════════════════
ИСТОЧНИКИ
══════════════════════════════════════════════════
- Datasheet каждого компонента (ссылки)
- Страницы продуктов
- Источник цен (дистрибьютор, дата)

Формат: структурированный текст + таблица
Язык: русский

ВАЖНО:
- Все данные только из datasheet — не додумывать
- Указывать условия измерения для корректного сравнения (КПД при какой нагрузке, Rds(on) при каком Vgs и т.д.)
- Если параметр не указан у какого-то компонента — писать "н/д"
- Не сравнивать Active с Obsolete как равноценные варианты"""
        }
        
        # Получаем базовый промпт
        if index == 5:  # Свой промпт
            custom = self.custom_prompt_edit.toPlainText().strip()
            if custom:
                # Заменяем {component} только если он есть в тексте
                base_prompt = custom.replace("{component}", component_name)
                # Добавляем ограничение по длине
                return base_prompt + "\n\nОБЯЗАТЕЛЬНО: Ответ должен быть ограничен 1000 символами."
            else:
                base_prompt = prompts[0]  # Fallback на стандартный
        else:
            base_prompt = prompts.get(index, prompts[0])
        
        # Добавляем уточняющую подсказку, если она есть (только для стандартных промптов)
        if hint:
            hint_instruction = f"""

══════════════════════════════════════════════════
КОНТЕКСТ ОТ ПОЛЬЗОВАТЕЛЯ
══════════════════════════════════════════════════

{hint}

──────────────────────────────────────────────────
ПРАВИЛА ПРИМЕНЕНИЯ КОНТЕКСТА:
──────────────────────────────────────────────────

1. Контекст ДОПОЛНЯЕТ основной запрос, но НЕ ОТМЕНЯЕТ его структуру
2. Если контекст сужает область поиска — применяй как фильтр
3. Если контекст расширяет требования — добавь информацию в соответствующий раздел
4. Если контекст противоречит основному промпту — уточни у пользователя
5. Пустой или нерелевантный контекст — игнорируй

ТИПИЧНЫЕ ПРИМЕНЕНИЯ КОНТЕКСТА:
- Ограничения проекта: "напряжение питания только 3.3В", "корпус не больше QFN-16"
- Приоритеты: "критичен низкий шум", "главное — цена"
- Исключения: "не предлагать компоненты от X", "без китайских производителей"
- Область применения: "для медицинского оборудования", "automotive -40...+125°C"
- Дополнительные требования: "нужна совместимость с существующей платой Rev.B"

ОБЯЗАТЕЛЬНО: Ответ должен быть ТОЛЬКО на русском языке."""
            return base_prompt + hint_instruction
        
        return base_prompt
    
    def on_search(self):
        """Запускает поиск"""
        query = self.search_input.text().strip()
        current_tab = self.tabs.currentIndex()
        
        # Проверка для локального поиска или стандартных AI промптов
        is_custom_ai = False
        if current_tab == 1 and self.prompt_combo.currentIndex() == 5:
            is_custom_ai = True
            
        if not query and not is_custom_ai:
            QMessageBox.warning(self, "Предупреждение", "Введите название компонента")
            return
            
        if is_custom_ai:
            # Для своего промпта проверяем, что он не пустой
            custom_text = self.custom_prompt_edit.toPlainText().strip()
            if not custom_text:
                QMessageBox.warning(self, "Предупреждение", "Введите текст вашего промпта")
                return
        
        if current_tab == 0:  # Локальный поиск
            self.run_local_search(query)
        elif self.unlocked and self.expert_mode:  # AI поиск - только для разблокированных экспертов
            self.run_ai_search(query)
        else:
            # Вкладка AI заблокирована
            QMessageBox.information(
                self,
                "AI поиск недоступен",
                "AI поиск доступен только в экспертном режиме после разблокировки приложения.\n\n"
                "Дважды кликните на имя разработчика для разблокировки."
            )
    
    def run_local_search(self, query: str):
        """Выполняет локальный поиск"""
        from .pdf_search import LocalPDFSearcher
        
        # Получаем список папок из интерфейса
        search_dirs = []
        for i in range(self.search_dirs_list.count()):
            item = self.search_dirs_list.item(i)
            path = item.data(Qt.UserRole)
            if path and os.path.exists(path):
                search_dirs.append(path)
        
        if not search_dirs:
            QMessageBox.warning(
                self,
                "Ошибка",
                "Список папок для поиска пуст!\n\n"
                "Нажмите '🔄 Сброс' для загрузки папок по умолчанию\n"
                "или '➕ Добавить папку' для выбора своей папки."
            )
            return
        
        # Выполняем поиск во всех директориях
        all_results = []
        for directory in search_dirs:
            searcher = LocalPDFSearcher(directory)
            results = searcher.search(query, min_match_length=3)
            all_results.extend(results)
        
        # Удаляем дубликаты по пути (если файл найден в нескольких директориях)
        seen_paths = set()
        unique_results = []
        for result in all_results:
            if result['path'] not in seen_paths:
                seen_paths.add(result['path'])
                unique_results.append(result)
        
        # Отображаем результаты
        self.local_results_list.clear()
        
        if not unique_results:
            item = QListWidgetItem(f"❌ Файлы не найдены в {len(search_dirs)} папках")
            item.setFlags(item.flags() & ~Qt.ItemIsEnabled)
            self.local_results_list.addItem(item)
        else:
            # Добавляем заголовок с количеством результатов
            header = QListWidgetItem(f"✅ Найдено {len(unique_results)} файлов в {len(search_dirs)} папках:")
            header.setFlags(header.flags() & ~Qt.ItemIsEnabled)
            header.setBackground(QColor("#313244"))
            self.local_results_list.addItem(header)
            
            for result in unique_results:
                item_text = f"📄 {result['filename']}\n   📁 {result['folder']} | 📊 {result['size']}"
                item = QListWidgetItem(item_text)
                item.setData(Qt.UserRole, result['path'])
                self.local_results_list.addItem(item)
    
    def run_ai_search(self, query: str):
        """Выполняет AI поиск"""
        from .pdf_search import AIPDFSearcher
        
        # Получаем API ключ из нового централизованного конфига
        api_keys = self.config.get("api_keys", {})
        telegram_security = self.config.get("telegram_security", {})
        provider = self.ai_provider_combo.currentText()
        api_key = None
        api_url = None
        use_encryption = False
        encryption_key = None
        app_id = telegram_security.get("app_id", "bomcategorizer-v5")
        
        if "Anthropic" in provider:
            api_key = api_keys.get("anthropic")
            provider_name = "anthropic"
        elif "Telegram" in provider:
            api_key = api_keys.get("telegram_key")
            api_url = api_keys.get("telegram_url")
            provider_name = "telegram_bot"
            # Получаем настройки шифрования для Telegram Bot
            use_encryption = api_keys.get("telegram_use_encryption", False)
            encryption_key = api_keys.get("telegram_enc_key")
        else:
            api_key = api_keys.get("openai")
            provider_name = "openai"
        
        if not api_key:
            QMessageBox.warning(
                self,
                "API ключ не найден",
                f"API ключ для {provider} не установлен.\n"
                "Откройте настройки и введите ваш API ключ."
            )
            return
        
        # Получаем промпт на основе выбранного типа
        custom_prompt = self._get_prompt_template(query)
        prompt_type = self.prompt_combo.currentText()
        
        # Проверяем наличие подсказки
        hint = ""
        if hasattr(self, 'hint_edit'):
            hint = self.hint_edit.toPlainText().strip()
        
        # Показываем индикатор загрузки с информацией о типе запроса
        encryption_info = " 🔒" if use_encryption else ""
        hint_info = f"<p style='color: #a6e3a1;'>💡 Подсказка: {hint[:50]}{'...' if len(hint) > 50 else ''}</p>" if hint else ""
        self.ai_results_browser.setHtml(
            f"<h3>⏳ Поиск...{encryption_info}</h3>"
            f"<p>Запрашиваем информацию у AI...</p>"
            f"<p style='color: #6c7086;'>Тип запроса: {prompt_type}</p>"
            f"<p style='color: #6c7086;'>Компонент: {query}</p>"
            f"{hint_info}"
        )
        
        # Запускаем поиск в отдельном потоке с кастомным промптом и шифрованием
        self.ai_worker = AISearchWorker(
            provider_name, api_key, query, api_url, custom_prompt,
            use_encryption=use_encryption, encryption_key=encryption_key,
            app_id=app_id
        )
        self.ai_worker.finished.connect(self.display_ai_results)
        self.ai_worker.start()
    
    def _open_external_link(self, url: QUrl):
        """Открывает ссылку во внешнем браузере"""
        QDesktopServices.openUrl(url)
    
    def display_ai_results(self, results: Dict):
        """Отображает результаты AI поиска"""
        if 'error' in results:
            html = f"""
            <h2 style="color: #f38ba8;">❌ Ошибка поиска</h2>
            <p><b>Компонент:</b> {results.get('component', 'N/A')}</p>
            <p><b>Ошибка:</b> {results['error']}</p>
            """
            if 'raw_response' in results:
                html += f"<h3>Сырой ответ:</h3><pre>{results['raw_response']}</pre>"
        else:
            html = self._format_ai_results_html(results)
        
        self.ai_results_browser.setHtml(html)
    
    def _format_ai_results_html(self, results: Dict) -> str:
        """Форматирует результаты AI в HTML"""
        if not results.get('found', False):
            return f"""
            <h2 style="color: #f9e2af;">⚠️ Компонент не найден</h2>
            <p><b>Запрос:</b> {results.get('component', 'N/A')}</p>
            <p>Информация о данном компоненте не найдена.</p>
            """
        
        # Базовые стили
        html = """
        <style>
            body { font-family: 'Segoe UI', Arial, sans-serif; line-height: 1.6; }
            h2 { color: #89b4fa; border-bottom: 2px solid #89b4fa; padding-bottom: 5px; }
            h3 { color: #a6e3a1; margin-top: 20px; }
            .spec-table { border-collapse: collapse; width: 100%; margin: 10px 0; }
            .spec-table td { padding: 8px; border: 1px solid #45475a; }
            .spec-table td:first-child { font-weight: bold; background-color: #313244; width: 30%; }
            .example { background-color: #1e1e2e; padding: 10px; margin: 5px 0; border-left: 3px solid #a6e3a1; }
            .datasheet-link { 
                display: inline-block;
                background-color: #89b4fa;
                color: #1e1e2e;
                padding: 10px 20px;
                text-decoration: none;
                border-radius: 5px;
                font-weight: bold;
                margin-top: 10px;
            }
            .datasheet-link:hover { background-color: #74c7ec; }
            .text-response { 
                background-color: #1e1e2e; 
                padding: 15px; 
                border-radius: 8px; 
                border-left: 4px solid #89b4fa;
                white-space: pre-wrap;
                font-size: 14px;
            }
            .text-response p { margin: 10px 0; }
            .text-response ul, .text-response ol { margin: 10px 0; padding-left: 20px; }
            .text-response li { margin: 5px 0; }
            a { color: #89b4fa; }
        </style>
        """
        
        # Проверяем, есть ли структурированные данные или только текст
        has_structured_data = any(key in results for key in ['full_name', 'manufacturer', 'type', 'specifications'])
        raw_response = results.get('raw_response', '')
        description = results.get('description', '')
        
        # Если есть только текстовый ответ (raw_response или description), показываем его красиво
        if not has_structured_data and (raw_response or description):
            text_content = raw_response or description
            
            # Преобразуем markdown-подобное форматирование в HTML
            formatted_text = self._format_markdown_to_html(text_content)
            
            # Информация о модели
            model_info = ""
            if results.get('model'):
                model_info = f"<p style='color: #6c7086; font-size: 11px;'>🤖 Модель: {results.get('model')} | Провайдер: {results.get('provider', 'N/A')}</p>"
            
            html += f"""
            <h2>📋 {results.get('component', 'Компонент')}</h2>
            {model_info}
            <div class="text-response">{formatted_text}</div>
            """
        else:
            # Стандартное структурированное отображение
            html += f"""
            <h2>📋 {results.get('full_name', results.get('component', 'Компонент'))}</h2>
            
            <table class="spec-table">
                <tr>
                    <td>🏭 Производитель</td>
                    <td>{results.get('manufacturer', 'N/A')}</td>
                </tr>
                <tr>
                    <td>🔧 Тип</td>
                    <td>{results.get('type', 'N/A')}</td>
                </tr>
            </table>
            
            <h3>📝 Описание</h3>
            <p>{description or 'Описание отсутствует'}</p>
            """
            
            # Характеристики
            specs = results.get('specifications', {})
            if specs:
                html += "<h3>⚙️ Основные характеристики</h3><table class='spec-table'>"
                for key, value in specs.items():
                    html += f"<tr><td>{key}</td><td>{value}</td></tr>"
                html += "</table>"
            
            # Примеры использования
            examples = results.get('examples', [])
            if examples:
                html += "<h3>💡 Примеры использования</h3>"
                for i, example in enumerate(examples, 1):
                    html += f"<div class='example'>{i}. {example}</div>"
            
            # Ссылка на datasheet
            datasheet_url = results.get('datasheet_url', '')
            if datasheet_url and datasheet_url.startswith('http'):
                html += f"""
                <h3>📄 Документация</h3>
                <a href="{datasheet_url}" class="datasheet-link" target="_blank">
                    📥 Скачать Datasheet (PDF)
                </a>
                """
        
        # Провайдер
        provider = results.get('provider', 'AI')
        html += f"<p style='margin-top: 30px; color: #6c7086; font-size: 0.9em;'>Информация предоставлена: {provider}</p>"
        
        return html
    
    def _format_markdown_to_html(self, text: str) -> str:
        """Преобразует markdown-подобный текст в HTML"""
        import re
        
        # Экранируем HTML
        text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        
        # Заголовки
        text = re.sub(r'^### (.+)$', r'<h4 style="color: #cba6f7;">\1</h4>', text, flags=re.MULTILINE)
        text = re.sub(r'^## (.+)$', r'<h3>\1</h3>', text, flags=re.MULTILINE)
        text = re.sub(r'^# (.+)$', r'<h2>\1</h2>', text, flags=re.MULTILINE)
        
        # Жирный и курсив
        text = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', text)
        text = re.sub(r'\*(.+?)\*', r'<em>\1</em>', text)
        
        # Ссылки
        text = re.sub(r'\[([^\]]+)\]\(([^\)]+)\)', r'<a href="\2" target="_blank">\1</a>', text)
        
        # URL без разметки
        text = re.sub(
            r'(https?://[^\s<>\"\)]+)',
            r'<a href="\1" target="_blank">\1</a>',
            text
        )
        
        # Списки с тире
        text = re.sub(r'^- (.+)$', r'<li>\1</li>', text, flags=re.MULTILINE)
        text = re.sub(r'(<li>.+</li>\n?)+', r'<ul>\g<0></ul>', text)
        
        # Нумерованные списки
        text = re.sub(r'^(\d+)\. (.+)$', r'<li>\2</li>', text, flags=re.MULTILINE)
        
        # Параграфы (двойные переносы)
        text = re.sub(r'\n\n+', '</p><p>', text)
        text = f'<p>{text}</p>'
        
        # Одиночные переносы строк
        text = text.replace('\n', '<br>')
        
        return text
    
    def _load_default_search_dirs(self):
        """Загружает папки для поиска по умолчанию"""
        from .pdf_search import get_default_pdf_directories
        
        self.search_dirs_list.clear()
        
        # Получаем пользовательские папки из конфига
        custom_dirs_from_config = self.config.get("pdf_search", {}).get("custom_directories", [])
        
        # Получаем все папки (включая пользовательские)
        all_dirs = get_default_pdf_directories(self.config)
        
        for directory in all_dirs:
            if os.path.exists(directory):
                # Проверяем, это пользовательская папка или системная
                is_custom = directory in custom_dirs_from_config
                
                # Добавляем иконку в зависимости от типа папки
                if is_custom:
                    icon = "👤"  # Пользовательская папка
                elif "pdf" in os.path.basename(directory).lower():
                    icon = "📄"
                elif "Project" in directory:
                    icon = "📁"
                elif "component_database" in directory or "BOMCategorizer" in directory:
                    icon = "💾"
                else:
                    icon = "📂"
                
                item_text = f"{icon} {directory}"
                item = QListWidgetItem(item_text)
                item.setData(Qt.UserRole, directory)
                item.setData(Qt.UserRole + 1, is_custom)  # Флаг: пользовательская папка
                
                tooltip = directory
                if is_custom:
                    tooltip += "\n(Пользовательская папка)"
                else:
                    tooltip += "\n(Системная папка по умолчанию)"
                item.setToolTip(tooltip)
                
                self.search_dirs_list.addItem(item)
        
        # Если список пустой, показываем предупреждение
        if self.search_dirs_list.count() == 0:
            item = QListWidgetItem("⚠️ Папки для поиска не найдены")
            item.setFlags(item.flags() & ~Qt.ItemIsEnabled)
            self.search_dirs_list.addItem(item)
    
    def add_search_directory(self):
        """Добавляет новую папку для поиска"""
        folder = QFileDialog.getExistingDirectory(
            self,
            "Выберите папку для поиска PDF",
            ""
        )
        
        if folder:
            # Проверяем, не добавлена ли уже эта папка
            for i in range(self.search_dirs_list.count()):
                item = self.search_dirs_list.item(i)
                existing_path = item.data(Qt.UserRole)
                if existing_path == folder:
                    QMessageBox.information(
                        self,
                        "Информация",
                        "Эта папка уже есть в списке!"
                    )
                    return
            
            # Добавляем новую пользовательскую папку
            item_text = f"👤 {folder}"
            item = QListWidgetItem(item_text)
            item.setData(Qt.UserRole, folder)
            item.setData(Qt.UserRole + 1, True)  # Помечаем как пользовательскую
            item.setToolTip(f"{folder}\n(Пользовательская папка - временно)")
            self.search_dirs_list.addItem(item)
    
    def remove_search_directory(self):
        """Удаляет выбранную папку из списка"""
        current_item = self.search_dirs_list.currentItem()
        if current_item:
            self.search_dirs_list.takeItem(self.search_dirs_list.row(current_item))
        else:
            QMessageBox.warning(
                self,
                "Предупреждение",
                "Выберите папку для удаления!"
            )
    
    def reset_search_directories(self):
        """Сбрасывает список папок к значениям по умолчанию"""
        reply = QMessageBox.question(
            self,
            "Подтверждение",
            "Вернуть список папок к значениям по умолчанию?\n\n"
            "Все временно добавленные папки будут удалены.",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No
        )
        if reply == QMessageBox.Yes:
            self._load_default_search_dirs()
    
    def save_search_dirs_to_config(self):
        """Сохраняет ТОЛЬКО пользовательские папки в config_qt.json"""
        # Собираем ТОЛЬКО пользовательские папки (с флагом is_custom = True)
        custom_dirs = []
        
        for i in range(self.search_dirs_list.count()):
            item = self.search_dirs_list.item(i)
            path = item.data(Qt.UserRole)
            is_custom = item.data(Qt.UserRole + 1)  # Флаг пользовательской папки
            
            if path and os.path.exists(path) and is_custom:
                # Нормализуем путь (убираем лишние слэши, приводим к абсолютному)
                normalized_path = os.path.normpath(os.path.abspath(path))
                custom_dirs.append(normalized_path)
        
        # Сохраняем файл - используем ту же логику, что и load_config()
        try:
            # Используем функцию get_config_path() из main_window
            from .main_window import get_config_path
            config_path = get_config_path()
            
            # Загружаем текущий конфиг из файла, чтобы сохранить все остальные настройки
            try:
                with open(config_path, 'r', encoding='utf-8') as f:
                    full_config = json.load(f)
            except FileNotFoundError:
                # Если файла нет, используем текущий конфиг
                full_config = self.config.copy()
            except Exception as e:
                # Если ошибка чтения, используем текущий конфиг
                print(f"Ошибка чтения конфига: {e}")
                full_config = self.config.copy()
            
            # Обновляем только секцию pdf_search
            if "pdf_search" not in full_config:
                full_config["pdf_search"] = {}
            full_config["pdf_search"]["custom_directories"] = custom_dirs
            
            # Создаем папку, если её нет
            config_dir = os.path.dirname(config_path)
            if config_dir and not os.path.exists(config_dir):
                os.makedirs(config_dir, exist_ok=True)
            
            # Сохраняем весь конфиг
            with open(config_path, 'w', encoding='utf-8') as f:
                json.dump(full_config, f, indent=2, ensure_ascii=False)
            
            # Формируем сообщение с путем к файлу
            if custom_dirs:
                msg = (f"✅ Сохранено {len(custom_dirs)} пользовательских папок в config_qt.json\n\n"
                       f"📁 Путь к файлу:\n{config_path}\n\n"
                       "Папки:\n" + "\n".join([f"  👤 {d}" for d in custom_dirs[:5]]) + 
                       (f"\n  ... и еще {len(custom_dirs) - 5}" if len(custom_dirs) > 5 else "") +
                       "\n\nСистемные папки (💾 📄 📁) не сохраняются - "
                       "они используются автоматически.")
            else:
                msg = (f"⚠️ Нет пользовательских папок для сохранения\n\n"
                       f"📁 Путь к файлу:\n{config_path}\n\n"
                       "Добавьте папки кнопкой ➕ - они будут помечены иконкой 👤\n"
                       "Системные папки (💾 📄 📁) сохранять не нужно.")
            
            QMessageBox.information(self, "Сохранено", msg)
            
            # Обновляем конфиг в памяти
            self.config = full_config
            
            # Обновляем конфиг в родительском окне
            if hasattr(self.parent_window, 'cfg'):
                self.parent_window.cfg = full_config
                self.parent_window.config = full_config  # Псевдоним для совместимости
            
            # Перезагружаем список, чтобы показать обновленные данные
            self._load_default_search_dirs()
            
        except Exception as e:
            QMessageBox.critical(
                self,
                "Ошибка",
                f"Не удалось сохранить конфигурацию:\n{str(e)}\n\n"
                f"Путь: {config_path if 'config_path' in locals() else 'не определен'}"
            )
    
    def open_local_file(self, item: QListWidgetItem):
        """Открывает PDF файл"""
        file_path = item.data(Qt.UserRole)
        if file_path and os.path.exists(file_path):
            self._open_file_in_system(file_path)
    
    def open_selected_local_file(self):
        """Открывает выбранный файл"""
        items = self.local_results_list.selectedItems()
        if items:
            self.open_local_file(items[0])
    
    def open_local_file_folder(self):
        """Открывает папку с выбранным файлом"""
        items = self.local_results_list.selectedItems()
        if not items:
            return
        
        file_path = items[0].data(Qt.UserRole)
        if file_path and os.path.exists(file_path):
            folder = os.path.dirname(file_path)
            self._open_file_in_system(folder)
    
    def open_search_directory(self, item: QListWidgetItem):
        """Открывает папку из списка поиска в проводнике/файловом менеджере"""
        directory = item.data(Qt.UserRole)
        if directory and os.path.exists(directory) and os.path.isdir(directory):
            self._open_file_in_system(directory)
        else:
            QMessageBox.warning(
                self,
                "Ошибка",
                f"Папка не найдена или недоступна:\n{directory}"
            )
    
    def _open_file_in_system(self, path: str):
        """Открывает файл или папку в системном приложении"""
        try:
            if platform.system() == 'Windows':
                os.startfile(path)
            elif platform.system() == 'Darwin':  # macOS
                subprocess.Popen(['open', path])
            else:  # Linux
                subprocess.Popen(['xdg-open', path])
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось открыть:\n{str(e)}")
    
    def save_ai_results(self):
        """Сохраняет результаты AI поиска"""
        # Проверка доступности (только для разблокированных экспертов)
        if not (self.unlocked and self.expert_mode):
            return
        
        if not hasattr(self, 'ai_results_browser'):
            return
        
        filters = (
            "HTML Files (*.html);;"
            "Text Files (*.txt);;"
            "Word Document (*.docx);;"
            "PDF (*.pdf)"
        )
        
        # Формируем безопасное имя файла (убираем спецсимволы)
        safe_name = "".join(c if c.isalnum() or c in "._- " else "_" for c in self.search_input.text())
        
        # Начальный путь — рабочий стол или домашняя директория
        desktop = os.path.join(os.path.expanduser("~"), "Desktop")
        if not os.path.exists(desktop):
            desktop = os.path.expanduser("~")
        
        default_path = os.path.join(desktop, f"ai_search_{safe_name}.html")
        
        file_path, selected_filter = QFileDialog.getSaveFileName(
            self,
            "Сохранить результаты",
            default_path,
            filters
        )
        
        if file_path:
            ext = os.path.splitext(file_path)[1].lower()
            
            if not ext:
                # Определяем расширение по выбранному фильтру
                if "Text Files" in selected_filter:
                    ext = ".txt"
                elif "Word Document" in selected_filter:
                    ext = ".docx"
                elif "PDF" in selected_filter:
                    ext = ".pdf"
                else:
                    ext = ".html"
                file_path += ext
            
            try:
                if ext == ".html":
                    content = self.ai_results_browser.toHtml()
                    with open(file_path, 'w', encoding='utf-8') as f:
                        f.write(content)
                elif ext == ".txt":
                    content = self.ai_results_browser.toPlainText()
                    with open(file_path, 'w', encoding='utf-8') as f:
                        f.write(content)
                elif ext == ".docx":
                    try:
                        from docx import Document
                    except ImportError:
                        QMessageBox.warning(
                            self,
                            "Отсутствует зависимость",
                            "Для сохранения в DOCX требуется пакет python-docx.\n"
                            "Установите его командой:\n\npip install python-docx"
                        )
                        return
                    
                    doc = Document()
                    doc.add_heading(f"AI поиск — {self.search_input.text()}", level=1)
                    for line in self.ai_results_browser.toPlainText().splitlines():
                        doc.add_paragraph(line if line else "")
                    doc.save(file_path)
                elif ext == ".pdf":
                    try:
                        from reportlab.lib.pagesizes import A4
                        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                        from reportlab.lib.styles import ParagraphStyle
                        from reportlab.lib.units import mm
                    except ImportError:
                        QMessageBox.warning(
                            self,
                            "Отсутствует зависимость",
                            "Для сохранения в PDF требуется пакет reportlab.\n"
                            "Установите его командой:\n\npip install reportlab"
                        )
                        return
                    
                    # Используем PDFExporter для правильной регистрации шрифтов
                    try:
                        from ..pdf_exporter import PDFExporter
                        pdf_exporter = PDFExporter()
                        font_name = pdf_exporter.cyrillic_font
                    except Exception as e:
                        print(f"Ошибка при инициализации PDFExporter: {e}")
                        font_name = 'Helvetica'
                    
                    # Создаём PDF документ
                    doc = SimpleDocTemplate(
                        file_path,
                        pagesize=A4,
                        leftMargin=15*mm,
                        rightMargin=15*mm,
                        topMargin=15*mm,
                        bottomMargin=15*mm
                    )
                    
                    # Стили
                    title_style = ParagraphStyle(
                        'Title',
                        fontName=font_name,
                        fontSize=12,
                        leading=14,
                        spaceAfter=10
                    )
                    
                    body_style = ParagraphStyle(
                        'Body',
                        fontName=font_name,
                        fontSize=9,
                        leading=11,
                        spaceAfter=3
                    )
                    
                    # Содержимое
                    story = []
                    
                    # Заголовок
                    title = f"AI поиск: {self.search_input.text()}"
                    story.append(Paragraph(title, title_style))
                    story.append(Spacer(1, 10))
                    
                    # Текст результатов
                    text = self.ai_results_browser.toPlainText()
                    for line in text.splitlines():
                        if line.strip():
                            # Экранируем HTML-специальные символы
                            safe_line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                            story.append(Paragraph(safe_line, body_style))
                        else:
                            story.append(Spacer(1, 6))
                    
                    doc.build(story)
                else:
                    QMessageBox.warning(self, "Неизвестный формат", f"Расширение {ext} не поддерживается.")
                    return
                
                QMessageBox.information(self, "Сохранено", f"Результаты сохранены:\n{file_path}")
            except Exception as e:
                QMessageBox.critical(self, "Ошибка сохранения", f"Не удалось сохранить файл:\n{e}")
    
    def open_settings(self):
        """Открывает настройки"""
        # Проверка доступности (настройки только для разблокированных экспертов)
        if not (self.unlocked and self.expert_mode):
            QMessageBox.information(
                self,
                "Настройки недоступны",
                "Настройки API доступны только в экспертном режиме после разблокировки приложения.\n\n"
                "Дважды кликните на имя разработчика для разблокировки."
            )
            return
        
        # Используем новое единое окно настроек напрямую
        dialog = UnifiedSettingsDialog(self, self.config)
        if dialog.exec() == QDialog.Accepted:
            # Обновляем конфиг
            self.config = dialog.get_config()
            # Сохраняем в родительском окне
            if hasattr(self.parent_window, 'save_pdf_search_config'):
                self.parent_window.save_pdf_search_config(self.config)
        
        # Force hide any lingering tooltips after dialog closes
        QToolTip.hideText()
        QTimer.singleShot(100, QToolTip.hideText)


class UnifiedSettingsDialog(QDialog):
    """Единое окно настроек с вкладками для API ключей и AI классификатора"""
    
    def __init__(self, parent, config: dict):
        super().__init__(parent)
        self.config = config.copy()
        self.parent_window = parent
        
        self.setWindowTitle("⚙️ Настройки API и AI")
        self.setModal(True)
        self.resize(730, 550)
        
        self._create_ui()
        self._load_settings()
    
    def _create_ui(self):
        """Создает интерфейс с вкладками"""
        layout = QVBoxLayout(self)
        
        self.tabs = QTabWidget()
        
        # Вкладка 1: Пути поиска PDF
        self.pdf_paths_tab = self._create_pdf_paths_tab()
        self.tabs.addTab(self.pdf_paths_tab, "📂 Пути PDF")
        
        # Вкладка 2: API ключи
        self.api_keys_tab = self._create_api_keys_tab()
        self.tabs.addTab(self.api_keys_tab, "🔑 API Ключи")
        
        # Вкладка 3: Настройки AI классификатора
        self.ai_classifier_tab = self._create_ai_classifier_tab()
        self.tabs.addTab(self.ai_classifier_tab, "🤖 AI Классификатор")
        
        layout.addWidget(self.tabs)
        
        # Кнопки
        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        buttons.accepted.connect(self._save_all_settings)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)

    def _create_pdf_paths_tab(self):
        """Создает вкладку управления путями поиска PDF"""
        tab = QWidget()
        layout = QVBoxLayout(tab)
        
        desc = QLabel(
            "📁 <b>Настройка пользовательских папок для поиска PDF</b><br><br>"
            "Добавьте свои папки, в которых будет выполняться рекурсивный поиск PDF файлов.<br>"
            "Эти папки будут использоваться дополнительно к стандартным (папка БД, Project/pdf* и т.д.)."
        )
        desc.setWordWrap(True)
        layout.addWidget(desc)
        
        # Список путей
        paths_group = QGroupBox("Пользовательские папки")
        paths_layout = QVBoxLayout()
        
        self.custom_paths_list = QListWidget()
        self.custom_paths_list.setMinimumHeight(200)
        paths_layout.addWidget(self.custom_paths_list)
        
        # Кнопки управления путями
        buttons_layout = QHBoxLayout()
        
        add_path_btn = QPushButton("➕ Добавить папку")
        add_path_btn.clicked.connect(self._add_custom_path)
        buttons_layout.addWidget(add_path_btn)
        
        remove_path_btn = QPushButton("➖ Удалить выбранную")
        remove_path_btn.clicked.connect(self._remove_custom_path)
        buttons_layout.addWidget(remove_path_btn)
        
        clear_paths_btn = QPushButton("🗑️ Очистить все")
        clear_paths_btn.clicked.connect(self._clear_custom_paths)
        buttons_layout.addWidget(clear_paths_btn)
        
        buttons_layout.addStretch()
        paths_layout.addLayout(buttons_layout)
        
        paths_group.setLayout(paths_layout)
        layout.addWidget(paths_group)
        
        # Подсказка
        hint_label = QLabel(
            "💡 <b>Совет:</b> Вы можете также вручную редактировать файл <code>config_qt.json</code><br>"
            "в разделе <code>\"pdf_search\" → \"custom_directories\"</code> для добавления путей."
        )
        hint_label.setWordWrap(True)
        layout.addWidget(hint_label)
        
        layout.addStretch()
        return tab
    
    def _add_custom_path(self):
        """Добавляет новую пользовательскую папку"""
        from PySide6.QtWidgets import QFileDialog
        
        folder = QFileDialog.getExistingDirectory(
            self,
            "Выберите папку для поиска PDF",
            "",
            QFileDialog.ShowDirsOnly | QFileDialog.DontResolveSymlinks
        )
        
        if folder:
            # Проверяем, не добавлена ли уже эта папка
            for i in range(self.custom_paths_list.count()):
                if self.custom_paths_list.item(i).text() == folder:
                    QMessageBox.information(
                        self,
                        "Информация",
                        "Эта папка уже добавлена в список!"
                    )
                    return
            
            self.custom_paths_list.addItem(folder)
    
    def _remove_custom_path(self):
        """Удаляет выбранную папку"""
        current_item = self.custom_paths_list.currentItem()
        if current_item:
            self.custom_paths_list.takeItem(self.custom_paths_list.row(current_item))
        else:
            QMessageBox.warning(
                self,
                "Предупреждение",
                "Выберите папку для удаления!"
            )
    
    def _clear_custom_paths(self):
        """Очищает весь список пользовательских папок"""
        if self.custom_paths_list.count() > 0:
            reply = QMessageBox.question(
                self,
                "Подтверждение",
                "Удалить все пользовательские папки из списка?",
                QMessageBox.Yes | QMessageBox.No,
                QMessageBox.No
            )
            if reply == QMessageBox.Yes:
                self.custom_paths_list.clear()

    def _create_api_keys_tab(self):
        """Создает единую вкладку для всех API ключей"""
        tab = QWidget()
        layout = QVBoxLayout(tab)
        
        api_group = QGroupBox("Ключи доступа для облачных сервисов")
        api_layout = QGridLayout()

        # Anthropic
        anthropic_label = QLabel("Anthropic Claude API Key:")
        self.anthropic_key_input = QLineEdit()
        self.anthropic_key_input.setEchoMode(QLineEdit.Password)
        self.anthropic_key_input.setPlaceholderText("sk-ant-...")
        
        show_anthropic_btn = QCheckBox("Показать")
        show_anthropic_btn.stateChanged.connect(
            lambda state: self.anthropic_key_input.setEchoMode(
                QLineEdit.Normal if state else QLineEdit.Password
            )
        )
        
        api_layout.addWidget(anthropic_label, 0, 0)
        api_layout.addWidget(self.anthropic_key_input, 0, 1)
        api_layout.addWidget(show_anthropic_btn, 0, 2)
        
        # OpenAI
        openai_label = QLabel("OpenAI GPT API Key:")
        self.openai_key_input = QLineEdit()
        self.openai_key_input.setEchoMode(QLineEdit.Password)
        self.openai_key_input.setPlaceholderText("sk-...")

        show_openai_btn = QCheckBox("Показать")
        show_openai_btn.stateChanged.connect(
            lambda state: self.openai_key_input.setEchoMode(
                QLineEdit.Normal if state else QLineEdit.Password
            )
        )
        
        api_layout.addWidget(openai_label, 1, 0)
        api_layout.addWidget(self.openai_key_input, 1, 1)
        api_layout.addWidget(show_openai_btn, 1, 2)
        
        api_group.setLayout(api_layout)
        layout.addWidget(api_group)

        # Ollama
        ollama_group = QGroupBox("Настройки для локальных моделей (Ollama)")
        ollama_layout = QGridLayout()

        ollama_label = QLabel("Ollama URL:")
        self.ollama_url_input = QLineEdit()
        self.ollama_url_input.setPlaceholderText("http://localhost:11434")

        ollama_layout.addWidget(ollama_label, 0, 0)
        ollama_layout.addWidget(self.ollama_url_input, 0, 1)

        ollama_group.setLayout(ollama_layout)
        layout.addWidget(ollama_group)

        # Telegram Bot
        telegram_group = QGroupBox("Настройки Telegram Bot API")
        telegram_layout = QGridLayout()

        telegram_url_label = QLabel("Bot API URL:")
        self.telegram_url_input = QLineEdit()
        self.telegram_url_input.setPlaceholderText("http://localhost:8000/ai_query")
        self.telegram_url_input.textChanged.connect(self._on_telegram_url_changed)
        
        telegram_port_label = QLabel("Порт:")
        self.telegram_port_input = QSpinBox()
        self.telegram_port_input.setRange(1, 65535)
        self.telegram_port_input.setValue(8000)
        self.telegram_port_input.valueChanged.connect(self._on_telegram_port_changed)
        
        # Checkbox for encryption (requested by user)
        self.use_encryption_cb = QCheckBox("Шифрование")
        self.use_encryption_cb.setToolTip("Если выключено, используется обычный HTTP без шифрования (небезопасно)")
        self.use_encryption_cb.setChecked(True) # Default to True
        self.use_encryption_cb.stateChanged.connect(self._on_encryption_toggled)
        
        telegram_key_label = QLabel("Bot API Key:")
        self.telegram_key_input = QLineEdit()
        self.telegram_key_input.setEchoMode(QLineEdit.Password)
        self.telegram_key_input.setPlaceholderText("secret_key")
        
        show_telegram_btn = QCheckBox("Показать")
        show_telegram_btn.stateChanged.connect(
            lambda state: self.telegram_key_input.setEchoMode(
                QLineEdit.Normal if state else QLineEdit.Password
            )
        )

        telegram_enc_label = QLabel("Encryption Key:")
        self.telegram_enc_input = QLineEdit()
        self.telegram_enc_input.setEchoMode(QLineEdit.Password)
        self.telegram_enc_input.setPlaceholderText("32-byte hex key")
        
        show_enc_btn = QCheckBox("Показать")
        show_enc_btn.stateChanged.connect(
            lambda state: self.telegram_enc_input.setEchoMode(
                QLineEdit.Normal if state else QLineEdit.Password
            )
        )

        telegram_layout.addWidget(telegram_url_label, 0, 0)
        telegram_layout.addWidget(self.telegram_url_input, 0, 1)
        telegram_layout.addWidget(telegram_port_label, 0, 2)
        telegram_layout.addWidget(self.telegram_port_input, 0, 3)
        telegram_layout.addWidget(self.use_encryption_cb, 0, 4) # Added checkbox here
        
        telegram_layout.addWidget(telegram_key_label, 1, 0)
        telegram_layout.addWidget(self.telegram_key_input, 1, 1, 1, 3) # Span across columns
        telegram_layout.addWidget(show_telegram_btn, 1, 4)

        telegram_layout.addWidget(telegram_enc_label, 2, 0)
        telegram_layout.addWidget(self.telegram_enc_input, 2, 1, 1, 3) # Span across columns
        telegram_layout.addWidget(show_enc_btn, 2, 4)
        
        # Test Connection Button
        test_conn_btn = QPushButton("🔄 Проверить соединение")
        test_conn_btn.setToolTip("Отправить тестовый запрос для проверки связи и шифрования")
        test_conn_btn.clicked.connect(self._test_connection)
        telegram_layout.addWidget(test_conn_btn, 3, 0, 1, 5) # Span full width

        telegram_group.setLayout(telegram_layout)
        layout.addWidget(telegram_group)
        
        # Помощь
        help_label = QLabel(
            "💡 <b>Как получить API ключи:</b><br>"
            "• <b>Anthropic:</b> <a href='https://console.anthropic.com/'>console.anthropic.com</a><br>"
            "• <b>OpenAI:</b> <a href='https://platform.openai.com/api-keys'>platform.openai.com/api-keys</a><br>"
            "• <b>Ollama:</b> <a href='https://ollama.ai/'>ollama.ai</a> (для локального запуска)<br>"
            "• <b>Telegram Bot:</b> команда <code>/api</code> в боте (только для админа)"
        )
        help_label.setOpenExternalLinks(True)
        help_label.setWordWrap(True)
        layout.addWidget(help_label)
        
        layout.addStretch()
        return tab

    def _create_ai_classifier_tab(self):
        """Создает вкладку настроек AI классификатора (без ключей)"""
        tab = QWidget()
        layout = QVBoxLayout(tab)
        
        desc = QLabel(
            "Настройте параметры для автоматической классификации компонентов.\n"
            "API ключи настраиваются на соседней вкладке 'API Ключи'."
        )
        desc.setWordWrap(True)
        layout.addWidget(desc)
        
        form_group = QGroupBox("Параметры классификатора")
        form = QFormLayout()
        
        # Провайдер
        self.provider_combo = QComboBox()
        self.provider_combo.addItems(["Anthropic Claude", "OpenAI GPT", "Ollama (локальный)"])
        form.addRow("Провайдер AI:", self.provider_combo)
        
        # Модель
        self.ai_model_input = QLineEdit()
        self.ai_model_input.setPlaceholderText("По умолчанию (оставьте пустым)")
        form.addRow("Модель (опционально):", self.ai_model_input)
        
        # Порог уверенности
        self.ai_confidence_combo = QComboBox()
        self.ai_confidence_combo.addItems(["Высокий (high)", "Средний (medium)", "Низкий (low)"])
        form.addRow("Порог уверенности:", self.ai_confidence_combo)

        form_group.setLayout(form)
        layout.addWidget(form_group)
        
        # Справка по моделям
        help_text = QTextBrowser()
        help_text.setReadOnly(True)
        help_text.setMaximumHeight(100)
        help_text.setOpenExternalLinks(True)
        help_text.setHtml("""
<b>Модели по умолчанию:</b><br>
• Anthropic: <code>claude-3-sonnet-20240229</code><br>
• OpenAI: <code>gpt-4</code><br>
• Ollama: <code>llama2</code>
        """)
        layout.addWidget(help_text)

        layout.addStretch()
        return tab
    
    def _on_telegram_url_changed(self, text: str):
        """Обработчик изменения URL: обновляет поле порта"""
        from urllib.parse import urlparse
        
        # Блокируем сигналы, чтобы не вызвать зацикливание
        self.telegram_port_input.blockSignals(True)
        try:
            # Пытаемся распарсить URL
            if not text.startswith(('http://', 'https://')):
                # Если нет схемы, urlparse может работать некорректно для наших целей
                # Но мы пока просто игнорируем
                pass
            else:
                parsed = urlparse(text)
                if parsed.port:
                    self.telegram_port_input.setValue(parsed.port)
                else:
                    # Если порта нет явно, ставим дефолтный для схемы
                    if parsed.scheme == 'https':
                        self.telegram_port_input.setValue(443)
                    elif parsed.scheme == 'http':
                        self.telegram_port_input.setValue(80)
        except Exception:
            pass
        finally:
            self.telegram_port_input.blockSignals(False)

    def _on_telegram_port_changed(self, port: int):
        """Обработчик изменения порта: обновляет URL"""
        from urllib.parse import urlparse, urlunparse
        
        text = self.telegram_url_input.text()
        if not text:
            return
            
        self.telegram_url_input.blockSignals(True)
        try:
            from urllib.parse import urlparse, urlunparse
            parsed = urlparse(text)
            # Reconstruct with new port
            netloc_parts = parsed.netloc.split(':')
            host = netloc_parts[0]
            new_netloc = f"{host}:{port}"
            
            new_parsed = parsed._replace(netloc=new_netloc)
            new_url = urlunparse(new_parsed)
            
            self.telegram_url_input.setText(new_url)
        except Exception as e:
            print(f"Error updating URL port: {e}")
        finally:
            self.telegram_url_input.blockSignals(False)

    def _on_encryption_toggled(self, state):
        """Обработчик переключения шифрования"""
        from PySide6.QtCore import Qt
        is_checked = (state == Qt.Checked)
        self.telegram_enc_input.setEnabled(is_checked)
        # Можно добавить логику изменения URL (secure/plain), но это может быть сложно,
        # так как пользователь может редактировать URL вручную.
        # Лучше оставить URL как есть, а логику выбора endpoint оставить в ai_classifier.py
    
    def _test_connection(self):
        """Проверяет соединение с TelegramHelper API"""
        url = self.telegram_url_input.text().strip()
        api_key = self.telegram_key_input.text().strip()
        enc_key = self.telegram_enc_input.text().strip()
        use_encryption = self.use_encryption_cb.isChecked()
        
        if not url:
            QMessageBox.warning(self, "Ошибка", "URL не может быть пустым")
            return
            
        # Визуальная индикация
        sender = self.sender()
        original_text = sender.text()
        sender.setText("⏳ Проверка...")
        sender.setEnabled(False)
        QApplication.processEvents()
        
        try:
            import requests
            import json
            import base64
            
            # Тестовые данные
            test_payload = {
                "prompt": "Test connection",
                "provider": "anthropic", # Используем легкий запрос
                "max_tokens": 10
            }
            
            headers = {
                "Content-Type": "application/json",
                "X-API-KEY": api_key,
                "X-APP-ID": "bomcategorizer-v5"
            }
            
            # Подготовка данных (шифрование если нужно)
            if use_encryption:
                if not enc_key:
                    raise ValueError("Ключ шифрования обязателен при включенном шифровании")
                
                try:
                    # Импортируем SecureMessenger из локального модуля
                    # Предполагаем что он доступен в bom_categorizer.encryption
                    from ..encryption import SecureMessenger
                    messenger = SecureMessenger(enc_key)
                    
                    encrypted_bytes = messenger.encrypt(test_payload)
                    b64_data = base64.b64encode(encrypted_bytes).decode('utf-8')
                    
                    json_data = {"data": b64_data}
                except Exception as e:
                    raise Exception(f"Ошибка шифрования: {e}")
            else:
                json_data = test_payload
            
            # Отправка запроса
            try:
                response = requests.post(url, json=json_data, headers=headers, timeout=10)
                response.raise_for_status()
                result = response.json()
                
                # Проверка ответа
                success_msg = "✅ Соединение успешно установлено!\n\n"
                
                if use_encryption:
                    if result.get("mode") != "encrypted":
                        success_msg += "⚠️ Внимание: Сервер ответил без шифрования!\n"
                    
                    if "data" in result:
                        try:
                            # Пробуем расшифровать ответ
                            encrypted_response = base64.b64decode(result["data"])
                            decrypted = messenger.decrypt(encrypted_response)
                            decrypted_json = json.loads(decrypted.decode('utf-8'))
                            success_msg += "🔐 Шифрование работает корректно (запрос и ответ).\n"
                            success_msg += f"Ответ сервера: {decrypted_json.get('status', 'OK')}"
                        except Exception as e:
                            success_msg += f"❌ Ошибка расшифровки ответа: {e}"
                    else:
                        success_msg += "❌ Ответ сервера не содержит зашифрованных данных"
                else:
                    success_msg += "📡 Обычное соединение (без шифрования) работает."
                
                QMessageBox.information(self, "Успех", success_msg)
                
            except requests.exceptions.HTTPError as e:
                status_code = e.response.status_code if e.response else "N/A"
                detail = e.response.text if e.response else str(e)
                QMessageBox.critical(self, "Ошибка API", f"Сервер вернул ошибку {status_code}:\n{detail}")
            except requests.exceptions.ConnectionError:
                QMessageBox.critical(self, "Ошибка сети", f"Не удалось подключиться к серверу.\nПроверьте URL и доступность сервера.")
            except requests.exceptions.Timeout:
                QMessageBox.critical(self, "Тайм-аут", f"Сервер не ответил вовремя (10 сек).")
            except Exception as e:
                QMessageBox.critical(self, "Ошибка запроса", f"Произошла ошибка при отправке:\n{e}")
                
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось выполнить проверку:\n{e}")
        finally:
            sender.setText(original_text)
            sender.setEnabled(True)
    
    def _load_settings(self):
        """Загружает настройки из config_qt.json"""
        # --- 0. Загрузка пользовательских путей PDF ---
        pdf_search_conf = self.config.get("pdf_search", {})
        custom_dirs = pdf_search_conf.get("custom_directories", [])
        self.custom_paths_list.clear()
        for path in custom_dirs:
            if path:  # Пропускаем пустые строки
                self.custom_paths_list.addItem(path)
        
        # --- 1. Загрузка API ключей ---
        # Сначала из новой централизованной секции
        api_keys = self.config.get("api_keys", {})
        
        # Для обратной совместимости, ищем в старых секциях, если в новой пусто
        ai_classifier_conf = self.config.get("ai_classifier", {})
        ai_api_keys = ai_classifier_conf.get("api_keys", {})
        
        # Anthropic
        anthropic_key = api_keys.get("anthropic") or \
                        pdf_search_conf.get("anthropic_api_key") or \
                        ai_api_keys.get("anthropic", "")
        self.anthropic_key_input.setText(anthropic_key)
        
        # OpenAI
        openai_key = api_keys.get("openai") or \
                     pdf_search_conf.get("openai_api_key") or \
                     ai_api_keys.get("openai", "")
        self.openai_key_input.setText(openai_key)
        
        # Ollama
        ollama_url = api_keys.get("ollama_url") or \
                     ai_api_keys.get("ollama") or \
                     "http://localhost:11434"
        self.ollama_url_input.setText(ollama_url)
        
        # Telegram Bot
        # Telegram settings
        # Use self.config for consistency, not self.full_config
        telegram_settings = self.config.get("ai_classifier", {}) 
        # api_keys is already defined at the start of _load_settings
        
        self.telegram_url_input.setText(api_keys.get("telegram_url", ""))
        self.telegram_key_input.setText(api_keys.get("telegram_key", ""))
        self.telegram_enc_input.setText(api_keys.get("telegram_enc_key", ""))
        
        # Load use_encryption setting (default True)
        use_encryption = api_keys.get("telegram_use_encryption", True)
        self.use_encryption_cb.setChecked(use_encryption)
        self.telegram_enc_input.setEnabled(use_encryption)
        
        # Try to extract port from URL
        url = api_keys.get("telegram_url", "") or ""
        # Initialize port from URL after setting the URL input
        self._on_telegram_url_changed(url)
        
        # --- 2. Загрузка настроек AI Классификатора ---
        settings = ai_classifier_conf # Используем уже загруженный конфиг
        
        provider_map = {"anthropic": 0, "openai": 1, "ollama": 2}
        self.provider_combo.setCurrentIndex(provider_map.get(settings.get("provider"), 0))
        
        self.ai_model_input.setText(settings.get("model", ""))
        
        confidence_map = {"high": 0, "medium": 1, "low": 2}
        self.ai_confidence_combo.setCurrentIndex(confidence_map.get(settings.get("confidence_threshold"), 1))

    def _save_all_settings(self):
        """Сохраняет все настройки в config_qt.json"""
        # --- 0. Сохраняем пользовательские пути PDF ---
        custom_dirs = []
        for i in range(self.custom_paths_list.count()):
            path = self.custom_paths_list.item(i).text()
            if path:
                custom_dirs.append(path)
        
        if "pdf_search" not in self.config:
            self.config["pdf_search"] = {}
        self.config["pdf_search"]["custom_directories"] = custom_dirs
        
        # --- 1. Сохраняем API ключи в централизованную секцию ---
        # Ensure api_keys section exists
        if "api_keys" not in self.config:
            self.config["api_keys"] = {}
            
        self.config["api_keys"]["anthropic"] = self.anthropic_key_input.text().strip()
        self.config["api_keys"]["openai"] = self.openai_key_input.text().strip()
        self.config["api_keys"]["ollama_url"] = self.ollama_url_input.text().strip()
        
        # Save Telegram settings to api_keys
        self.config["api_keys"]["telegram_url"] = self.telegram_url_input.text().strip()
        self.config["api_keys"]["telegram_key"] = self.telegram_key_input.text().strip()
        self.config["api_keys"]["telegram_enc_key"] = self.telegram_enc_input.text().strip()
        self.config["api_keys"]["telegram_use_encryption"] = self.use_encryption_cb.isChecked()
        
        # --- 2. Сохраняем настройки AI классификатора ---
        # Удаляем старые ключи из секции pdf_search для очистки
        if "pdf_search" in self.config:
            self.config["pdf_search"].pop("anthropic_api_key", None)
            self.config["pdf_search"].pop("openai_api_key", None)

        ai_provider_map = {0: "anthropic", 1: "openai", 2: "ollama"}
        ai_confidence_map = {0: "high", 1: "medium", 2: "low"}
        
        ai_settings = {
            "enabled": self.config.get("ai_classifier", {}).get("enabled", False),
            "provider": ai_provider_map[self.provider_combo.currentIndex()],
            "model": self.ai_model_input.text().strip(),
            "auto_classify": self.config.get("ai_classifier", {}).get("auto_classify", False),
            "confidence_threshold": ai_confidence_map[self.ai_confidence_combo.currentIndex()],
            # ВАЖНО: секция api_keys здесь больше не нужна, т.к. они хранятся централизованно
        }
        self.config["ai_classifier"] = ai_settings
        
        # --- 3. Сохраняем весь файл config_qt.json ---
        try:
            # __file__ = bom_categorizer/gui/pdf_search_dialogs.py
            # Нужно 3 уровня вверх: gui -> bom_categorizer -> корень проекта
            project_root = os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
            config_path = os.path.join(project_root, "config_qt.json")
            with open(config_path, 'w', encoding='utf-8') as f:
                json.dump(self.config, f, indent=2, ensure_ascii=False)
            
            if hasattr(self.parent_window, 'log_text') and self.parent_window.log_text:
                self.parent_window.log_text.append("✅ Настройки API и AI сохранены")
            
            if hasattr(self.parent_window, 'update_ai_status'):
                self.parent_window.update_ai_status()
            
            # Скрываем подсказки перед закрытием
            focus_widget = QApplication.focusWidget()
            if focus_widget:
                focus_widget.clearFocus()
            
            # Используем таймер для отложенного скрытия, чтобы гарантировать очистку после закрытия окна
            QToolTip.hideText() # Скрываем сразу
            QTimer.singleShot(100, QToolTip.hideText) # Скрываем через 100мс
            QTimer.singleShot(500, QToolTip.hideText) # Второй контрольный выстрел
            
            self.accept()
        except Exception as e:
            QMessageBox.warning(self, "Ошибка", f"Не удалось сохранить настройки: {e}")

    def closeEvent(self, event):
        """Обработчик закрытия окна"""
        # Скрываем любые активные подсказки с задержкой
        QToolTip.hideText()
        QTimer.singleShot(100, QToolTip.hideText)
        event.accept()

    def get_config(self) -> dict:
        """Возвращает обновленный конфиг"""
        return self.config


class PDFSearchSettingsDialog(QDialog):
    """Диалог настроек поиска PDF (устаревший, используйте UnifiedSettingsDialog)"""
    
    def __init__(self, parent, config: dict):
        super().__init__(parent)
        # Перенаправляем на единое окно настроек
        unified_dialog = UnifiedSettingsDialog(parent, config)
        result = unified_dialog.exec()
        # Для совместимости возвращаем конфиг
        self.config = unified_dialog.get_config() if result == QDialog.Accepted else config
        # Устанавливаем результат для этого диалога
        if result == QDialog.Accepted:
            self.accept()
        else:
            self.reject()
    
    def get_config(self) -> dict:
        """Возвращает обновленный конфиг"""
        return self.config


class AISearchWorker(QThread):
    """Worker для AI поиска в отдельном потоке"""
    finished = Signal(dict)
    
    def __init__(self, provider: str, api_key: str, query: str, api_url: str = None, 
                 custom_prompt: str = None, use_encryption: bool = False, encryption_key: str = None,
                 app_id: str = "bomcategorizer-v5"):
        super().__init__()
        self.provider = provider
        self.api_key = api_key
        self.query = query
        self.api_url = api_url
        self.custom_prompt = custom_prompt
        self.use_encryption = use_encryption
        self.encryption_key = encryption_key
        self.app_id = app_id
    
    def run(self):
        """Выполняет AI поиск"""
        from .pdf_search import AIPDFSearcher
        
        searcher = AIPDFSearcher(
            self.provider, 
            self.api_key, 
            self.api_url,
            use_encryption=self.use_encryption,
            encryption_key=self.encryption_key,
            app_id=self.app_id
        )
        
        # Используем кастомный промпт если передан
        if self.custom_prompt:
            results = searcher.search_with_prompt(self.query, self.custom_prompt)
        else:
            results = searcher.search(self.query)
        
        self.finished.emit(results)

