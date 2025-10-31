# -*- coding: utf-8 -*-
"""
Парсеры для различных форматов BOM файлов

Поддерживаемые форматы:
- TXT: текстовые файлы с разделителями
- DOCX: документы Word с таблицами
- Excel: XLSX файлы
"""

import os
import re
from typing import List, Dict, Any, Optional
import pandas as pd

from .utils import LINE_SPLIT_RE, POS_PREFIX_RE

try:
    from docx import Document  # python-docx
except Exception:
    Document = None  # optional; raise if used without installed


def normalize_dashes(text: str) -> str:
    """
    Нормализует различные виды тире и дефисов к обычному дефису
    
    При конвертации .doc → .docx Word может заменять дефисы на типографские тире.
    Это функция приводит все варианты к единому формату для правильного объединения компонентов.
    
    Args:
        text: Исходный текст
        
    Returns:
        Текст с нормализованными дефисами
    """
    if not text:
        return text
    
    # Заменяем все виды тире на обычный дефис
    # U+2013: EN DASH (–)
    # U+2014: EM DASH (—)
    # U+2212: MINUS SIGN (−)
    # U+2010: HYPHEN (‐)
    # U+2011: NON-BREAKING HYPHEN (‑)
    text = text.replace('\u2013', '-')  # EN DASH
    text = text.replace('\u2014', '-')  # EM DASH
    text = text.replace('\u2212', '-')  # MINUS SIGN
    text = text.replace('\u2010', '-')  # HYPHEN
    text = text.replace('\u2011', '-')  # NON-BREAKING HYPHEN
    
    return text


def normalize_cell(s: Any) -> str:
    """Нормализует содержимое ячейки таблицы"""
    text = str(s or "").strip()
    # Нормализуем тире для корректного объединения компонентов
    text = normalize_dashes(text)
    # Удаляем непечатные символы (включая �)
    text = ''.join(char for char in text if char.isprintable() or char in '\n\r\t')
    return text


def count_from_reference(ref: str) -> int:
    """
    Подсчитывает количество элементов из позиционного обозначения
    
    Examples:
        R1 -> 1
        R1, R2 -> 2
        R1-R6 -> 6
        FU1-FU6 -> 6
        C1, C2, C3-C5 -> 5
    
    Args:
        ref: Позиционное обозначение
        
    Returns:
        Количество элементов
    """
    if not ref or not ref.strip():
        return 1
    
    ref = ref.strip()
    total = 0
    
    # Разделяем по запятым
    parts = [p.strip() for p in ref.split(',')]
    
    for part in parts:
        # Проверяем на диапазон (R1-R6, FU1-FU6, и т.д.)
        range_match = re.match(r'([A-Za-z]+)(\d+)\s*[-–—]\s*([A-Za-z]+)?(\d+)', part)
        if range_match:
            prefix1 = range_match.group(1)
            num1 = int(range_match.group(2))
            prefix2 = range_match.group(3) or prefix1  # Если второй префикс отсутствует, используем первый
            num2 = int(range_match.group(4))
            
            # Подсчитываем диапазон
            if prefix1 == prefix2 and num2 >= num1:
                total += (num2 - num1 + 1)
            else:
                # Если префиксы разные или порядок неправильный, считаем как 2 элемента
                total += 2
        else:
            # Одиночный элемент
            total += 1
    
    return max(total, 1)


def parse_txt_like(path: str) -> pd.DataFrame:
    """
    Парсит текстовый файл с разделителями
    
    Args:
        path: Путь к TXT файлу
        
    Returns:
        DataFrame с колонками: reference, description, qty
    """
    rows: List[Dict[str, Any]] = []
    try:
        with open(path, "r", encoding="utf-8") as f:
            text = f.read()
    except UnicodeDecodeError:
        with open(path, "r", encoding="cp1251", errors="ignore") as f:
            text = f.read()
    
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        parts = [p.strip() for p in LINE_SPLIT_RE.split(line) if p.strip()]
        if not parts:
            continue
        
        # Heuristic mapping: [pos?, name/desc, qty?]
        pos = parts[0] if POS_PREFIX_RE.match(parts[0]) else None
        qty = None
        
        # Try explicit "Nшт" or "N шт" patterns first
        m = re.search(r"(\d+)\s*(шт\.?|pcs|pieces)", line, flags=re.IGNORECASE)
        if m:
            qty = int(m.group(1))
        else:
            for p in parts[::-1]:
                if re.fullmatch(r"\d+", p):
                    qty = int(p)
                    break
        
        desc = " ".join(parts[1:-1]) if pos and len(parts) >= 2 else (" ".join(parts))
        row = {"reference": pos or "", "description": desc, "qty": qty if qty is not None else 1}
        rows.append(row)
    
    if not rows:
        # fallback: whole text in a single row
        rows = [{"description": text, "qty": 1}]
    
    return pd.DataFrame(rows)


def parse_docx(path: str) -> pd.DataFrame:
    """
    Парсит DOCX файл с таблицами
    
    Args:
        path: Путь к DOCX файлу
        
    Returns:
        DataFrame с колонками: zone, reference, description, qty, note
    """
    if Document is None:
        raise SystemExit("python-docx is required to parse DOCX. Install with pip install python-docx")
    
    doc = Document(path)
    extracted: List[Dict[str, Any]] = []

    def guess_header_index(table) -> int:
        """Определяет индекс строки-заголовка в таблице"""
        max_scan = min(5, len(table.rows))
        header_keywords = ["наимен", "обознач", "кол.", "кол ", "кол", "примеч"]
        for i in range(max_scan):
            row_texts = [normalize_cell(c.text).lower() for c in table.rows[i].cells]
            hits = sum(any(hk in t for hk in header_keywords) for t in row_texts)
            if hits >= 2:
                return i
        return 0

    # Parse tables with header detection
    for table in doc.tables:
        if not table.rows:
            continue
        
        header_idx = guess_header_index(table)
        header_cells = [normalize_cell(c.text).strip() for c in table.rows[header_idx].cells]
        header_norm = [h.lower() for h in header_cells]

        # Find column indices by common names
        def find_col_idx(candidates: List[str]) -> Optional[int]:
            for i, h in enumerate(header_norm):
                for cand in candidates:
                    if cand in h:
                        return i
            return None

        idx_zone = find_col_idx(["зона"])  # optional
        idx_ref = find_col_idx(["поз", "обозн"])  # позиционное обозначение
        idx_name = find_col_idx(["наимен"])  # наименование
        idx_qty = find_col_idx(["кол.", "кол ", "кол", "количество"])  # количество
        idx_note = find_col_idx(["примеч"])  # опционально

        # Переменные для хранения текущего ТУ и типа компонента из заголовка группы
        current_group_tu = ""
        current_group_type = ""
        # Словарь для хранения ТУ для разных типов компонентов (например, К10, К53, GRM)
        # Это нужно для конденсаторов/резисторов с несколькими групповыми заголовками
        component_type_tu_map = {}

        for tr in table.rows[header_idx + 1:]:
            vals = [normalize_cell(c.text) for c in tr.cells]
            if not any(v.strip() for v in vals):
                continue
            
            zone = vals[idx_zone] if idx_zone is not None and idx_zone < len(vals) else ""
            ref = vals[idx_ref] if idx_ref is not None and idx_ref < len(vals) else ""
            name = vals[idx_name] if idx_name is not None and idx_name < len(vals) else ""
            qty_raw = vals[idx_qty] if idx_qty is not None and idx_qty < len(vals) else ""
            cell_note = vals[idx_note] if idx_note is not None and idx_note < len(vals) else ""
            original_note = cell_note  # Сохранить оригинальное примечание для извлечения подборов
            note = cell_note  # Инициализируем note содержимым ячейки
            
            # Проверить, является ли это строкой-заголовком группы
            section_headers = [
                'конденсаторы', 'конденсаторов', 'резисторы', 'резисторов', 
                'микросхемы', 'микросхем', 'дроссели', 'дросселей', 
                'индуктивности', 'индуктивностей',
                'разъемы', 'разъемов', 'диоды', 'диодов', 
                'транзисторы', 'транзисторов', 'кабели', 'кабелей', 
                'модули', 'модулей',
                'набор резисторов', 'набор конденсаторов', 'набор микросхем',
                'трансформаторы', 'трансформаторов', 'датчики', 'датчиков', 
                'реле', 'предохранители', 'предохранителей',
                'оптопары', 'оптопар', 'оптроны', 'оптронов', 
                'светодиоды', 'светодиодов', 'стабилитроны', 'стабилитронов',
                'вариаторы', 'вариаторов', 'переключатели', 'переключателей', 
                'кнопки', 'кнопок', 'тумблеры', 'тумблеров',
                'фильтры', 'фильтров', 'антенны', 'антенн', 
                'радиаторы', 'радиаторов', 'крепеж', 'крепежа',
                'провода', 'проводов', 'жгуты', 'жгутов', 
                'шлейфы', 'шлейфов', 'платы', 'плат',
                'корпуса', 'корпусов', 'панели', 'панелей', 
                'винты', 'винтов', 'гайки', 'гаек',
                'изделия', 'изделий', 'детали', 'деталей', 'прочие элементы'
            ]
            name_lower = name.strip().lower()
            
            is_group_header = False
            
            # Проверка группового заголовка: должно быть БЕЗ ref И содержать тип компонента
            # qty может быть любым (иногда в docx заголовки имеют qty=1)
            if not ref.strip():
                # Проверяем наличие типа компонента В ЛЮБОЙ ЧАСТИ строки (не только в начале!)
                # Это важно для заголовков типа "Чип катушки индуктивности", "Набор резисторов"
                if any(section in name_lower for section in section_headers):
                    # Дополнительная проверка: если есть ТУ-код, это точно заголовок
                    # Если нет ТУ и есть qty, это может быть компонент
                    # ВАЖНО: между номером и "ТУ" может быть пробел (например, "ШКАБ.434110.018 ТУ")
                    has_tu_in_name = bool(re.search(r'[А-ЯЁ]{2,}[\.\d]+\s*ТУ', name) or re.search(r'ТУ\s+[\d\-]+', name))
                    if has_tu_in_name or not qty_raw.strip():
                        is_group_header = True
                # ИЛИ строка содержит ТОЛЬКО обозначение типа (К10-, К53-, Р1-, и т.д.) + ТУ-код
                # Паттерн: буквы+цифры (тип) + пробелы + ТУ-код, БЕЗ детального описания
                elif re.match(r'^[А-ЯЁ]+[\d\-]+\s+[А-ЯЁ]{2,}[\.\d]+\s*ТУ\s*$', name.strip(), re.IGNORECASE):
                    is_group_header = True
                # ИЛИ строка содержит ТОЛЬКО ТУ-код без детального наименования (короткая строка)
                elif len(name.strip()) < 30 and not qty_raw.strip():
                    has_tu_code = re.search(r'([А-ЯЁ]{2,}[\.\d]+ТУ)', name) or re.search(r'ТУ\s+[\d\-]+', name)
                    if has_tu_code:
                        is_group_header = True
            
            if is_group_header:
                # Извлекаем ТУ
                tu_pattern1 = r'([А-ЯЁ]{2,}\.\d+[\d\.\-]*\s*ТУ)'
                tu_match1 = re.search(tu_pattern1, name)
                if tu_match1:
                    current_group_tu = tu_match1.group(1)
                else:
                    tu_pattern2 = r'([А-ЯЁ]{2,}[\d\.]+\s*ТУ)'
                    tu_match2 = re.search(tu_pattern2, name)
                    if tu_match2:
                        current_group_tu = tu_match2.group(1)
                    else:
                        tu_pattern3 = r'ТУ\s+([\d\-]+)'
                        tu_match3 = re.search(tu_pattern3, name)
                        if tu_match3:
                            current_group_tu = 'ТУ ' + tu_match3.group(1)
                        else:
                            # Если нет ТУ-кода, проверяем производителя
                            mfr_pattern_header = r'ф\.\s*([A-Za-z][A-Za-z0-9\s\-]+)'
                            mfr_match_header = re.search(mfr_pattern_header, name)
                            if mfr_match_header:
                                # Используем производителя как current_group_tu
                                current_group_tu = mfr_match_header.group(1).strip()
                            else:
                                current_group_tu = ""
                
                # Извлекаем тип компонента
                type_text = name
                if current_group_tu:
                    type_text = type_text.replace(current_group_tu, '')
                type_text = re.sub(r'\s+[А-ЯЁ]+\d+[\dА-ЯЁ\-]*', '', type_text)
                type_text = re.sub(r'\s+[А-ЯЁ]+\.\d+[\d\.]*', '', type_text)
                current_group_type = type_text.strip()
                
                # Извлекаем префикс типа компонента для маппинга (К10, К53, GRM, НР1 и т.д.)
                # Паттерн: буквы + цифры + необязательный дефис и еще буквы/цифры
                component_type_pattern = r'([А-ЯЁ]+[\d\-]+[А-ЯЁ]*|[A-Z]+[\d]*)'
                component_type_match = re.search(component_type_pattern, name)
                if component_type_match and current_group_tu:
                    component_type_prefix = component_type_match.group(1).strip()
                    # Нормализуем: К53 – 65 → К53-65
                    component_type_prefix = component_type_prefix.replace(' ', '').replace('–', '-').replace('—', '-')
                    component_type_tu_map[component_type_prefix] = current_group_tu
                # Также извлекаем производителя из заголовка (например, "GRM ф. Murata")
                mfr_in_header = ""
                mfr_pattern_header = r'ф\.\s*([A-Za-z][A-Za-z0-9\s\-]+)'
                mfr_match_header = re.search(mfr_pattern_header, name)
                if mfr_match_header:
                    mfr_in_header = mfr_match_header.group(1).strip()
                    # Для производителя тоже создаем маппинг
                    component_type_match2 = re.search(r'([A-Z]+)', name, re.IGNORECASE)
                    if component_type_match2:
                        component_type_prefix2 = component_type_match2.group(1).upper()
                        component_type_tu_map[component_type_prefix2] = mfr_in_header
                
                continue  # Пропускаем строку-заголовок
            
            # ===================================================================
            # ОБЪЕДИНЕНИЕ МНОГОСТРОЧНЫХ ЗАПИСЕЙ
            # ===================================================================
            
            # СЛУЧАЙ 1: Строка БЕЗ reference, БЕЗ name, НО С note - это продолжение примечания
            if not ref.strip() and not name.strip() and note.strip() and extracted:
                # Это продолжение примечания для последнего компонента
                last_item = extracted[-1]
                current_note = last_item.get('original_note', '')
                if current_note:
                    # Объединяем через пробел (запятая уже есть в конце предыдущей части)
                    last_item['original_note'] = current_note.strip() + ' ' + note.strip()
                    # ВАЖНО: НЕ обновляем note, если там уже установлен ТУ из заголовка группы!
                    # Также НЕ обновляем note если там производитель (для замен/подборов)
                    # Проверяем, содержит ли current note ТУ-код (паттерн: АЛЯР.434110.005ТУ)
                    last_note = last_item.get('note', '')
                    orig_note_val = last_item.get('original_note', '')
                    # Проверяем что original_note содержит замену - тогда в note производитель
                    is_replacement_continuation = bool(orig_note_val and 'замена' in orig_note_val.lower())
                    
                    # Проверяем: не перезаписываем note если там ТУ, производитель или замена
                    has_tu_in_last_note = bool(last_note and ('ТУ' in last_note or re.search(r'[А-ЯЁ]{2,}\.\d+[\d\.\-]*ТУ', last_note)))
                    looks_like_manufacturer = bool(last_note and len(last_note) < 50 and ',' not in last_note)
                    
                    if not has_tu_in_last_note and not is_replacement_continuation and not looks_like_manufacturer:
                        # В note нет ТУ/производителя и это не замена - можно обновить из original_note
                        last_item['note'] = last_item['original_note']
                else:
                    last_item['original_note'] = note.strip()
                    # Не перезаписываем note если там уже есть ТУ или производитель
                    last_note = last_item.get('note', '')
                    has_tu = bool(last_note and ('ТУ' in last_note or re.search(r'[А-ЯЁ]{2,}\.\d+[\d\.\-]*ТУ', last_note)))
                    looks_like_mfr = bool(last_note and len(last_note) < 50 and ',' not in last_note)
                    is_replacement = bool('замена' in note.lower())
                    
                    if not has_tu and not looks_like_mfr and not is_replacement:
                        last_item['note'] = note.strip()
                continue  # Пропускаем эту строку, т.к. мы уже добавили примечание
            
            # СЛУЧАЙ 2: Строка БЕЗ reference, но С name/note И с qty - продолжение компонента
            # Это случай когда:
            # - Первая строка: L1 | Дроссель высокочастотный ДМ – 3 – 10 ± 5% – В | (пусто) | (пусто)
            # - Вторая строка:    | «Н» ЦКСН.671342.001ТУ                        | 1      | (пусто)
            if not ref.strip() and (name.strip() or note.strip()) and qty_raw.strip() and extracted:
                last_item = extracted[-1]
                # Проверяем, что предыдущая строка имела reference но НЕ имела явного qty
                if last_item.get('reference', '').strip() and not last_item.get('has_explicit_qty', False):
                    # Объединяем description
                    current_desc = last_item.get('description', '').strip()
                    additional_desc = name.strip() if name.strip() else note.strip()
                    if current_desc and additional_desc:
                        last_item['description'] = current_desc + ' ' + additional_desc
                    elif additional_desc:
                        last_item['description'] = additional_desc
                    
                    # Устанавливаем количество
                    try:
                        qty_val = int(re.sub(r'\D', '', qty_raw))
                        last_item['qty'] = qty_val if qty_val > 0 else 1
                        last_item['has_explicit_qty'] = True
                    except (ValueError, AttributeError):
                        last_item['qty'] = 1
                    
                    # Обновляем note/original_note если они были во второй строке
                    if note.strip():
                        last_item['note'] = note.strip()
                        last_item['original_note'] = note.strip()
                    
                    continue  # Пропускаем эту строку, т.к. мы уже объединили с предыдущей

            # If header wasn't detected, try fallback mapping
            if not any([ref, name, qty_raw]) and len(vals) >= 2:
                name = " ".join(vals[:-1])
                qty_raw = vals[-1]

            # parse qty
            qty = None
            has_explicit_qty = False  # Флаг: была ли qty ячейка НЕ пустой
            m = re.search(r"(\d+)", str(qty_raw))
            if m:
                try:
                    qty = int(m.group(1))
                    has_explicit_qty = True
                except Exception:
                    qty = 1
                    has_explicit_qty = False

            # Проверить, есть ли в наименовании информация о производителе (формат "ф. Производитель")
            manufacturer_in_name = ""
            if name:
                # Ищем паттерн "ф. Производитель" в конце или середине строки
                mfr_pattern = r'\s+ф\.\s*([A-Za-z0-9\s\-]+)'
                mfr_match = re.search(mfr_pattern, name)
                if mfr_match:
                    manufacturer_in_name = mfr_match.group(1).strip()
                    # ВАЖНО: НЕ удаляем производителя из description, если в примечании есть ПОДБОРЫ номиналов
                    # (но НЕ замены! Замены - это другое)
                    # Подборные элементы должны наследовать производителя из оригинального description
                    # Удаляем производителя только если нет подборов номиналов
                    has_podbor_preview = bool(cell_note and ref.strip() and (',' in cell_note or ';' in cell_note) 
                                             and 'замена' not in cell_note.lower())
                    if not has_podbor_preview:
                        # Удаляем информацию о производителе из наименования
                        name = re.sub(mfr_pattern, '', name).strip()
            
            # Проверяем, содержит ли ячейка примечания подборы (номиналы с запятыми/точками с запятой)
            # Подборы - это список номиналов или артикулов, разделенных запятыми
            has_podbor_in_note = bool(cell_note and ref.strip() and (',' in cell_note or ';' in cell_note))
            
            # Ищем подходящую ТУ в словаре component_type_tu_map по описанию компонента
            # Например, для "К53-65-..." ищем "К53-65" или "К53" в словаре
            # ВАЖНО: Делаем это ПЕРЕД проверкой should_use_group_tu
            matched_tu_from_map = ""
            if component_type_tu_map and name:
                # Нормализуем описание для поиска
                name_normalized = name.replace(' ', '').replace('–', '-').replace('—', '-')
                # Ищем совпадения по убыванию длины (сначала точные, потом короткие)
                for comp_type in sorted(component_type_tu_map.keys(), key=len, reverse=True):
                    if comp_type in name_normalized:
                        matched_tu_from_map = component_type_tu_map[comp_type]
                        break
            
            # Используем matched_tu_from_map если найдено, иначе current_group_tu
            effective_group_tu = matched_tu_from_map if matched_tu_from_map else current_group_tu
            
            # Проверяем, принадлежит ли текущий компонент к текущей группе
            # Если префикс позиционного обозначения не соответствует типу группы - сбрасываем ТУ группы
            ref_prefix = re.sub(r'\d.*$', '', ref.split()[0].upper()) if ref else ""
            should_use_group_tu = True
            
            if effective_group_tu and ref_prefix:
                # Проверяем соответствие префикса типу группы
                group_type_lower = current_group_type.lower() if current_group_type else ""
                
                # Резисторы (R) - ТУ резисторов не применяется к PAT, D, C и т.д.
                if 'резистор' in group_type_lower and not ref_prefix.startswith('R'):
                    should_use_group_tu = False
                # Конденсаторы (C) - ТУ конденсаторов не применяется к R, D и т.д.
                elif 'конденсатор' in group_type_lower and not ref_prefix.startswith('C'):
                    should_use_group_tu = False
                # Микросхемы (DA, DD, U, IC) - ТУ микросхем не применяется к R, C, D и т.д.
                elif 'микросхем' in group_type_lower and not ref_prefix.startswith(('DA', 'DD', 'U', 'IC')):
                    should_use_group_tu = False
                # Индуктивности (L) - ТУ индуктивностей не применяется к R, C, D и т.д.
                elif ('дроссел' in group_type_lower or 'индуктивност' in group_type_lower) and not ref_prefix.startswith('L'):
                    should_use_group_tu = False
            
            # Фильтруем служебные фразы из cell_note
            # "допускается отсутствие", "справ." и т.д. - не являются ТУ
            service_phrases = ['допускается отсутствие', 'справ.', 'см. примечание']
            is_service_note = any(phrase in cell_note.lower() for phrase in service_phrases) if cell_note else False
            
            # Проверяем, является ли примечание заменой (содержит "замена")
            is_replacement_note = bool(cell_note and 'замена' in cell_note.lower())
            
            # Добавить ТОЛЬКО ТУ из заголовка группы в note (НЕ тип компонента!)
            # Тип компонента из заголовка может быть неточным и перевесить правильную классификацию
            if has_podbor_in_note and effective_group_tu and should_use_group_tu and not is_replacement_note:
                # Ячейка примечания содержит подборы - используем ТУ из заголовка для note
                # Подборы остаются в original_note для последующего извлечения
                note = effective_group_tu
                if manufacturer_in_name:
                    note = note + ' | ' + manufacturer_in_name
            elif has_podbor_in_note and manufacturer_in_name and not is_replacement_note:
                # Ячейка примечания содержит подборы, но нет ТУ из заголовка
                # (например, для PAT компонентов нет группового заголовка)
                # Используем производитель из наименования
                note = manufacturer_in_name
            elif is_replacement_note and manufacturer_in_name:
                # Ячейка примечания содержит замены - используем производитель из наименования для note
                # Примечание о замене остается в original_note для последующего извлечения
                note = manufacturer_in_name
            elif effective_group_tu and manufacturer_in_name and should_use_group_tu:
                # Есть и ТУ из заголовка, и производитель из наименования
                note = effective_group_tu + ' | ' + manufacturer_in_name
            elif effective_group_tu and not cell_note and should_use_group_tu:
                # Только ТУ из заголовка (и ячейка примечания пустая)
                note = effective_group_tu
            elif manufacturer_in_name:
                # Только производитель из наименования
                note = manufacturer_in_name
            elif is_service_note:
                # Ячейка содержит служебную фразу (например, "допускается отсутствие")
                # Для основного компонента используем ТУ из заголовка (если есть)
                if effective_group_tu and should_use_group_tu:
                    note = effective_group_tu
                else:
                    note = ""
            # Если cell_note не содержит подборы и нет ТУ из заголовка - оставляем note=cell_note (уже установлено)

            # Не добавлять строку без данных
            if not ref.strip() and not name.strip():
                continue
            
            # Фильтровать служебные записи (Изм., Лист регистрации, и т.д.)
            name_check = name.lower().strip()
            service_keywords = [
                'изм.', 'изме-ненных', 'заме-ненных', 'аннули-рован', 'всего листов', 
                'номер докум', 'входя-щий', 'сопрово-дитель', 'подп.',
                'лист регистрации', 'регистрации изменений'
            ]
            if any(kw in name_check for kw in service_keywords):
                continue
            
            # Обработать строки "ф. Производитель" (без reference)
            # Это строка с производителем для ПРЕДЫДУЩЕГО элемента
            if not ref.strip() and name.strip().startswith('ф.'):
                if extracted:
                    manufacturer_text = name.strip()
                    # Извлечь производителя (после "ф.")
                    mfr_match = re.search(r'ф\.\s*(.+)', manufacturer_text)
                    if mfr_match:
                        manufacturer = mfr_match.group(1).strip()
                        # Добавить производителя к note последнего элемента
                        # ИСПРАВЛЕНО: Порядок изменён на "ТУ | manufacturer" для правильной обработки
                        if extracted[-1]['note']:
                            # Если note уже содержит ТУ-код, добавляем производителя ПОСЛЕ разделителя
                            extracted[-1]['note'] = extracted[-1]['note'] + ' | ' + manufacturer
                        else:
                            # Если note пустой, добавляем только производителя
                            extracted[-1]['note'] = manufacturer
                continue
            
            # Убираем запятую в конце названия (если есть)
            name = name.rstrip(',').strip()

            # Если количество не указано явно, пытаемся посчитать из reference (например, FU1-FU6 = 6)
            if qty is None or qty == 0:
                qty = count_from_reference(ref)
            
            # Определить нужно ли использовать group_type для этого элемента
            # Сбрасываем group_type если элемент явно не принадлежит к текущей группе
            use_group_type = current_group_type
            if ref.strip():
                # Проверяем префикс позиционного обозначения
                ref_prefix = re.sub(r'\d.*$', '', ref.split()[0].upper()) if ref else ""
                
                # Если группа "Резисторы", но префикс НЕ R - сбрасываем
                if current_group_type and 'резистор' in current_group_type.lower() and not ref_prefix.startswith('R'):
                    use_group_type = ""
                # Если группа "Конденсаторы", но префикс НЕ C - сбрасываем
                elif current_group_type and 'конденсатор' in current_group_type.lower() and not ref_prefix.startswith('C'):
                    use_group_type = ""
                # Если группа "Микросхемы", но префикс НЕ DA/DD/U - сбрасываем
                elif current_group_type and 'микросхем' in current_group_type.lower() and not ref_prefix.startswith(('DA', 'DD', 'U', 'IC')):
                    use_group_type = ""
                # Если группа "Индуктивности", но префикс НЕ L - сбрасываем
                elif current_group_type and ('дроссел' in current_group_type.lower() or 'индуктивност' in current_group_type.lower()) and not ref_prefix.startswith('L'):
                    use_group_type = ""
                # Если группа "Разъемы", но префикс НЕ X/XS/J/P - сбрасываем
                elif current_group_type and 'разъем' in current_group_type.lower() and not ref_prefix.startswith(('X', 'J', 'P')):
                    use_group_type = ""

            row = {
                "zone": zone,
                "reference": ref,
                "description": name if name else note,
                "qty": qty,
                "note": note,  # ТУ и производитель, НЕ тип компонента
                "group_type": use_group_type,  # Тип компонента для классификации (может быть сброшен)
                "original_note": original_note,  # Оригинальное примечание для подборов
                "has_explicit_qty": has_explicit_qty,  # Флаг для отслеживания многострочных элементов
            }
            extracted.append(row)

    # Additionally parse free text paragraphs (fallback)
    for p in doc.paragraphs:
        t = normalize_cell(p.text)
        if not t:
            continue
        
        # Фильтровать служебные записи из параграфов
        t_check = t.lower().strip()
        service_keywords = [
            'изм.', 'изме-ненных', 'заме-ненных', 'аннули-рован', 'всего листов', 
            'номер докум', 'входя-щий', 'сопрово-дитель', 'подп.',
            'лист регистрации', 'регистрации изменений'
        ]
        if any(kw in t_check for kw in service_keywords):
            continue
        
        parts = [s.strip() for s in LINE_SPLIT_RE.split(t) if s.strip()]
        if parts:
            pos = parts[0] if POS_PREFIX_RE.match(parts[0]) else ""
            m = re.search(r"(\d+)\s*(шт\.?|pcs|pieces)?", t, flags=re.IGNORECASE)
            qty = int(m.group(1)) if m else 1
            desc = " ".join(parts[1:]) if pos else " ".join(parts)
            extracted.append({"reference": pos, "description": desc, "qty": qty})

    if not extracted:
        extracted = [{"description": " ".join(normalize_cell(p.text) for p in doc.paragraphs), "qty": 1}]
    
    return pd.DataFrame(extracted)
