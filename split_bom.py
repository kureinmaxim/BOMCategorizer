import argparse
import os
import sys
import re
import json
from typing import List, Optional, Dict, Any, Iterable, Tuple, Union

import pandas as pd

try:
    from docx import Document  # python-docx
except Exception:
    Document = None  # optional; raise if used without installed

# Исправление кодировки для корректного вывода русских символов в терминале
if sys.stdout.encoding != 'utf-8':
    try:
        sys.stdout.reconfigure(encoding='utf-8')
        sys.stderr.reconfigure(encoding='utf-8')
    except AttributeError:
        # Для старых версий Python
        import codecs
        sys.stdout = codecs.getwriter('utf-8')(sys.stdout.buffer, 'strict')
        sys.stderr = codecs.getwriter('utf-8')(sys.stderr.buffer, 'strict')


def normalize_column_names(columns: List[str]) -> List[str]:
    normalized = []
    for name in columns:
        if name is None:
            normalized.append("")
            continue
        normalized.append(str(name).strip().lower())
    return normalized


def find_column(possible_names: List[str], columns: List[str]) -> Optional[str]:
    # Сначала ищем точное совпадение
    for candidate in possible_names:
        if candidate in columns:
            return candidate
    # Если не нашли точное совпадение, ищем частичное (колонка начинается с candidate)
    for candidate in possible_names:
        for col in columns:
            if col.startswith(candidate):
                return col
    return None


def has_any(text: str, keywords: List[str]) -> bool:
    if not isinstance(text, str):
        return False
    lower = text.lower()
    return any(k in lower for k in keywords)


RESISTOR_VALUE_RE = re.compile(r"(?i)\b\d+(?:[\.,]\d+)?\s*(?:ом|ohm|k\s*ohm|kohm|к\s*ом|ком|m\s*ohm|mohm|м\s*ом|мом)\b")
CAP_VALUE_RE = re.compile(r"(?i)\b\d+(?:[\.,]\d+)?\s*(?:pf|nf|uf|µf|μf|ф|пф|нф|мкф)\b")
IND_VALUE_RE = re.compile(r"(?i)\b\d+(?:[\.,]\d+)?\s*(?:nh|uh|µh|μh|mh|h|нгн|мкгн|мгн|гн)\b")


# NEW: lightweight parser for TXT/DOCX tables and lines
LINE_SPLIT_RE = re.compile(r"\s{2,}|\t|;|,\s?(?=\S)" )
POS_PREFIX_RE = re.compile(r"^(?:[A-ZА-Я]+\d+(?:[-,;\s]*[A-ZА-Я]*\d+)*)$", re.IGNORECASE)


def parse_txt_like(path: str) -> pd.DataFrame:
    rows: List[Dict[str, Any]] = []
    try:
        with open(path, "r", encoding="utf-8") as f:
            text = f.read()
    except UnicodeDecodeError:
        with open(path, "r", encoding="cp1251", errors="ignore") as f:
            text = f.read()
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        parts = [p.strip() for p in LINE_SPLIT_RE.split(line) if p.strip()]
        if not parts:
            continue
        # Heuristic mapping: [pos?, name/desc, qty?]
        pos = parts[0] if POS_PREFIX_RE.match(parts[0]) else None
        qty = None
        # Try explicit "Nшт" or "N шт" patterns first
        m = re.search(r"(\d+)\s*(шт\.?|pcs|pieces)", line, flags=re.IGNORECASE)
        if m:
            qty = int(m.group(1))
        else:
            for p in parts[::-1]:
                if re.fullmatch(r"\d+", p):
                    qty = int(p)
                    break
        desc = " ".join(parts[1:-1]) if pos and len(parts) >= 2 else (" ".join(parts))
        row = {"reference": pos or "", "description": desc, "qty": qty if qty is not None else 1}
        rows.append(row)
    if not rows:
        # fallback: whole text in a single row
        rows = [{"description": text, "qty": 1}]
    return pd.DataFrame(rows)


def parse_docx(path: str) -> pd.DataFrame:
    import re  # Импортируем re в начале функции
    if Document is None:
        raise SystemExit("python-docx is required to parse DOCX. Install with pip install python-docx")
    doc = Document(path)
    extracted: List[Dict[str, Any]] = []

    def guess_header_index(table) -> int:
        max_scan = min(5, len(table.rows))
        header_keywords = ["наимен", "обознач", "кол.", "кол ", "кол", "примеч"]
        for i in range(max_scan):
            row_texts = [normalize_cell(c.text).lower() for c in table.rows[i].cells]
            hits = sum(any(hk in t for hk in header_keywords) for t in row_texts)
            if hits >= 2:
                return i
        return 0

    # Parse tables with header detection
    for table in doc.tables:
        if not table.rows:
            continue
        header_idx = guess_header_index(table)
        header_cells = [normalize_cell(c.text).strip() for c in table.rows[header_idx].cells]
        header_norm = [h.lower() for h in header_cells]

        # Find column indices by common names
        def find_col_idx(candidates: List[str]) -> Optional[int]:
            for i, h in enumerate(header_norm):
                for cand in candidates:
                    if cand in h:
                        return i
            return None

        idx_zone = find_col_idx(["зона"])  # optional
        idx_ref = find_col_idx(["поз", "обозн"])  # позиционное обозначение
        idx_name = find_col_idx(["наимен"])  # наименование
        idx_qty = find_col_idx(["кол.", "кол ", "кол", "количество"])  # количество
        idx_note = find_col_idx(["примеч"])  # опционально

        # Переменные для хранения текущего ТУ и типа компонента из заголовка группы
        current_group_tu = ""
        current_group_type = ""
        
        for tr in table.rows[header_idx + 1:]:
            vals = [normalize_cell(c.text) for c in tr.cells]
            if not any(v.strip() for v in vals):
                continue
            zone = vals[idx_zone] if idx_zone is not None and idx_zone < len(vals) else ""
            ref = vals[idx_ref] if idx_ref is not None and idx_ref < len(vals) else ""
            name = vals[idx_name] if idx_name is not None and idx_name < len(vals) else ""
            qty_raw = vals[idx_qty] if idx_qty is not None and idx_qty < len(vals) else ""
            note = vals[idx_note] if idx_note is not None and idx_note < len(vals) else ""
            
            # Проверить, является ли это строкой-заголовком группы
            # Примеры: "Резисторы Р1-12 ШКАБ.434110.002 ТУ", "Набор резисторов НР1-4Р ШКАБ.434110.018 ТУ"
            section_headers = ['конденсаторы', 'резисторы', 'микросхемы', 'дроссели', 'индуктивности',
                             'разъемы', 'диоды', 'транзисторы', 'кабели', 'модули', 
                             'набор резисторов', 'набор конденсаторов', 'набор микросхем']
            name_lower = name.strip().lower()
            
            # Если строка начинается с названия раздела И нет позиционного обозначения И нет количества
            # ИЛИ если в строке есть ТУ и нет позиционного обозначения и количества
            is_group_header = False
            has_tu_code = re.search(r'([А-ЯЁ]{2,}[\.\d]+ТУ)', name) or re.search(r'ТУ\s+[\d\-]+', name)
            
            if not ref.strip() and not qty_raw.strip():
                # Проверяем: начинается ли с раздела ИЛИ содержит ТУ
                if any(name_lower.startswith(section) for section in section_headers) or has_tu_code:
                    is_group_header = True
            
            if is_group_header:
                # Это заголовок группы - извлечь ТУ и тип компонента
                
                # Извлекаем ТУ (различные форматы)
                # Важно: ТУ может быть с пробелом или без ("ШКАБ.434110.018 ТУ" или "ШКАБ.434110.018ТУ")
                tu_pattern1 = r'([А-ЯЁ]{2,}\.\d+[\d\.\-]*\s*ТУ)'
                tu_match1 = re.search(tu_pattern1, name)
                if tu_match1:
                    current_group_tu = tu_match1.group(1)
                else:
                    # Вариант 2: без точек (например, "СМ3.362.805ТУ")
                    tu_pattern2 = r'([А-ЯЁ]{2,}[\d\.]+\s*ТУ)'
                    tu_match2 = re.search(tu_pattern2, name)
                    if tu_match2:
                        current_group_tu = tu_match2.group(1)
                    else:
                        # Вариант 3: ТУ в начале (например, "ТУ 6329-019-07614320-99")
                        tu_pattern3 = r'ТУ\s+([\d\-]+)'
                        tu_match3 = re.search(tu_pattern3, name)
                        if tu_match3:
                            current_group_tu = 'ТУ ' + tu_match3.group(1)
                        else:
                            current_group_tu = ""
                
                # Извлекаем тип компонента (например, "Набор резисторов", "Резисторы")
                # Для строк типа "Набор резисторов НР1-4Р ШКАБ.434110.018 ТУ"
                # Нужно оставить только "Набор резисторов"
                type_text = name
                # Удаляем ТУ
                if current_group_tu:
                    type_text = type_text.replace(current_group_tu, '')
                # Удаляем партномеры (буквы+цифры+дефис+буквы/цифры) и коды ТУ
                type_text = re.sub(r'\s+[А-ЯЁ]+\d+[\dА-ЯЁ\-]*', '', type_text)  # НР1-4Р, Р1-12
                type_text = re.sub(r'\s+[А-ЯЁ]+\.\d+[\d\.]*', '', type_text)  # ШКАБ.434110.018
                current_group_type = type_text.strip()
                
                continue  # Пропускаем строку-заголовок, не добавляем в результат

            # If header wasn't detected, try fallback mapping by last column digits
            if not any([ref, name, qty_raw]) and len(vals) >= 2:
                name = " ".join(vals[:-1])
                qty_raw = vals[-1]

            # parse qty
            qty = None
            m = re.search(r"(\d+)", str(qty_raw))
            if m:
                try:
                    qty = int(m.group(1))
                except Exception:
                    qty = 1

            # Добавить ТУ и тип из заголовка группы в note
            # ТУ идет в note (колонка "Примечания" в исходном файле или для хранения ТУ)
            # Тип компонента тоже идет в note (будет обработан позже в clean_component_name)
            if current_group_tu or current_group_type:
                parts = []
                if current_group_type:
                    parts.append(current_group_type)
                if current_group_tu:
                    parts.append(current_group_tu)
                # НЕ добавляем исходный note, если есть ТУ из заголовка (чтобы не дублировать)
                # if note.strip():
                #     parts.append(note.strip())
                note = ' | '.join(parts) if parts else note

            # Не добавлять строку, если нет ни reference, ни description (name)
            # (note может содержать ТУ из заголовка группы, но это не основные данные)
            if not ref.strip() and not name.strip():
                continue
            
            row = {
                "zone": zone,
                "reference": ref,
                "description": name if name else note,
                "qty": qty if qty is not None else 1,
                "note": note,
            }
            extracted.append(row)

    # Additionally parse free text paragraphs (fallback)
    for p in doc.paragraphs:
        t = normalize_cell(p.text)
        if not t:
            continue
        parts = [s.strip() for s in LINE_SPLIT_RE.split(t) if s.strip()]
        if parts:
            pos = parts[0] if POS_PREFIX_RE.match(parts[0]) else ""
            m = re.search(r"(\d+)\s*(шт\.?|pcs|pieces)?", t, flags=re.IGNORECASE)
            qty = int(m.group(1)) if m else 1
            desc = " ".join(parts[1:]) if pos else " ".join(parts)
            extracted.append({"reference": pos, "description": desc, "qty": qty})

    if not extracted:
        extracted = [{"description": " ".join(normalize_cell(p.text) for p in doc.paragraphs), "qty": 1}]
    return pd.DataFrame(extracted)


def normalize_cell(s: Any) -> str:
    return (str(s or "").strip())


def classify_row(ref: Optional[str], description: Optional[str], value: Optional[str], partname: Optional[str], strict: bool, source_file: Optional[str] = None, note: Optional[str] = None) -> str:
    def to_text(x: Any) -> str:
        if x is None:
            return ""
        try:
            import math
            if isinstance(x, float) and math.isnan(x):
                return ""
        except Exception:
            pass
        s = str(x)
        return s.strip()

    ref = to_text(ref)
    desc = to_text(description)
    val = to_text(value)
    part = to_text(partname)
    src_file = to_text(source_file)
    note_text = to_text(note)

    # Create text blob early for use in reference-based checks (теперь включая note!)
    text_blob = " ".join([desc, val, part, note_text])

    # Refdes first where reliable
    ref_prefix = ref.split(" ")[0].upper() if ref else ""
    ref_prefix = re.sub(r"\d.*$", "", ref_prefix)  # take letters before digits

    # PRIORITY 1: Check context-specific categories FIRST (before generic prefixes)
    # Check if this is a board/PCB file (self-reference: description is just the filename)
    if src_file and desc:
        # Extract filename without extension
        file_base = src_file.split('/')[-1].split('\\')[-1].rsplit('.', 1)[0].lower()
        desc_lower = desc.lower()
        
        # If description is ONLY the filename (board referencing itself), it's our development
        # AND doesn't contain component keywords
        component_keywords = ['резистор', 'конденсатор', 'микросхема', 'разъем', 'диод', 'индуктор', 'дроссель',
                             'транзистор', 'стабилитрон', 'генератор', 'вилка', 'розетка', 'кабель']
        is_component = any(kw in desc_lower for kw in component_keywords)
        
        if not is_component and file_base in desc_lower.replace('.xlsx', '').replace('.xls', ''):
            # Check if it's just the filename without other text (with tolerance for spaces/extensions)
            desc_clean = desc_lower.replace('.xlsx', '').replace('.xls', '').replace(' ', '').replace('_', '')
            file_clean = file_base.replace(' ', '').replace('_', '')
            if desc_clean == file_clean or desc_clean.startswith(file_clean) or file_clean in desc_clean:
                return "our_developments"
    
    # Our developments - check before "A" prefix
    if has_any(text_blob, ["мвок", "наша разработ", "собственной разработ", "шск-м", "плата контроллера шск"]):
        return "our_developments"
    
    # Optical modules with U prefix - check before "U" prefix for ICs
    if ref and ref_prefix.startswith("U"):
        if has_any(text_blob, ["оптический модуль", "optical module", "передающий оптический", "приемный оптический", "mp2320"]):
            return "optics"
    
    # PRIORITY 2: Heuristics by ref (only if we have a real ref column)
    if ref:
        if ref_prefix.startswith("R"):
            return "resistors"
        if ref_prefix.startswith("C"):
            return "capacitors"
        if ref_prefix.startswith("L"):
            return "inductors"
        if ref_prefix.startswith(("U", "DD", "DA", "IC")):
            return "ics"
        if ref_prefix.startswith(("J", "X", "P", "K", "XS", "XP", "JTAG")):
            return "connectors"
        # Prefix "A" or "А" (latin or cyrillic) -> отладочные платы (dev boards)
        if ref_prefix in ("A", "А"):
            # Simple prefix A1, A2, etc.
            return "dev_boards"
        # Russian prefix "А" for attenuators (optics) - only if longer than 2 chars or has optical keywords
        if ref_prefix.startswith(("А", "A")) and len(ref_prefix) > 2:
            # Check if it's really an attenuator
            if has_any(text_blob, ["аттенюат", "ослабител", "attenuator", "fc/apc", "fc/upc", "оптич"]):
                return "optics"
        # Prefix "W" often used for RF modules, waveguides, delay lines
        if ref_prefix.startswith("W"):
            if has_any(text_blob, ["свч", "rf", "линия задержек", "delay line", "усилитель", "делитель", "сумматор", "splitter", "combiner", "amplifier"]):
                return "rf_modules"
        # Prefix "WS" for splitters/dividers
        if ref_prefix.startswith("WS"):
            return "rf_modules"
        # Prefix "WU" for RF components
        if ref_prefix.startswith("WU"):
            return "rf_modules"
        # Prefix "H" for indicators/LEDs
        if ref_prefix.startswith("H"):
            return "semiconductors"
        # Prefix "V", "VT", "Q" for transistors (BUT check text first - "Микросхема" has priority)
        if ref_prefix.startswith(("V", "VT", "Q")):
            # Если в тексте явно указано "Микросхема", то это микросхема
            if has_any(text_blob, ["микросхем", "микросхема"]):
                return "ics"
            return "semiconductors"
        # Prefix "D" for diodes (BUT check text first)
        if ref_prefix.startswith("D"):
            # Если в тексте явно указано "Микросхема" или другой IC тип
            if has_any(text_blob, ["микросхем", "микросхема"]):
                return "ics"
            return "semiconductors"
        # Prefix "S" for switches/buttons (when not connectors)
        if ref_prefix.startswith("S"):
            if has_any(text_blob, ["переключ", "тумблер", "кнопка", "switch", "button", "toggle"]):
                return "others"

    # Russian and English keywords
    if RESISTOR_VALUE_RE.search(text_blob) or has_any(text_blob, ["резист", "resistor"]):
        return "resistors"

    if CAP_VALUE_RE.search(text_blob) or has_any(text_blob, ["конденс", "capacitor", "tantalum", "ceramic", "к10-", "к53-"]):
        return "capacitors"

    if IND_VALUE_RE.search(text_blob) or has_any(text_blob, ["дросс", "индукт", "inductor", "ferrite", "феррит", "катушка", "choke"]):
        return "inductors"
    
    # Предохранители - check BEFORE semiconductors and ICs (чтобы имели приоритет)
    if has_any(text_blob, ["предохранитель", "fuse", "fuzetec"]):
        return "others"
    
    # Semiconductors (диоды, транзисторы, стабилитроны, оптроны) - check BEFORE ICs
    if has_any(text_blob, [
        "диод", "стабилитрон", "транзистор", "оптрон", "оптопар", "2с630", "2т630", "индикатор", 
        "led ", "svetodiod", "indicator", "transistor", "optocoupler", "thyristor", "тиристор",
        "mosfet", "igbt", "triac", "симистор", "полевой транзистор", "биполярный транзистор"
    ]):
        return "semiconductors"

    if has_any(text_blob, [
        "микросхем", " ic", "mcu", "контроллер", "процессор", "оп-амп", "op-amp", "opamp", "adc", "dac", "fpga",
        "asic", "драйвер ", "компаратор", "стабил", "регулятор", "transceiver", "sn74", "ti ", "stm32", "lmk", "ad9"
    ]):
        return "ics"

    if has_any(text_blob, [
        "разъем", "разъём", "connector", "header", "socket", "rj45", "rj11", "sma", "bnc", "terminal", "клемм",
        "штырь", "pin header", "fpc", "ffc", "din", "dc jack", "barrel", "штекер", "вилка", "розетка", "d-sub", "harting"
    ]):
        return "connectors"

    if has_any(text_blob, [
        "отладоч", " dev board", "evaluation", "eval", "nucleo", "arduino", "raspberry",
        "esp32", "stm32 nucleo", "breakout", "fmc", "carrier", "ultrazed", "microzed", "picozed", "zedboard",
        "zynq", "som ", "system on module", "voyager", "tinypilot"
    ]):
        return "dev_boards"

    # New categories
    if has_any(text_blob, [
        "оптичес", "лазер", "оптопара", "led ", "светодиод", "fiber", "оптоволок", "sfp", "qsfp", "transceiver module",
        "аттенюат", "ослабител", "fc/apc", "fc/upc", "sc/apc", "lc/apc", "pigtail", "патч-корд оптич"
    ]):
        return "optics"

    if has_any(text_blob, [
        "свч", "вч ", "rf ", "microwave", "mini-circuits", "planar monolithics", "pmi", "qualwave", "ghz", "lna", "rf amp",
        "линия задержек", "delay line", "делитель", "сумматор", "splitter", "combiner", "аттенюатор свч", "усилител",
        "polaris", "gigabaudics", "etl systems", "vat-", "zx60", "pne-l"
    ]):
        return "rf_modules"

    if has_any(text_blob, [
        "кабель", "cable", "шлейф", "провод", "wire", "patch cord", "jumper"
    ]):
        return "cables"

    if has_any(text_blob, [
        "модуль питания", "power module", "dc-dc", "ac-dc", "buck", "boost", "источник питания", "блок питания", "psu",
        "converter"
    ]):
        return "power_modules"

    # Our developments
    if has_any(text_blob, [
        "мвок", "наша разработ", "собственной разработ", "шск-м", "плата контроллера шск"
    ]):
        return "our_developments"

    # OTHER general hardware to bucket into 'others' (cabinets, bolts, shelves, keyboards etc.)
    if has_any(text_blob, [
        "rittal", "шкаф", "станция", "полка", "кронштейн", "ролик", "болт", "гайка", "шайба", "клавиатура", "моноблок",
        "кабель", "клеммная", "корпус", "шасси", "стеллаж", "стойка", "провод", "розетка", "вентилятор", "генератор",
        "предохранитель", "держател", "зажим", "fuzetec"
    ]):
        return "others"

    # In strict mode, avoid loose matches
    return "unclassified"


def main():
    parser = argparse.ArgumentParser(description="Split BOM Excel into category CSVs or a single XLSX with sheets")
    parser.add_argument("--input", help="Path to input XLSX file (deprecated if --inputs used)")
    parser.add_argument("--inputs", nargs="+", help="Paths to input XLSX files")
    parser.add_argument("--sheet", default=None, help="Sheet name or index (default: first if not using --sheets)")
    parser.add_argument("--sheets", default=None, help="Comma-separated sheet names or indices to process")
    parser.add_argument("--out", default=None, help="Output directory for CSVs (optional)")
    parser.add_argument("--xlsx", default="categorized.xlsx", help="Path to output XLSX with category sheets")
    parser.add_argument("--merge-into", dest="merge_into", default=None, help="Existing categorized XLSX to merge into (append)")
    parser.add_argument("--combine", action="store_true", help="Create a SUMMARY sheet with aggregated quantities")
    parser.add_argument("--interactive", action="store_true", help="Interactive classification for unclassified rows")
    parser.add_argument("--loose", action="store_true", help="Use looser heuristics (may misclassify non-electronics)")
    parser.add_argument("--assign-json", dest="assign_json", default="rules.json", help="Path to JSON rules for assigning categories to unclassified rows. Format: [{'contains': 'text', 'category': 'ics'}]")
    parser.add_argument("--txt-dir", dest="txt_dir", default=None, help="Output directory for TXT files per category (in addition to XLSX)")
    args = parser.parse_args()

    if not args.inputs and not args.input:
        raise SystemExit("Provide --inputs <files...> or --input <file>")

    inputs: List[str] = args.inputs or ([args.input] if args.input else [])
    selected_sheets: Optional[List[Any]] = None
    if args.sheets:
        selected_sheets = []
        for token in args.sheets.split(','):
            token = token.strip()
            if not token:
                continue
            try:
                selected_sheets.append(int(token))
            except ValueError:
                selected_sheets.append(token)

    if args.out:
        os.makedirs(args.out, exist_ok=True)

    # Load workbook
    read_kwargs = {"engine": "openpyxl"}
    all_rows: List[pd.DataFrame] = []
    for input_path in inputs:
        try:
            ext = os.path.splitext(input_path)[1].lower()
            if ext in [".txt"]:
                df_local = parse_txt_like(input_path)
                df_local["source_file"] = os.path.basename(input_path)
                df_local["source_sheet"] = "txt"
                all_rows.append(df_local)
                continue
            if ext in [".docx"]:
                df_local = parse_docx(input_path)
                df_local["source_file"] = os.path.basename(input_path)
                df_local["source_sheet"] = "docx"
                all_rows.append(df_local)
                continue
            if ext in [".doc"]:
                # Try to convert .doc -> .docx via Word COM (if available), else fallback to text
                try:
                    from win32com.client import Dispatch  # type: ignore
                    word = Dispatch("Word.Application")
                    word.Visible = False
                    doc = word.Documents.Open(os.path.abspath(input_path))
                    tmp_docx = os.path.splitext(os.path.abspath(input_path))[0] + "_conv.docx"
                    wdFormatXMLDocument = 12
                    doc.SaveAs(tmp_docx, FileFormat=wdFormatXMLDocument)
                    doc.Close(False)
                    word.Quit()
                    df_local = parse_docx(tmp_docx)
                except Exception:
                    df_local = parse_txt_like(input_path)
                df_local["source_file"] = os.path.basename(input_path)
                df_local["source_sheet"] = "doc"
                all_rows.append(df_local)
                continue

            if selected_sheets is not None:
                # read all selected sheets explicitly
                for sh in selected_sheets:
                    df_local = pd.read_excel(input_path, sheet_name=sh, **{k: v for k, v in read_kwargs.items() if k != 'sheet_name'})
                    if isinstance(df_local, dict):
                        # unlikely when specifying a single sheet, but guard anyway
                        for _, dfi in df_local.items():
                            # Проверка на пустую первую строку
                            if all(str(col).lower().startswith('unnamed') for col in dfi.columns):
                                if not dfi.empty and dfi.iloc[0].notna().any():
                                    new_headers = dfi.iloc[0].fillna('').astype(str)
                                    dfi = dfi[1:].reset_index(drop=True)
                                    dfi.columns = new_headers
                            dfi["source_file"] = os.path.basename(input_path)
                            dfi["source_sheet"] = str(sh)
                            all_rows.append(dfi)
                    else:
                        # Проверка на пустую первую строку
                        if all(str(col).lower().startswith('unnamed') for col in df_local.columns):
                            if not df_local.empty and df_local.iloc[0].notna().any():
                                new_headers = df_local.iloc[0].fillna('').astype(str)
                                df_local = df_local[1:].reset_index(drop=True)
                                df_local.columns = new_headers
                        df_local["source_file"] = os.path.basename(input_path)
                        df_local["source_sheet"] = str(sh)
                        all_rows.append(df_local)
            else:
                # single sheet or first sheet
                sheet = args.sheet
                if sheet is not None:
                    try:
                        sheet = int(sheet)
                    except ValueError:
                        pass
                    read_kwargs["sheet_name"] = sheet

                df = pd.read_excel(input_path, **read_kwargs)
                if isinstance(df, dict):
                    # take first sheet
                    first_key = next(iter(df))
                    df = df[first_key]
                    src_sheet = first_key
                else:
                    src_sheet = sheet if sheet is not None else 0
                
                # Проверка: если все колонки unnamed и первая строка содержит текст - это заголовки
                # (первая строка файла была пустая)
                if all(str(col).lower().startswith('unnamed') for col in df.columns):
                    if not df.empty and df.iloc[0].notna().any():
                        # Первая строка данных содержит настоящие заголовки
                        new_headers = df.iloc[0].fillna('').astype(str)
                        df = df[1:].reset_index(drop=True)
                        df.columns = new_headers
                
                df["source_file"] = os.path.basename(input_path)
                df["source_sheet"] = str(src_sheet)
                all_rows.append(df)
        except Exception as exc:
            raise SystemExit(f"Failed to read Excel '{input_path}': {exc}")

    if not all_rows:
        raise SystemExit("No data loaded from inputs")

    df = pd.concat(all_rows, ignore_index=True)

    # Normalize columns
    original_cols = list(df.columns)
    lower_cols = normalize_column_names(original_cols)
    rename_map = {orig: norm for orig, norm in zip(original_cols, lower_cols)}
    df = df.rename(columns=rename_map)

    # Common column guesses
    ref_col = find_column(["ref", "reference", "designator", "refdes", "reference designator", "обозначение", "позиционное обозначение"], list(df.columns))
    # Для description: сначала ищем точные совпадения, потом частичные (наименование мр)
    desc_col = find_column(["description", "desc", "наименование", "имя", "item", "part", "part name", "наим."], list(df.columns))
    value_col = find_column(["value", "значение", "номинал"], list(df.columns))
    part_col = find_column(["partnumber", "mfr part", "mpn", "pn", "art", "артикул", "part", "part name"], list(df.columns))
    qty_col = find_column([
        "qty", "quantity", "количество", "кол.", "кол-во", "кол. в ктд", "кол в ктд", "кол. в спецификации", "кол. в кдт",
        "кол. в ктд", "кол. в ктд, шт", "кол. в ктд (шт)", "кол. в ктд, шт."
    ], list(df.columns))
    mr_col = find_column([
        "код мр", "код ивп", "код мр/ивп", "код позиции", "код изделия", "код мр позиции", "код мр ивп"
    ], list(df.columns))
    
    # Проверяем, есть ли несколько колонок с description/наименование (из разных файлов)
    possible_desc_cols = [col for col in df.columns if any(
        col.startswith(prefix) for prefix in ["description", "наименование", "desc", "имя"]
    )]
    
    if len(possible_desc_cols) > 1:
        # Объединяем несколько колонок description в одну
        def merge_desc(row):
            for col in possible_desc_cols:
                val = row.get(col)
                if pd.notna(val) and str(val).strip():
                    return val
            return None
        
        df["_merged_description_"] = df.apply(merge_desc, axis=1)
        # Удаляем старые колонки description/наименование
        for col in possible_desc_cols:
            if col in df.columns:
                df = df.drop(columns=[col])
        desc_col = "_merged_description_"
    
    # Проверяем, есть ли несколько колонок с qty/количество (из разных файлов)
    possible_qty_cols = [col for col in df.columns if any(
        col.startswith(prefix) for prefix in ["qty", "quantity", "количество", "кол"]
    )]
    
    if len(possible_qty_cols) > 1:
        # Объединяем несколько колонок qty в одну
        def merge_qty(row):
            for col in possible_qty_cols:
                val = row.get(col)
                if pd.notna(val):
                    try:
                        return float(val) if val != 0 or str(val).strip() == '0' else None
                    except:
                        pass
            return None
        
        df["_merged_qty_"] = df.apply(merge_qty, axis=1)
        # Удаляем старые колонки qty/количество
        for col in possible_qty_cols:
            if col in df.columns:
                df = df.drop(columns=[col])
        qty_col = "_merged_qty_"

    # Ensure we have at least some text to classify against
    if not any([ref_col, desc_col, value_col, part_col]):
        # create a synthetic description column as a fallback using all columns joined
        df["_row_text_"] = df.apply(lambda r: " ".join(str(x) for x in r.values if pd.notna(x)), axis=1)
        desc_col = "_row_text_"

    def run_classification(input_df: pd.DataFrame) -> pd.DataFrame:
        categories_local: List[str] = []
        for _, row in input_df.iterrows():
            ref = row.get(ref_col) if ref_col else None
            desc = row.get(desc_col) if desc_col else None
            val = row.get(value_col) if value_col else None
            part = row.get(part_col) if part_col else None
            src_file = row.get('source_file') if 'source_file' in input_df.columns else None
            note_val = row.get('note') if 'note' in input_df.columns else None
            categories_local.append(classify_row(ref, desc, val, part, strict=not args.loose, source_file=src_file, note=note_val))
        input_df = input_df.copy()
        input_df["category"] = categories_local
        return input_df

    df = run_classification(df)

    # Interactive reassignment for unclassified
    if args.interactive:
        cat_names = [
            ("resistors", "Резисторы"),
            ("capacitors", "Конденсаторы"),
            ("inductors", "Дроссели"),
            ("ics", "Микросхемы"),
            ("connectors", "Разъемы"),
            ("dev_boards", "Отладочные платы"),
            ("diods", "Диоды"),
            ("our_developments", "Наши разработки"),
            ("others", "Другие"),
            ("unclassified", "Не распределено"),
        ]
        uncls = df[df["category"] == "unclassified"].copy()
        max_preview = min(len(uncls), 50)
        print(f"Нераспределено: {len(uncls)}. Покажу первые {max_preview} для разметки.")
        # Load existing rules (append new choices)
        existing_rules: List[Dict[str, Any]] = []
        if args.assign_json and os.path.exists(args.assign_json):
            try:
                with open(args.assign_json, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    if isinstance(data, list):
                        existing_rules = data
            except Exception:
                pass
        for idx, (_, row) in enumerate(uncls.head(max_preview).iterrows(), start=1):
            text_blob = " ".join(str(x) for x in [row.get(desc_col), row.get(value_col), row.get(part_col)] if pd.notna(x))
            print(f"[{idx}] {text_blob}")
            for i, (_, ru) in enumerate(cat_names, start=1):
                print(f"  {i}. {ru}")
            choice = input("Выберите номер категории (Enter чтобы пропустить): ").strip()
            if choice.isdigit():
                ci = int(choice)
                if 1 <= ci <= len(cat_names):
                    selected_key = cat_names[ci - 1][0]
                    df.loc[uncls.index[idx - 1], "category"] = selected_key
                    # Persist rule by 'contains' text blob
                    rule = {"contains": text_blob[:160], "category": selected_key}
                    existing_rules.append(rule)
        # Save updated rules if any
        try:
            if args.assign_json:
                with open(args.assign_json, "w", encoding="utf-8") as f:
                    json.dump(existing_rules, f, ensure_ascii=False, indent=2)
                print(f"Сохранил правила: {args.assign_json}")
        except Exception as exc:
            print(f"Не удалось сохранить правила: {exc}")
        # Optionally redo classification for remaining items if needed (skipped for simplicity)

    # Объединяем категории для "Отладочные модули"
    debug_modules_parts = []
    
    # 1. Наши разработки
    our_dev = df[df["category"] == "our_developments"]
    if not our_dev.empty:
        debug_modules_parts.append(our_dev)
    
    # 2. Пустая строка
    if debug_modules_parts:
        empty_row = pd.DataFrame([{col: '' for col in df.columns}])
        debug_modules_parts.append(empty_row)
    
    # 3. Отладочные платы
    dev_boards = df[df["category"] == "dev_boards"]
    if not dev_boards.empty:
        debug_modules_parts.append(dev_boards)
    
    # 4. Пустая строка
    if len(debug_modules_parts) > 0 and not dev_boards.empty:
        empty_row2 = pd.DataFrame([{col: '' for col in df.columns}])
        debug_modules_parts.append(empty_row2)
    
    # 5. СВЧ модули
    rf_mods = df[df["category"] == "rf_modules"]
    if not rf_mods.empty:
        debug_modules_parts.append(rf_mods)
    
    # Объединяем все части
    debug_modules_combined = pd.concat(debug_modules_parts, ignore_index=True) if debug_modules_parts else pd.DataFrame()
    
    # Новый порядок категорий согласно требованиям
    outputs = {
        "debug_modules": debug_modules_combined,  # Отладочные модули (объединенная)
        "ics": df[df["category"] == "ics"],  # Микросхемы
        "resistors": df[df["category"] == "resistors"],  # Резисторы
        "capacitors": df[df["category"] == "capacitors"],  # Конденсаторы
        "inductors": df[df["category"] == "inductors"],  # Индуктивности
        "semiconductors": df[df["category"] == "semiconductors"],  # Полупроводники
        "connectors": df[df["category"] == "connectors"],  # Разъемы
        "optics": df[df["category"] == "optics"],  # Оптические компоненты
        "power_modules": df[df["category"] == "power_modules"],  # Модули питания
        "cables": df[df["category"] == "cables"],  # Кабели
        "others": df[df["category"] == "others"],  # Другие
        "unclassified": df[df["category"] == "unclassified"],  # Не распределено
    }

    # Apply assignment rules if provided
    if args.assign_json and os.path.exists(args.assign_json):
        try:
            with open(args.assign_json, "r", encoding="utf-8") as f:
                rules = json.load(f)
        except Exception as exc:
            print(f"Failed to read assign rules: {exc}")
            rules = []
        if isinstance(rules, list):
            df = df.copy()
            lower_cols = set(df.columns)
            for i, rule in enumerate(rules, start=1):
                cat = str(rule.get("category", "")).strip()
                contains = str(rule.get("contains", "")).strip().lower()
                regex = rule.get("regex")
                if not cat or (not contains and not regex):
                    continue
                mask = df["category"] == "unclassified"
                if contains:
                    blob = (
                        df.get("description", pd.Series([""] * len(df))).astype(str).str.lower().fillna("") + " " +
                        df.get("value", pd.Series([""] * len(df))).astype(str).str.lower().fillna("") + " " +
                        df.get("part", pd.Series([""] * len(df))).astype(str).str.lower().fillna("") + " " +
                        df.get("reference", pd.Series([""] * len(df))).astype(str).str.lower().fillna("")
                    )
                    mask = mask & blob.str.contains(re.escape(contains), na=False)
                if regex:
                    try:
                        r = re.compile(regex, re.IGNORECASE)
                        text_series = (
                            df.get("description", "").astype(str).fillna("") + " " +
                            df.get("value", "").astype(str).fillna("") + " " +
                            df.get("part", "").astype(str).fillna("") + " " +
                            df.get("reference", "").astype(str).fillna("")
                        )
                        mask = mask & text_series.apply(lambda t: bool(r.search(t)))
                    except Exception:
                        pass
                df.loc[mask, "category"] = cat
            # refresh outputs
            outputs = {k: df[df["category"] == k] for k in outputs.keys()}

    # Enrich each output with MR code and total quantity per item
    def enrich_with_mr_and_total(frame: pd.DataFrame) -> pd.DataFrame:
        enriched = frame.copy()
        # MR code column
        if mr_col and mr_col in enriched.columns:
            mr_series = enriched[mr_col].fillna("-")
        else:
            mr_series = pd.Series(["-"] * len(enriched), index=enriched.index)
        enriched["Код МР"] = mr_series.astype(str)

        # quantity per row
        if qty_col and qty_col in enriched.columns and pd.api.types.is_numeric_dtype(enriched[qty_col]):
            qty_series = enriched[qty_col]
        else:
            qty_series = pd.Series([1] * len(enriched), index=enriched.index)

        # grouping keys to compute total quantity per item
        group_keys: List[str] = []
        if mr_col and mr_col in enriched.columns and (enriched[mr_col].notna().any()):
            group_keys = [mr_col]
        else:
            for cand in [part_col, value_col, desc_col]:
                if cand and cand in enriched.columns:
                    group_keys.append(cand)
        if not group_keys:
            group_keys = ["category"]

        tmp = enriched.copy()
        tmp["__qty__"] = qty_series
        totals = tmp.groupby(group_keys, dropna=False)["__qty__"].sum().reset_index().rename(columns={"__qty__": "Общее количество"})
        enriched = enriched.merge(totals, on=group_keys, how="left")
        # fill missing
        if "Общее количество" not in enriched.columns:
            enriched["Общее количество"] = 0
        return enriched

    outputs = {name: enrich_with_mr_and_total(df_part) for name, df_part in outputs.items()}

    # Russian category names for output
    rus_sheet_names = {
        "debug_modules": "Отладочные модули",  # Объединенная категория
        "resistors": "Резисторы",
        "capacitors": "Конденсаторы",
        "inductors": "Индуктивности",  # Переименовано
        "ics": "Микросхемы",
        "connectors": "Разъемы",
        "dev_boards": "Отладочные платы",
        "optics": "Оптические компоненты",
        "rf_modules": "СВЧ модули",
        "cables": "Кабели",
        "power_modules": "Модули питания",
        "semiconductors": "Полупроводники",
        "our_developments": "Наши разработки",
        "others": "Другие",
        "unclassified": "Не распределено",
    }

    # Optional CSVs
    if args.out:
        for name, part_df in outputs.items():
            out_path = os.path.join(args.out, f"{name}.csv")
            save_df = part_df.copy()
            inverse_rename = {v: k for k, v in rename_map.items()}
            save_df = save_df.rename(columns=inverse_rename)
            save_df.to_csv(out_path, index=False, encoding="utf-8-sig")


    # XLSX with sheets
    def merge_existing_if_needed(writer, new_outputs: Dict[str, pd.DataFrame], existing_path: Optional[str]):
        if not existing_path or not os.path.exists(existing_path):
            # just write new_outputs as-is
            for key, part_df in new_outputs.items():
                save_df = part_df.copy()
                inverse_rename = {v: k for k, v in rename_map.items()}
                save_df = save_df.rename(columns=inverse_rename)
                sheet_name = rus_sheet_names.get(key, key)[:31]
                if save_df.empty:
                    save_df = pd.DataFrame(columns=list(inverse_rename.values()) + ["category", "source_file", "source_sheet", "Код МР", "Общее количество"])  # type: ignore
                save_df.to_excel(writer, sheet_name=sheet_name, index=False)
            return

        # merge with existing workbook
        try:
            xls = pd.ExcelFile(existing_path, engine="openpyxl")
        except Exception:
            # if cannot read, fallback to new write
            for key, part_df in new_outputs.items():
                save_df = part_df.copy()
                inverse_rename = {v: k for k, v in rename_map.items()}
                save_df = save_df.rename(columns=inverse_rename)
                sheet_name = rus_sheet_names.get(key, key)[:31]
                if save_df.empty:
                    save_df = pd.DataFrame(columns=list(inverse_rename.values()) + ["category", "source_file", "source_sheet", "Код МР", "Общее количество"])  # type: ignore
                save_df.to_excel(writer, sheet_name=sheet_name, index=False)
            return

        for key, part_df in new_outputs.items():
            inverse_rename = {v: k for k, v in rename_map.items()}
            new_df = part_df.rename(columns=inverse_rename).copy()
            sheet_name = rus_sheet_names.get(key, key)[:31]
            if sheet_name in xls.sheet_names:
                try:
                    old_df = pd.read_excel(existing_path, sheet_name=sheet_name, engine="openpyxl")
                    combined = pd.concat([old_df, new_df], ignore_index=True, sort=False)
                except Exception:
                    combined = new_df
            else:
                combined = new_df
            if combined.empty:
                combined = pd.DataFrame(columns=list(inverse_rename.values()) + ["category", "source_file", "source_sheet", "Код МР", "Общее количество"])  # type: ignore
            combined.to_excel(writer, sheet_name=sheet_name, index=False)

    def clean_component_name(row, desc_col_name):
        """Очистить наименование компонента и извлечь ТУ и тип компонента"""
        import re
        
        if desc_col_name not in row or pd.isna(row[desc_col_name]):
            return '', '', ''
        
        text = str(row[desc_col_name])
        original_text = text
        
        # Проверить note на наличие ТУ и типа из заголовка группы
        note_tu = ''
        note_type = ''
        
        if 'note' in row and pd.notna(row['note']) and ' | ' in str(row['note']):
            note_parts = str(row['note']).split(' | ')
            for part in note_parts:
                part = part.strip()
                # Если содержит ТУ - это код ТУ (запомним для fallback)
                if 'ТУ' in part:
                    note_tu = part
                # Если начинается с "Набор" или других типов - это тип компонента (запомним для fallback)
                elif any(part.startswith(t) for t in ['Набор', 'Резистор', 'Конденсатор', 'Микросхем']):
                    note_type = part
        
        # ПРИОРИТЕТ 1: Искать ТУ в тексте компонента (собственное ТУ)
        tu_code = ''
        tu_found_in_text = False
        
        # Вариант 1: ТУ с опциональным суффиксом типа "/Д6", "/02" (например, "И93.456.000 ТУ/Д6")
        tu_pattern1 = r'([А-ЯЁ]+[\d\.]+\s*ТУ[/А-ЯЁ\d]*)'
        tu_match1 = re.search(tu_pattern1, text)
        if tu_match1:
            tu_code = tu_match1.group(1)
            tu_found_in_text = True
            text = text.replace(tu_code, '').strip()
        else:
            # Вариант 2: ТУ с точками и дефисами (например, "АЕНВ.431320.515-01ТУ", "АЛЯР.434110.005ТУ")
            tu_pattern2 = r'([А-ЯЁ]{2,}\.\d+[\d\.\-]*ТУ)'
            tu_match2 = re.search(tu_pattern2, text)
            if tu_match2:
                tu_code = tu_match2.group(1)
                tu_found_in_text = True
                text = text.replace(tu_code, '').strip()
            else:
                # Вариант 3: ТУ в начале (например, "ТУ 6329-019-07614320-99")
                tu_pattern3 = r'ТУ\s+([\d\-]+)'
                tu_match3 = re.search(tu_pattern3, text)
                if tu_match3:
                    tu_code = 'ТУ ' + tu_match3.group(1)
                    tu_found_in_text = True
                    text = text.replace(tu_code, '').strip()
                else:
                    # Вариант 4: Без дефиса но с точками (например, "СМ3.362.805ТУ")
                    tu_pattern4 = r'([А-ЯЁ]{2,}[\d\.]+ТУ)'
                    tu_match4 = re.search(tu_pattern4, text)
                    if tu_match4:
                        tu_code = tu_match4.group(1)
                        tu_found_in_text = True
                        text = text.replace(tu_code, '').strip()
                    else:
                        # Вариант 5: Общий паттерн (на случай других форматов)
                        tu_pattern5 = r'([А-ЯЁ]{2,}[\d\.]+[А-ЯЁ]{0,3})'
                        tu_match5 = re.search(tu_pattern5, text)
                        if tu_match5:
                            tu_code = tu_match5.group(1)
                            tu_found_in_text = True
                            text = text.replace(tu_code, '').strip()
        
        # ПРИОРИТЕТ 2: Если ТУ не найдено в тексте, использовать ТУ из note (заголовка группы)
        if not tu_found_in_text and note_tu:
            tu_code = note_tu
        
        # Извлечь и удалить типы компонентов в начале
        # ПРИОРИТЕТ 1: Тип из самого компонента (в начале описания)
        component_type = ''
        component_types = ['Резистор', 'Конденсатор', 'Микросхема', 'Разъем', 'Диод', 
                          'Транзистор', 'Индуктивность', 'Дроссель', 'Кабель', 'Модуль',
                          'Стабилитрон', 'Вилка', 'Розетка', 'Генератор', 'Транзисторная матрица']
        # Сортируем по длине (от самых длинных к самым коротким), чтобы сначала проверять более специфичные типы
        component_types_sorted = sorted(component_types, key=len, reverse=True)
        type_found_in_text = False
        for comp_type in component_types_sorted:
            if original_text.startswith(comp_type):
                component_type = comp_type
                type_found_in_text = True
                text = text[len(comp_type):].strip() if text.startswith(comp_type) else text
                break
        
        # ПРИОРИТЕТ 2: Если тип не найден в тексте, использовать тип из note (заголовка группы)
        # НО ТОЛЬКО если это тип "Набор ..." (не применяем тип заголовка к компонентам с явным типом)
        if not type_found_in_text and note_type and note_type.startswith('Набор'):
            component_type = note_type
        
        # Очистить от лишних пробелов
        text = ' '.join(text.split())
        
        return text, tu_code, component_type
    
    def extract_nominal_value(text):
        """
        Извлечь номинал из текста и преобразовать в число для сортировки.
        Примеры: "180 Ом" -> 180, "1 кОм" -> 1000, "100 пФ" -> 100e-12
        Также парсит SMD коды: GRM1555C1H102J (конденсатор 1 нФ), AC0402JR-0710KL (резистор 710 кОм)
        """
        import re
        
        if not isinstance(text, str):
            return 0
        
        # 1. Попробовать парсить SMD код резистора (например AC0402JR-0710KL)
        # Формат: цифры + R/K/M (R=Ом, K=кОм, M=МОм)
        resistor_smd_pattern = r'-0*(\d+(?:\.\d+)?)\s*([RKM])(?:\d|L)'
        resistor_match = re.search(resistor_smd_pattern, text.upper())
        if resistor_match:
            value = float(resistor_match.group(1))
            unit = resistor_match.group(2)
            if unit == 'R':
                return value  # Ом
            elif unit == 'K':
                return value * 1e3  # кОм
            elif unit == 'M':
                return value * 1e6  # МОм
        
        # 1a. Попробовать парсить SMD код индуктивности (например 0603HP-47NXJ_LW, 1206CS-821XJB)
        # Формат 1: цифры-буква-цифры (например 3N3 = 3.3 нГн, 47N = 47 нГн)
        inductor_pattern1 = r'-(\d+)N(\d+)(?:X|_)'
        inductor_match1 = re.search(inductor_pattern1, text.upper())
        if inductor_match1:
            # Формат XNY означает X.Y нГн
            value = float(inductor_match1.group(1) + '.' + inductor_match1.group(2))
            return value * 1e-9  # нГн в Генри
        
        # Формат 2: просто цифры+N (например 47N = 47 нГн)
        inductor_pattern2 = r'-(\d+)N(?:X|J)'
        inductor_match2 = re.search(inductor_pattern2, text.upper())
        if inductor_match2:
            value = float(inductor_match2.group(1))
            return value * 1e-9  # нГн в Генри
        
        # Формат 3: трёхзначный код (например 821 = 82 × 10^1 = 820 мкГн)
        inductor_pattern3 = r'-(\d{3})(?:X|J)'
        inductor_match3 = re.search(inductor_pattern3, text.upper())
        if inductor_match3:
            code = inductor_match3.group(1)
            mantissa = int(code[:2])
            exponent = int(code[2])
            uh_value = mantissa * (10 ** exponent)  # В микрогенри
            return uh_value * 1e-6  # мкГн в Генри
        
        # 2. Попробовать парсить SMD код конденсатора (например GRM1555C1H102J, NFM21CC102R1H3D)
        # Последние 3 цифры перед буквой - это код емкости (первые 2 цифры × 10^последняя в пФ)
        # Ищем ВСЕ вхождения 3 цифр перед буквой, берем последнее
        cap_smd_pattern = r'(\d{3})(?=[A-Z])'
        cap_matches = re.findall(cap_smd_pattern, text.upper())
        if cap_matches:
            # Берем последнее вхождение
            code = cap_matches[-1]
            if code.isdigit() and len(code) == 3:
                # Первые 2 цифры - значение, последняя - степень десяти
                mantissa = int(code[:2])
                exponent = int(code[2])
                pf_value = mantissa * (10 ** exponent)  # В пикофарадах
                return pf_value * 1e-12  # Преобразуем в фарады для сортировки
        
        # 3. Множители для разных единиц (для текстовых описаний)
        multipliers = {
            # Сопротивление
            'мом': 1e6, 'мегаом': 1e6,
            'ком': 1e3, 'килоом': 1e3,
            'ом': 1,
            # Емкость
            'ф': 1, 'фарад': 1,
            'мф': 1e-3, 'миллифарад': 1e-3,
            'мкф': 1e-6, 'микрофарад': 1e-6, 'µф': 1e-6,
            'нф': 1e-9, 'нанофарад': 1e-9,
            'пф': 1e-12, 'пикофарад': 1e-12,
            # Индуктивность
            'гн': 1, 'генри': 1,
            'мгн': 1e-3, 'миллигенри': 1e-3,
            'мкгн': 1e-6, 'микрогенри': 1e-6, 'µгн': 1e-6,
        }
        
        # Попытаться найти число с единицей измерения
        # Паттерн: число (с возможной точкой/запятой) + пробелы + единица
        pattern = r'([\d.,]+)\s*([а-яА-Яa-zA-Z]+)'
        matches = re.findall(pattern, text.lower())
        
        # Ищем значение с правильной единицей измерения (приоритет: пФ, нФ, мкФ, мФ, Ф)
        # Сначала ищем единицы емкости, потом сопротивления, потом индуктивности
        priority_units = ['пф', 'нф', 'мкф', 'мф', 'ф', 'ом', 'ком', 'мом', 'гн', 'мгн', 'мкгн']
        
        for priority_unit in priority_units:
            for num_str, unit in matches:
                if priority_unit in unit.lower():
                    try:
                        num = float(num_str.replace(',', '.'))
                        # Найти множитель для единицы
                        # ВАЖНО: проверяем от самых ДЛИННЫХ к КОРОТКИМ чтобы "пф" не схватило "ф"
                        sorted_multipliers = sorted(multipliers.items(), key=lambda x: len(x[0]), reverse=True)
                        for unit_key, mult in sorted_multipliers:
                            if unit_key in unit.lower():
                                return num * mult
                    except ValueError:
                        continue
        
        return 0

    with pd.ExcelWriter(args.xlsx, engine="openpyxl") as writer:
        # Записать каждую категорию с № п/п
        for key, part_df in outputs.items():
            sheet_name = rus_sheet_names.get(key, key)[:31]
            result_df = part_df.copy()
            
            # Удалить ВСЕ старые колонки с номерами (могут быть варианты написания)
            cols_to_remove = [col for col in result_df.columns 
                            if col in ['№ п/п', '№ п\п', 'п/п', 'номер', '№']]
            if cols_to_remove:
                result_df = result_df.drop(columns=cols_to_remove)
            
            # НЕ удаляем 'note' здесь - он нужен для clean_component_name!
            # Удалить ненужные столбцы (технические, служебные, цена, стоимость)
            cols_to_remove = [col for col in result_df.columns 
                            if col in ['количество', 'первоначальная цена, тыс.руб.', 
                                      'первоначальная стоимость, тыс.руб.', 
                                      'source_sheet', '_row_text_', 'category',
                                      'zone', 'reference', 'qty', 'value', 
                                      'партномер', 'partname', 'part', 'примечания',
                                      'номинал', 'тип', 'позиционное обозначение']]
            if cols_to_remove:
                result_df = result_df.drop(columns=cols_to_remove)
            
            # Переименовать столбцы
            if 'Общее количество' in result_df.columns:
                result_df = result_df.rename(columns={'Общее количество': 'Кол-во'})
            if 'наименование ивп' in result_df.columns:
                result_df = result_df.rename(columns={'наименование ивп': 'Наименование ИВП'})
            # Переименовать нормализованные английские колонки в русские
            if '_merged_description_' in result_df.columns and 'Наименование ИВП' not in result_df.columns:
                result_df = result_df.rename(columns={'_merged_description_': 'Наименование ИВП'})
            elif 'description' in result_df.columns and 'Наименование ИВП' not in result_df.columns:
                result_df = result_df.rename(columns={'description': 'Наименование ИВП'})
            if 'qty' in result_df.columns and 'Кол-во' not in result_df.columns:
                result_df = result_df.rename(columns={'qty': 'Кол-во'})
            
            # Обработать наименование - очистить и извлечь ТУ
            # Найти столбец с наименованием (ВАЖНО: сначала ищем переименованный столбец!)
            desc_col_name = None
            for possible_name in ['Наименование ИВП', 'наименование ивп', 'описание', 'наименование', desc_col]:
                if possible_name and possible_name in result_df.columns:
                    desc_col_name = possible_name
                    break
            
            if desc_col_name:
                # Применить функцию очистки к каждой строке
                cleaned_data = result_df.apply(lambda row: clean_component_name(row, desc_col_name), axis=1)
                result_df[desc_col_name] = [item[0] for item in cleaned_data]
                
                # Для "Наших разработок" - если название пустое, взять из source_file
                if sheet_name == 'Наши разработки' and 'source_file' in result_df.columns:
                    for idx in result_df.index:
                        if not result_df.loc[idx, desc_col_name] or pd.isna(result_df.loc[idx, desc_col_name]) or str(result_df.loc[idx, desc_col_name]).strip() == '':
                            source_file = result_df.loc[idx, 'source_file']
                            if source_file and pd.notna(source_file):
                                # Извлечь название файла без расширения
                                file_name = os.path.splitext(os.path.basename(str(source_file)))[0]
                                result_df.loc[idx, desc_col_name] = file_name
                
                # Вставить ТУ сразу после наименования (а не в конец)
                tu_data = [item[1] for item in cleaned_data]
                desc_idx = list(result_df.columns).index(desc_col_name)
                result_df.insert(desc_idx + 1, 'ТУ', tu_data)
                
                # Вставить Примечание после ТУ
                component_types = [item[2] for item in cleaned_data]
                
                # Определить стандартный тип для категории
                category_standard_types = {
                    'Резисторы': 'Резистор',
                    'Конденсаторы': 'Конденсатор',
                    'Индуктивности': 'Дроссель',  # Переименовано
                    'Микросхемы': 'Микросхема',
                    'Разъемы': 'Разъем',
                    'Полупроводники': '',  # Нет стандартного типа - тут разные (диоды, транзисторы, стабилитроны)
                }
                
                standard_type = category_standard_types.get(sheet_name, '')
                
                # Если тип компонента совпадает со стандартным для категории - ставим "-"
                primechanie = []
                for comp_type in component_types:
                    if not comp_type or comp_type == standard_type:
                        primechanie.append('-')
                    else:
                        primechanie.append(comp_type)
                
                result_df.insert(desc_idx + 2, 'Примечание', primechanie)
                
                # Сортировка зависит от категории
                # Для конденсаторов, дросселей - сортируем ВСЕ по номиналу
                # Для микросхем - сначала импортные, потом отечественные по первым цифрам
                
                if sheet_name in ['Конденсаторы', 'Дроссели', 'Резисторы']:
                    # Сортировка только по номиналу для всех компонентов
                    result_df['_nominal'] = result_df[desc_col_name].apply(extract_nominal_value)
                    result_df = result_df.sort_values(
                        by=['_nominal', desc_col_name],
                        ascending=[True, True]
                    )
                    result_df = result_df.drop(columns=['_nominal'])
                else:
                    # Сортировка: сначала импортные, потом отечественные
                    result_df['_is_domestic'] = result_df['ТУ'].apply(
                        lambda x: 1 if (x and x != '-' and str(x).strip() != '') else 0
                    )
                    
                    # Для отечественных: извлечь первые цифры из названия
                    def get_domestic_number(row):
                        if row['_is_domestic'] == 1:
                            import re
                            name = str(row[desc_col_name])
                            match = re.match(r'(\d+)', name)
                            if match:
                                return int(match.group(1))
                            else:
                                return 999999
                        else:
                            return 0
                    
                    result_df['_domestic_num'] = result_df.apply(get_domestic_number, axis=1)
                    
                    # Для импортных: сортировка по номиналу
                    result_df['_nominal'] = result_df[desc_col_name].apply(extract_nominal_value)
                    
                    # Сортируем: сначала по типу (импортные/отечественные), потом по номиналу/номеру, потом по имени
                    result_df = result_df.sort_values(
                        by=['_is_domestic', '_domestic_num', '_nominal', desc_col_name],
                        ascending=[True, True, True, True]
                    )
                    
                    result_df = result_df.drop(columns=['_is_domestic', '_domestic_num', '_nominal'])
                
                result_df = result_df.reset_index(drop=True)
            
            # Добавить новую колонку № п/п в начало (ПОСЛЕ сортировки!)
            result_df.insert(0, '№ п/п', range(1, len(result_df) + 1))
            
            # Упорядочить колонки в правильном порядке
            # Стандартный порядок: № п/п, Наименование ИВП, ТУ, Примечание, source_file, Код МР, Кол-во
            desired_order = ['№ п/п', 'Наименование ИВП', 'ТУ', 'Примечание', 'source_file', 'Код МР', 'Кол-во']
            
            # Сначала добавляем колонки в нужном порядке (если они есть)
            ordered_cols = [col for col in desired_order if col in result_df.columns]
            
            # Затем добавляем оставшиеся колонки (если вдруг есть)
            remaining_cols = [col for col in result_df.columns if col not in ordered_cols]
            
            # Финальный порядок
            final_cols = ordered_cols + remaining_cols
            result_df = result_df[final_cols]
            
            # Удалить 'note' перед записью (если он остался)
            if 'note' in result_df.columns:
                result_df = result_df.drop(columns=['note'])
            
            result_df.to_excel(writer, sheet_name=sheet_name, index=False)

        if args.combine:
            # Не создаем SUMMARY - он не нужен или создаем простую сводку
            summary_rows = []
            for key, part_df in outputs.items():
                if len(part_df) == 0:
                    continue
                category_name = rus_sheet_names.get(key, key)
                # Суммировать количество по категории
                total_qty = part_df['Общее количество'].sum() if 'Общее количество' in part_df.columns else len(part_df)
                summary_rows.append({
                    '№ п/п': len(summary_rows) + 1,
                    'Категория': category_name,
                    'Кол-во позиций': len(part_df),
                    'Общее количество': int(total_qty)
                })
            
            summary = pd.DataFrame(summary_rows)
            summary.to_excel(writer, sheet_name="SUMMARY", index=False)
        # SOURCES sheet
        sources = pd.DataFrame(sorted({(r.get("source_file", ""), r.get("source_sheet", "")) for _, r in df.iterrows()}), columns=["source_file", "source_sheet"])
        sources.to_excel(writer, sheet_name="SOURCES", index=False)
        
        # Применить форматирование к каждому листу
        from openpyxl.styles import Alignment
        
        for sheet_name in writer.book.sheetnames:
            ws = writer.book[sheet_name]
            
            # Найти индексы столбцов "Наименование ИВП" и "ТУ"
            desc_col_idx = None
            tu_col_idx = None
            for idx, cell in enumerate(ws[1], start=1):
                cell_val = str(cell.value).lower() if cell.value else ''
                if 'наименование ивп' in cell_val or 'наименование' in cell_val:
                    desc_col_idx = idx
                elif cell_val == 'ту':
                    tu_col_idx = idx
            
            # Центрировать все ячейки, кроме "наименование ивп" и "ТУ"
            for row_idx, row in enumerate(ws.iter_rows(), start=1):
                for col_idx, cell in enumerate(row, start=1):
                    if col_idx == desc_col_idx or col_idx == tu_col_idx:
                        # Наименование ИВП и ТУ - выравнивание по левому краю
                        cell.alignment = Alignment(horizontal='left', vertical='center')
                    else:
                        # Все остальные - по центру
                        cell.alignment = Alignment(horizontal='center', vertical='center')
            
            # Автоматически установить ширину столбцов по содержимому
            for column in ws.columns:
                max_length = 0
                column_letter = column[0].column_letter
                
                for cell in column:
                    try:
                        cell_value = str(cell.value) if cell.value is not None else ""
                        if len(cell_value) > max_length:
                            max_length = len(cell_value)
                    except:
                        pass
                
                # Установить ширину с небольшим запасом
                adjusted_width = min(max_length + 2, 50)  # максимум 50 символов
                ws.column_dimensions[column_letter].width = adjusted_width

    # Optional TXT files per category - создаем ПОСЛЕ записи Excel из обработанных данных
    if args.txt_dir:
        os.makedirs(args.txt_dir, exist_ok=True)
        # Читаем только что созданный Excel файл с обработанными данными
        for key, _ in outputs.items():
            rus_name = rus_sheet_names.get(key, key)
            sheet_name = rus_sheet_names.get(key, key)[:31]
            
            try:
                # Читаем обработанные данные из Excel
                processed_df = pd.read_excel(args.xlsx, sheet_name=sheet_name, engine='openpyxl')
                if processed_df.empty:
                    continue
                    
                txt_path = os.path.join(args.txt_dir, f"{rus_name}.txt")
                with open(txt_path, "w", encoding="utf-8") as f:
                    f.write(f"=== {rus_name.upper()} ===\n")
                    f.write(f"Всего элементов: {len(processed_df)}\n")
                    f.write("=" * 80 + "\n\n")
                    
                    for num, (_, row) in enumerate(processed_df.iterrows(), start=1):
                        # Используем обработанные столбцы
                        name_val = row.get('Наименование ИВП', '')
                        tu_val = row.get('ТУ', '')
                        qty_val = row.get('Кол-во', '')
                        mr_val = row.get('Код МР', '')
                        
                        # Форматируем вывод
                        f.write(f"{num}. ")
                        
                        if pd.notna(name_val) and str(name_val).strip():
                            f.write(f"{name_val}")
                        
                        if tu_val and pd.notna(tu_val) and str(tu_val).strip() and str(tu_val) != '-' and str(tu_val).lower() != 'nan':
                            f.write(f" | ТУ: {tu_val}")
                        
                        if pd.notna(qty_val):
                            try:
                                qty_int = int(float(qty_val))
                                f.write(f" | Кол-во: {qty_int} шт")
                            except (ValueError, TypeError):
                                pass
                        
                        if pd.notna(mr_val) and str(mr_val).strip() and str(mr_val) != "-":
                            f.write(f" | Код МР: {mr_val}")
                        
                        f.write("\n")
                    
                    f.write("\n" + "=" * 80 + "\n")
                    f.write(f"Всего записей: {len(processed_df)}\n")
            except Exception as e:
                print(f"Warning: Could not create TXT for {rus_name}: {e}")
        
        print(f"TXT files written to: {args.txt_dir}")
    
    # Print counts
    print("Split complete:")
    for name, part_df in outputs.items():
        print(f"  {name}: {len(part_df)}")
    print(f"XLSX written: {args.xlsx}")


def build_summary(df: pd.DataFrame, ref_col: Optional[str], desc_col: Optional[str], value_col: Optional[str], part_col: Optional[str], qty_col: Optional[str], mr_col: Optional[str]) -> pd.DataFrame:
    work = df.copy()
    if qty_col and qty_col in work.columns and pd.api.types.is_numeric_dtype(work[qty_col]):
        work["_qty_"] = work[qty_col]
    else:
        work["_qty_"] = 1

    group_keys: List[str] = []
    if mr_col and mr_col in work.columns and (work[mr_col].notna().any()):
        group_keys = [mr_col]
    else:
        for cand in [part_col, value_col, desc_col]:
            if cand and cand in work.columns:
                group_keys.append(cand)
    if not group_keys:
        group_keys = ["category"]

    summary = work.groupby(group_keys, dropna=False, as_index=False)["_qty_"].sum()
    summary = summary.rename(columns={"_qty_": "Общее количество"})

    # Add MR code column
    if mr_col and mr_col in work.columns:
        mr_map = work.groupby(group_keys)[mr_col].agg(lambda s: next((x for x in s if pd.notna(x)), None)).reset_index()
        summary = summary.merge(mr_map, on=group_keys, how="left")
        summary = summary.rename(columns={mr_col: "Код МР"})
    else:
        summary["Код МР"] = "-"

    cols = list(summary.columns)
    if "category" in cols:
        cols = ["category"] + [c for c in cols if c != "category"]
        summary = summary[cols]
    return summary


if __name__ == "__main__":
    main()


