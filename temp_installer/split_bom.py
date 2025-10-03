import argparse
import os
import sys
import re
import json
from typing import List, Optional, Dict, Any, Iterable, Tuple, Union

import pandas as pd

try:
    from docx import Document  # python-docx
except Exception:
    Document = None  # optional; raise if used without installed


def normalize_column_names(columns: List[str]) -> List[str]:
    normalized = []
    for name in columns:
        if name is None:
            normalized.append("")
            continue
        normalized.append(str(name).strip().lower())
    return normalized


def find_column(possible_names: List[str], columns: List[str]) -> Optional[str]:
    for candidate in possible_names:
        if candidate in columns:
            return candidate
    return None


def has_any(text: str, keywords: List[str]) -> bool:
    if not isinstance(text, str):
        return False
    lower = text.lower()
    return any(k in lower for k in keywords)


RESISTOR_VALUE_RE = re.compile(r"(?i)\b\d+(?:[\.,]\d+)?\s*(?:ом|ohm|k\s*ohm|kohm|к\s*ом|ком|m\s*ohm|mohm|м\s*ом|мом)\b")
CAP_VALUE_RE = re.compile(r"(?i)\b\d+(?:[\.,]\d+)?\s*(?:pf|nf|uf|µf|μf|ф|пф|нф|мкф)\b")
IND_VALUE_RE = re.compile(r"(?i)\b\d+(?:[\.,]\d+)?\s*(?:nh|uh|µh|μh|mh|h|нгн|мкгн|мгн|гн)\b")


# NEW: lightweight parser for TXT/DOCX tables and lines
LINE_SPLIT_RE = re.compile(r"\s{2,}|\t|;|,\s?(?=\S)" )
POS_PREFIX_RE = re.compile(r"^(?:[A-ZА-Я]+\d+(?:[-,;\s]*[A-ZА-Я]*\d+)*)$", re.IGNORECASE)


def parse_txt_like(path: str) -> pd.DataFrame:
    rows: List[Dict[str, Any]] = []
    try:
        with open(path, "r", encoding="utf-8") as f:
            text = f.read()
    except UnicodeDecodeError:
        with open(path, "r", encoding="cp1251", errors="ignore") as f:
            text = f.read()
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        parts = [p.strip() for p in LINE_SPLIT_RE.split(line) if p.strip()]
        if not parts:
            continue
        # Heuristic mapping: [pos?, name/desc, qty?]
        pos = parts[0] if POS_PREFIX_RE.match(parts[0]) else None
        qty = None
        # Try explicit "Nшт" or "N шт" patterns first
        m = re.search(r"(\d+)\s*(шт\.?|pcs|pieces)", line, flags=re.IGNORECASE)
        if m:
            qty = int(m.group(1))
        else:
            for p in parts[::-1]:
                if re.fullmatch(r"\d+", p):
                    qty = int(p)
                    break
        desc = " ".join(parts[1:-1]) if pos and len(parts) >= 2 else (" ".join(parts))
        row = {"reference": pos or "", "description": desc, "qty": qty if qty is not None else 1}
        rows.append(row)
    if not rows:
        # fallback: whole text in a single row
        rows = [{"description": text, "qty": 1}]
    return pd.DataFrame(rows)


def parse_docx(path: str) -> pd.DataFrame:
    if Document is None:
        raise SystemExit("python-docx is required to parse DOCX. Install with pip install python-docx")
    doc = Document(path)
    extracted: List[Dict[str, Any]] = []

    def guess_header_index(table) -> int:
        max_scan = min(5, len(table.rows))
        header_keywords = ["наимен", "обознач", "кол.", "кол ", "кол", "примеч"]
        for i in range(max_scan):
            row_texts = [normalize_cell(c.text).lower() for c in table.rows[i].cells]
            hits = sum(any(hk in t for hk in header_keywords) for t in row_texts)
            if hits >= 2:
                return i
        return 0

    # Parse tables with header detection
    for table in doc.tables:
        if not table.rows:
            continue
        header_idx = guess_header_index(table)
        header_cells = [normalize_cell(c.text).strip() for c in table.rows[header_idx].cells]
        header_norm = [h.lower() for h in header_cells]

        # Find column indices by common names
        def find_col_idx(candidates: List[str]) -> Optional[int]:
            for i, h in enumerate(header_norm):
                for cand in candidates:
                    if cand in h:
                        return i
            return None

        idx_zone = find_col_idx(["зона"])  # optional
        idx_ref = find_col_idx(["поз", "обозн"])  # позиционное обозначение
        idx_name = find_col_idx(["наимен"])  # наименование
        idx_qty = find_col_idx(["кол.", "кол ", "кол", "количество"])  # количество
        idx_note = find_col_idx(["примеч"])  # опционально

        for tr in table.rows[header_idx + 1:]:
            vals = [normalize_cell(c.text) for c in tr.cells]
            if not any(v.strip() for v in vals):
                continue
            zone = vals[idx_zone] if idx_zone is not None and idx_zone < len(vals) else ""
            ref = vals[idx_ref] if idx_ref is not None and idx_ref < len(vals) else ""
            name = vals[idx_name] if idx_name is not None and idx_name < len(vals) else ""
            qty_raw = vals[idx_qty] if idx_qty is not None and idx_qty < len(vals) else ""
            note = vals[idx_note] if idx_note is not None and idx_note < len(vals) else ""

            # If header wasn't detected, try fallback mapping by last column digits
            if not any([ref, name, qty_raw]) and len(vals) >= 2:
                name = " ".join(vals[:-1])
                qty_raw = vals[-1]

            # parse qty
            qty = None
            m = re.search(r"(\d+)", str(qty_raw))
            if m:
                try:
                    qty = int(m.group(1))
                except Exception:
                    qty = 1

            row = {
                "zone": zone,
                "reference": ref,
                "description": name if name else note,
                "qty": qty if qty is not None else 1,
                "note": note,
            }
            extracted.append(row)

    # Additionally parse free text paragraphs (fallback)
    for p in doc.paragraphs:
        t = normalize_cell(p.text)
        if not t:
            continue
        parts = [s.strip() for s in LINE_SPLIT_RE.split(t) if s.strip()]
        if parts:
            pos = parts[0] if POS_PREFIX_RE.match(parts[0]) else ""
            m = re.search(r"(\d+)\s*(шт\.?|pcs|pieces)?", t, flags=re.IGNORECASE)
            qty = int(m.group(1)) if m else 1
            desc = " ".join(parts[1:]) if pos else " ".join(parts)
            extracted.append({"reference": pos, "description": desc, "qty": qty})

    if not extracted:
        extracted = [{"description": " ".join(normalize_cell(p.text) for p in doc.paragraphs), "qty": 1}]
    return pd.DataFrame(extracted)


def normalize_cell(s: Any) -> str:
    return (str(s or "").strip())


def classify_row(ref: Optional[str], description: Optional[str], value: Optional[str], partname: Optional[str], strict: bool) -> str:
    def to_text(x: Any) -> str:
        if x is None:
            return ""
        try:
            import math
            if isinstance(x, float) and math.isnan(x):
                return ""
        except Exception:
            pass
        s = str(x)
        return s.strip()

    ref = to_text(ref)
    desc = to_text(description)
    val = to_text(value)
    part = to_text(partname)

    # Refdes first where reliable
    ref_prefix = ref.split(" ")[0].upper() if ref else ""
    ref_prefix = re.sub(r"\d.*$", "", ref_prefix)  # take letters before digits

    # Heuristics by ref (only if we have a real ref column)
    if ref:
        if ref_prefix.startswith("R"):
            return "resistors"
        if ref_prefix.startswith("C"):
            return "capacitors"
        if ref_prefix.startswith("L"):
            return "inductors"
        if ref_prefix.startswith(("U", "DD", "DA", "IC")):
            return "ics"
        if ref_prefix.startswith(("J", "X", "P", "K", "XS", "XP", "JTAG")):
            return "connectors"

    text_blob = " ".join([desc, val, part])

    # Russian and English keywords
    if RESISTOR_VALUE_RE.search(text_blob) or has_any(text_blob, ["резист", "resistor"]):
        return "resistors"

    if CAP_VALUE_RE.search(text_blob) or has_any(text_blob, ["конденс", "capacitor", "tantalum", "ceramic"]):
        return "capacitors"

    if IND_VALUE_RE.search(text_blob) or has_any(text_blob, ["дросс", "индукт", "inductor", "ferrite", "феррит", "катушка", "choke"]):
        return "inductors"

    if has_any(text_blob, [
        "микросхем", " ic", "mcu", "контроллер", "процессор", "оп-амп", "op-amp", "opamp", "adc", "dac", "fpga",
        "asic", "драйвер ", "компаратор", "стабил", "регулятор", "transceiver", "sn74", "ti ", "stm32", "lmk", "ad9"
    ]):
        return "ics"

    if has_any(text_blob, [
        "разъем", "разъём", "connector", "header", "socket", "rj45", "rj11", "sma", "bnc", "terminal", "клемм",
        "штырь", "pin header", "fpc", "ffc", "din", "dc jack", "barrel", "штекер", "вилка", "розетка", "d-sub", "harting"
    ]):
        return "connectors"

    if has_any(text_blob, [
        "отладоч", " dev board", "evaluation", "eval", "nucleo", "arduino", "raspberry",
        "esp32", "stm32 nucleo", "breakout", "fmc", "carrier", "ultrazed", "microzed", "picozed", "zedboard",
        "zynq", "som ", "system on module", "voyager", "tinypilot"
    ]):
        return "dev_boards"

    # New categories
    if has_any(text_blob, [
        "оптичес", "лазер", "оптопара", "led ", "светодиод", "fiber", "оптоволок", "sfp", "qsfp", "transceiver module"
    ]):
        return "optics"

    if has_any(text_blob, [
        "свч", "вч ", "rf ", "microwave", "mini-circuits", "planar monolithics", "pmi", "qualwave", "ghz", "lna", "rf amp"
    ]):
        return "rf_modules"

    if has_any(text_blob, [
        "кабель", "cable", "шлейф", "провод", "wire", "patch cord", "jumper"
    ]):
        return "cables"

    if has_any(text_blob, [
        "модуль питания", "power module", "dc-dc", "ac-dc", "buck", "boost", "источник питания", "блок питания", "psu",
        "converter"
    ]):
        return "power_modules"

    # Diodes/indicators
    if has_any(text_blob, [
        "диод", "индикатор", "led ", "svetodiod", "indicator"
    ]):
        return "diods"

    # Our developments
    if has_any(text_blob, [
        "мвок", "наша разработ", "собственной разработ", "шск-м", "плата контроллера шск"
    ]):
        return "our_developments"

    # OTHER general hardware to bucket into 'others' (cabinets, bolts, shelves, keyboards etc.)
    if has_any(text_blob, [
        "rittal", "шкаф", "станция", "полка", "кронштейн", "ролик", "болт", "гайка", "шайба", "клавиатура", "моноблок",
        "кабель", "клеммная", "корпус", "шасси", "стеллаж", "стойка", "провод", "розетка", "вентилятор"
    ]):
        return "others"

    # In strict mode, avoid loose matches
    return "unclassified"


def main():
    parser = argparse.ArgumentParser(description="Split BOM Excel into category CSVs or a single XLSX with sheets")
    parser.add_argument("--input", help="Path to input XLSX file (deprecated if --inputs used)")
    parser.add_argument("--inputs", nargs="+", help="Paths to input XLSX files")
    parser.add_argument("--sheet", default=None, help="Sheet name or index (default: first if not using --sheets)")
    parser.add_argument("--sheets", default=None, help="Comma-separated sheet names or indices to process")
    parser.add_argument("--out", default=None, help="Output directory for CSVs (optional)")
    parser.add_argument("--xlsx", default="categorized.xlsx", help="Path to output XLSX with category sheets")
    parser.add_argument("--merge-into", dest="merge_into", default=None, help="Existing categorized XLSX to merge into (append)")
    parser.add_argument("--combine", action="store_true", help="Create a SUMMARY sheet with aggregated quantities")
    parser.add_argument("--interactive", action="store_true", help="Interactive classification for unclassified rows")
    parser.add_argument("--loose", action="store_true", help="Use looser heuristics (may misclassify non-electronics)")
    parser.add_argument("--assign-json", dest="assign_json", default="rules.json", help="Path to JSON rules for assigning categories to unclassified rows. Format: [{'contains': 'text', 'category': 'ics'}]")
    args = parser.parse_args()

    if not args.inputs and not args.input:
        raise SystemExit("Provide --inputs <files...> or --input <file>")

    inputs: List[str] = args.inputs or ([args.input] if args.input else [])
    selected_sheets: Optional[List[Any]] = None
    if args.sheets:
        selected_sheets = []
        for token in args.sheets.split(','):
            token = token.strip()
            if not token:
                continue
            try:
                selected_sheets.append(int(token))
            except ValueError:
                selected_sheets.append(token)

    if args.out:
        os.makedirs(args.out, exist_ok=True)

    # Load workbook
    read_kwargs = {"engine": "openpyxl"}
    all_rows: List[pd.DataFrame] = []
    for input_path in inputs:
        try:
            ext = os.path.splitext(input_path)[1].lower()
            if ext in [".txt"]:
                df_local = parse_txt_like(input_path)
                df_local["source_file"] = os.path.basename(input_path)
                df_local["source_sheet"] = "txt"
                all_rows.append(df_local)
                continue
            if ext in [".docx"]:
                df_local = parse_docx(input_path)
                df_local["source_file"] = os.path.basename(input_path)
                df_local["source_sheet"] = "docx"
                all_rows.append(df_local)
                continue
            if ext in [".doc"]:
                # Try to convert .doc -> .docx via Word COM (if available), else fallback to text
                try:
                    from win32com.client import Dispatch  # type: ignore
                    word = Dispatch("Word.Application")
                    word.Visible = False
                    doc = word.Documents.Open(os.path.abspath(input_path))
                    tmp_docx = os.path.splitext(os.path.abspath(input_path))[0] + "_conv.docx"
                    wdFormatXMLDocument = 12
                    doc.SaveAs(tmp_docx, FileFormat=wdFormatXMLDocument)
                    doc.Close(False)
                    word.Quit()
                    df_local = parse_docx(tmp_docx)
                except Exception:
                    df_local = parse_txt_like(input_path)
                df_local["source_file"] = os.path.basename(input_path)
                df_local["source_sheet"] = "doc"
                all_rows.append(df_local)
                continue

            if selected_sheets is not None:
                # read all selected sheets explicitly
                for sh in selected_sheets:
                    df_local = pd.read_excel(input_path, sheet_name=sh, **{k: v for k, v in read_kwargs.items() if k != 'sheet_name'})
                    if isinstance(df_local, dict):
                        # unlikely when specifying a single sheet, but guard anyway
                        for _, dfi in df_local.items():
                            dfi["source_file"] = os.path.basename(input_path)
                            dfi["source_sheet"] = str(sh)
                            all_rows.append(dfi)
                    else:
                        df_local["source_file"] = os.path.basename(input_path)
                        df_local["source_sheet"] = str(sh)
                        all_rows.append(df_local)
            else:
                # single sheet or first sheet
                sheet = args.sheet
                if sheet is not None:
                    try:
                        sheet = int(sheet)
                    except ValueError:
                        pass
                    read_kwargs["sheet_name"] = sheet

                df = pd.read_excel(input_path, **read_kwargs)
                if isinstance(df, dict):
                    # take first sheet
                    first_key = next(iter(df))
                    df = df[first_key]
                    src_sheet = first_key
                else:
                    src_sheet = sheet if sheet is not None else 0
                df["source_file"] = os.path.basename(input_path)
                df["source_sheet"] = str(src_sheet)
                all_rows.append(df)
        except Exception as exc:
            raise SystemExit(f"Failed to read Excel '{input_path}': {exc}")

    if not all_rows:
        raise SystemExit("No data loaded from inputs")

    df = pd.concat(all_rows, ignore_index=True)

    # Normalize columns
    original_cols = list(df.columns)
    lower_cols = normalize_column_names(original_cols)
    rename_map = {orig: norm for orig, norm in zip(original_cols, lower_cols)}
    df = df.rename(columns=rename_map)

    # Common column guesses
    ref_col = find_column(["ref", "reference", "designator", "refdes", "reference designator", "обозначение", "позиционное обозначение"], list(df.columns))
    desc_col = find_column(["description", "desc", "наименование", "имя", "item", "part", "part name", "наим."], list(df.columns))
    value_col = find_column(["value", "значение", "номинал"], list(df.columns))
    part_col = find_column(["partnumber", "mfr part", "mpn", "pn", "art", "артикул", "part", "part name"], list(df.columns))
    qty_col = find_column([
        "qty", "quantity", "количество", "кол.", "кол-во", "кол. в ктд", "кол в ктд", "кол. в спецификации", "кол. в кдт",
        "кол. в ктд", "кол. в ктд, шт", "кол. в ктд (шт)", "кол. в ктд, шт."
    ], list(df.columns))
    mr_col = find_column([
        "код мр", "код ивп", "код мр/ивп", "код позиции", "код изделия", "код мр позиции", "код мр ивп"
    ], list(df.columns))

    # Ensure we have at least some text to classify against
    if not any([ref_col, desc_col, value_col, part_col]):
        # create a synthetic description column as a fallback using all columns joined
        df["_row_text_"] = df.apply(lambda r: " ".join(str(x) for x in r.values if pd.notna(x)), axis=1)
        desc_col = "_row_text_"

    def run_classification(input_df: pd.DataFrame) -> pd.DataFrame:
        categories_local: List[str] = []
        for _, row in input_df.iterrows():
            ref = row.get(ref_col) if ref_col else None
            desc = row.get(desc_col) if desc_col else None
            val = row.get(value_col) if value_col else None
            part = row.get(part_col) if part_col else None
            categories_local.append(classify_row(ref, desc, val, part, strict=not args.loose))
        input_df = input_df.copy()
        input_df["category"] = categories_local
        return input_df

    df = run_classification(df)

    # Interactive reassignment for unclassified
    if args.interactive:
        cat_names = [
            ("resistors", "Резисторы"),
            ("capacitors", "Конденсаторы"),
            ("inductors", "Дроссели"),
            ("ics", "Микросхемы"),
            ("connectors", "Разъемы"),
            ("dev_boards", "Отладочные платы"),
            ("diods", "Диоды"),
            ("our_developments", "Наши разработки"),
            ("others", "Другие"),
            ("unclassified", "Не распределено"),
        ]
        uncls = df[df["category"] == "unclassified"].copy()
        max_preview = min(len(uncls), 50)
        print(f"Нераспределено: {len(uncls)}. Покажу первые {max_preview} для разметки.")
        # Load existing rules (append new choices)
        existing_rules: List[Dict[str, Any]] = []
        if args.assign_json and os.path.exists(args.assign_json):
            try:
                with open(args.assign_json, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    if isinstance(data, list):
                        existing_rules = data
            except Exception:
                pass
        for idx, (_, row) in enumerate(uncls.head(max_preview).iterrows(), start=1):
            text_blob = " ".join(str(x) for x in [row.get(desc_col), row.get(value_col), row.get(part_col)] if pd.notna(x))
            print(f"[{idx}] {text_blob}")
            for i, (_, ru) in enumerate(cat_names, start=1):
                print(f"  {i}. {ru}")
            choice = input("Выберите номер категории (Enter чтобы пропустить): ").strip()
            if choice.isdigit():
                ci = int(choice)
                if 1 <= ci <= len(cat_names):
                    selected_key = cat_names[ci - 1][0]
                    df.loc[uncls.index[idx - 1], "category"] = selected_key
                    # Persist rule by 'contains' text blob
                    rule = {"contains": text_blob[:160], "category": selected_key}
                    existing_rules.append(rule)
        # Save updated rules if any
        try:
            if args.assign_json:
                with open(args.assign_json, "w", encoding="utf-8") as f:
                    json.dump(existing_rules, f, ensure_ascii=False, indent=2)
                print(f"Сохранил правила: {args.assign_json}")
        except Exception as exc:
            print(f"Не удалось сохранить правила: {exc}")
        # Optionally redo classification for remaining items if needed (skipped for simplicity)

    outputs = {
        "resistors": df[df["category"] == "resistors"],
        "capacitors": df[df["category"] == "capacitors"],
        "inductors": df[df["category"] == "inductors"],
        "ics": df[df["category"] == "ics"],
        "connectors": df[df["category"] == "connectors"],
        "dev_boards": df[df["category"] == "dev_boards"],
        "optics": df[df["category"] == "optics"],
        "rf_modules": df[df["category"] == "rf_modules"],
        "cables": df[df["category"] == "cables"],
        "power_modules": df[df["category"] == "power_modules"],
        "diods": df[df["category"] == "diods"],
        "our_developments": df[df["category"] == "our_developments"],
        "others": df[df["category"] == "others"],
        "unclassified": df[df["category"] == "unclassified"],
    }

    # Apply assignment rules if provided
    if args.assign_json and os.path.exists(args.assign_json):
        try:
            with open(args.assign_json, "r", encoding="utf-8") as f:
                rules = json.load(f)
        except Exception as exc:
            print(f"Failed to read assign rules: {exc}")
            rules = []
        if isinstance(rules, list):
            df = df.copy()
            lower_cols = set(df.columns)
            for i, rule in enumerate(rules, start=1):
                cat = str(rule.get("category", "")).strip()
                contains = str(rule.get("contains", "")).strip().lower()
                regex = rule.get("regex")
                if not cat or (not contains and not regex):
                    continue
                mask = df["category"] == "unclassified"
                if contains:
                    blob = (
                        df.get("description", pd.Series([""] * len(df))).astype(str).str.lower().fillna("") + " " +
                        df.get("value", pd.Series([""] * len(df))).astype(str).str.lower().fillna("") + " " +
                        df.get("part", pd.Series([""] * len(df))).astype(str).str.lower().fillna("") + " " +
                        df.get("reference", pd.Series([""] * len(df))).astype(str).str.lower().fillna("")
                    )
                    mask = mask & blob.str.contains(re.escape(contains), na=False)
                if regex:
                    try:
                        r = re.compile(regex, re.IGNORECASE)
                        text_series = (
                            df.get("description", "").astype(str).fillna("") + " " +
                            df.get("value", "").astype(str).fillna("") + " " +
                            df.get("part", "").astype(str).fillna("") + " " +
                            df.get("reference", "").astype(str).fillna("")
                        )
                        mask = mask & text_series.apply(lambda t: bool(r.search(t)))
                    except Exception:
                        pass
                df.loc[mask, "category"] = cat
            # refresh outputs
            outputs = {k: df[df["category"] == k] for k in outputs.keys()}

    # Enrich each output with MR code and total quantity per item
    def enrich_with_mr_and_total(frame: pd.DataFrame) -> pd.DataFrame:
        enriched = frame.copy()
        # MR code column
        if mr_col and mr_col in enriched.columns:
            mr_series = enriched[mr_col].fillna("-")
        else:
            mr_series = pd.Series(["-"] * len(enriched), index=enriched.index)
        enriched["Код МР"] = mr_series.astype(str)

        # quantity per row
        if qty_col and qty_col in enriched.columns and pd.api.types.is_numeric_dtype(enriched[qty_col]):
            qty_series = enriched[qty_col]
        else:
            qty_series = pd.Series([1] * len(enriched), index=enriched.index)

        # grouping keys to compute total quantity per item
        group_keys: List[str] = []
        if mr_col and mr_col in enriched.columns and (enriched[mr_col].notna().any()):
            group_keys = [mr_col]
        else:
            for cand in [part_col, value_col, desc_col]:
                if cand and cand in enriched.columns:
                    group_keys.append(cand)
        if not group_keys:
            group_keys = ["category"]

        tmp = enriched.copy()
        tmp["__qty__"] = qty_series
        totals = tmp.groupby(group_keys, dropna=False)["__qty__"].sum().reset_index().rename(columns={"__qty__": "Общее количество"})
        enriched = enriched.merge(totals, on=group_keys, how="left")
        # fill missing
        if "Общее количество" not in enriched.columns:
            enriched["Общее количество"] = 0
        return enriched

    outputs = {name: enrich_with_mr_and_total(df_part) for name, df_part in outputs.items()}

    # Optional CSVs
    if args.out:
        for name, part_df in outputs.items():
            out_path = os.path.join(args.out, f"{name}.csv")
            save_df = part_df.copy()
            inverse_rename = {v: k for k, v in rename_map.items()}
            save_df = save_df.rename(columns=inverse_rename)
            save_df.to_csv(out_path, index=False, encoding="utf-8-sig")

    # XLSX with sheets
    rus_sheet_names = {
        "resistors": "Резисторы",
        "capacitors": "Конденсаторы",
        "inductors": "Дроссели",
        "ics": "Микросхемы",
        "connectors": "Разъемы",
        "dev_boards": "Отладочные платы",
        "optics": "Оптические компоненты",
        "rf_modules": "СВЧ модули",
        "cables": "Кабели",
        "power_modules": "Модули питания",
        "diods": "Диоды",
        "our_developments": "Наши разработки",
        "others": "Другие",
        "unclassified": "Не распределено",
    }

    def merge_existing_if_needed(writer, new_outputs: Dict[str, pd.DataFrame], existing_path: Optional[str]):
        if not existing_path or not os.path.exists(existing_path):
            # just write new_outputs as-is
            for key, part_df in new_outputs.items():
                save_df = part_df.copy()
                inverse_rename = {v: k for k, v in rename_map.items()}
                save_df = save_df.rename(columns=inverse_rename)
                sheet_name = rus_sheet_names.get(key, key)[:31]
                if save_df.empty:
                    save_df = pd.DataFrame(columns=list(inverse_rename.values()) + ["category", "source_file", "source_sheet", "Код МР", "Общее количество"])  # type: ignore
                save_df.to_excel(writer, sheet_name=sheet_name, index=False)
            return

        # merge with existing workbook
        try:
            xls = pd.ExcelFile(existing_path, engine="openpyxl")
        except Exception:
            # if cannot read, fallback to new write
            for key, part_df in new_outputs.items():
                save_df = part_df.copy()
                inverse_rename = {v: k for k, v in rename_map.items()}
                save_df = save_df.rename(columns=inverse_rename)
                sheet_name = rus_sheet_names.get(key, key)[:31]
                if save_df.empty:
                    save_df = pd.DataFrame(columns=list(inverse_rename.values()) + ["category", "source_file", "source_sheet", "Код МР", "Общее количество"])  # type: ignore
                save_df.to_excel(writer, sheet_name=sheet_name, index=False)
            return

        for key, part_df in new_outputs.items():
            inverse_rename = {v: k for k, v in rename_map.items()}
            new_df = part_df.rename(columns=inverse_rename).copy()
            sheet_name = rus_sheet_names.get(key, key)[:31]
            if sheet_name in xls.sheet_names:
                try:
                    old_df = pd.read_excel(existing_path, sheet_name=sheet_name, engine="openpyxl")
                    combined = pd.concat([old_df, new_df], ignore_index=True, sort=False)
                except Exception:
                    combined = new_df
            else:
                combined = new_df
            if combined.empty:
                combined = pd.DataFrame(columns=list(inverse_rename.values()) + ["category", "source_file", "source_sheet", "Код МР", "Общее количество"])  # type: ignore
            combined.to_excel(writer, sheet_name=sheet_name, index=False)

    with pd.ExcelWriter(args.xlsx, engine="openpyxl") as writer:
        merge_existing_if_needed(writer, outputs, args.merge_into)

        if args.combine:
            summary = build_summary(df, ref_col, desc_col, value_col, part_col, qty_col, mr_col)
            summary.to_excel(writer, sheet_name="SUMMARY", index=False)
        # SOURCES sheet
        sources = pd.DataFrame(sorted({(r.get("source_file", ""), r.get("source_sheet", "")) for _, r in df.iterrows()}), columns=["source_file", "source_sheet"])
        sources.to_excel(writer, sheet_name="SOURCES", index=False)

    # Print counts
    print("Split complete:")
    for name, part_df in outputs.items():
        print(f"  {name}: {len(part_df)}")
    print(f"XLSX written: {args.xlsx}")


def build_summary(df: pd.DataFrame, ref_col: Optional[str], desc_col: Optional[str], value_col: Optional[str], part_col: Optional[str], qty_col: Optional[str], mr_col: Optional[str]) -> pd.DataFrame:
    work = df.copy()
    if qty_col and qty_col in work.columns and pd.api.types.is_numeric_dtype(work[qty_col]):
        work["_qty_"] = work[qty_col]
    else:
        work["_qty_"] = 1

    group_keys: List[str] = []
    if mr_col and mr_col in work.columns and (work[mr_col].notna().any()):
        group_keys = [mr_col]
    else:
        for cand in [part_col, value_col, desc_col]:
            if cand and cand in work.columns:
                group_keys.append(cand)
    if not group_keys:
        group_keys = ["category"]

    summary = work.groupby(group_keys, dropna=False, as_index=False)["_qty_"].sum()
    summary = summary.rename(columns={"_qty_": "Общее количество"})

    # Add MR code column
    if mr_col and mr_col in work.columns:
        mr_map = work.groupby(group_keys)[mr_col].agg(lambda s: next((x for x in s if pd.notna(x)), None)).reset_index()
        summary = summary.merge(mr_map, on=group_keys, how="left")
        summary = summary.rename(columns={mr_col: "Код МР"})
    else:
        summary["Код МР"] = "-"

    cols = list(summary.columns)
    if "category" in cols:
        cols = ["category"] + [c for c in cols if c != "category"]
        summary = summary[cols]
    return summary


if __name__ == "__main__":
    main()


