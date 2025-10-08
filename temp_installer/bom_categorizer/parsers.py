# -*- coding: utf-8 -*-
"""
Парсеры для различных форматов BOM файлов

Поддерживаемые форматы:
- TXT: текстовые файлы с разделителями
- DOCX: документы Word с таблицами
- Excel: XLSX файлы
"""

import os
import re
from typing import List, Dict, Any, Optional
import pandas as pd

from .utils import LINE_SPLIT_RE, POS_PREFIX_RE

try:
    from docx import Document  # python-docx
except Exception:
    Document = None  # optional; raise if used without installed


def normalize_cell(s: Any) -> str:
    """Нормализует содержимое ячейки таблицы"""
    return (str(s or "").strip())


def parse_txt_like(path: str) -> pd.DataFrame:
    """
    Парсит текстовый файл с разделителями
    
    Args:
        path: Путь к TXT файлу
        
    Returns:
        DataFrame с колонками: reference, description, qty
    """
    rows: List[Dict[str, Any]] = []
    try:
        with open(path, "r", encoding="utf-8") as f:
            text = f.read()
    except UnicodeDecodeError:
        with open(path, "r", encoding="cp1251", errors="ignore") as f:
            text = f.read()
    
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        parts = [p.strip() for p in LINE_SPLIT_RE.split(line) if p.strip()]
        if not parts:
            continue
        
        # Heuristic mapping: [pos?, name/desc, qty?]
        pos = parts[0] if POS_PREFIX_RE.match(parts[0]) else None
        qty = None
        
        # Try explicit "Nшт" or "N шт" patterns first
        m = re.search(r"(\d+)\s*(шт\.?|pcs|pieces)", line, flags=re.IGNORECASE)
        if m:
            qty = int(m.group(1))
        else:
            for p in parts[::-1]:
                if re.fullmatch(r"\d+", p):
                    qty = int(p)
                    break
        
        desc = " ".join(parts[1:-1]) if pos and len(parts) >= 2 else (" ".join(parts))
        row = {"reference": pos or "", "description": desc, "qty": qty if qty is not None else 1}
        rows.append(row)
    
    if not rows:
        # fallback: whole text in a single row
        rows = [{"description": text, "qty": 1}]
    
    return pd.DataFrame(rows)


def parse_docx(path: str) -> pd.DataFrame:
    """
    Парсит DOCX файл с таблицами
    
    Args:
        path: Путь к DOCX файлу
        
    Returns:
        DataFrame с колонками: zone, reference, description, qty, note
    """
    if Document is None:
        raise SystemExit("python-docx is required to parse DOCX. Install with pip install python-docx")
    
    doc = Document(path)
    extracted: List[Dict[str, Any]] = []

    def guess_header_index(table) -> int:
        """Определяет индекс строки-заголовка в таблице"""
        max_scan = min(5, len(table.rows))
        header_keywords = ["наимен", "обознач", "кол.", "кол ", "кол", "примеч"]
        for i in range(max_scan):
            row_texts = [normalize_cell(c.text).lower() for c in table.rows[i].cells]
            hits = sum(any(hk in t for hk in header_keywords) for t in row_texts)
            if hits >= 2:
                return i
        return 0

    # Parse tables with header detection
    for table in doc.tables:
        if not table.rows:
            continue
        
        header_idx = guess_header_index(table)
        header_cells = [normalize_cell(c.text).strip() for c in table.rows[header_idx].cells]
        header_norm = [h.lower() for h in header_cells]

        # Find column indices by common names
        def find_col_idx(candidates: List[str]) -> Optional[int]:
            for i, h in enumerate(header_norm):
                for cand in candidates:
                    if cand in h:
                        return i
            return None

        idx_zone = find_col_idx(["зона"])  # optional
        idx_ref = find_col_idx(["поз", "обозн"])  # позиционное обозначение
        idx_name = find_col_idx(["наимен"])  # наименование
        idx_qty = find_col_idx(["кол.", "кол ", "кол", "количество"])  # количество
        idx_note = find_col_idx(["примеч"])  # опционально

        # Переменные для хранения текущего ТУ и типа компонента из заголовка группы
        current_group_tu = ""
        current_group_type = ""

        for tr in table.rows[header_idx + 1:]:
            vals = [normalize_cell(c.text) for c in tr.cells]
            if not any(v.strip() for v in vals):
                continue
            
            zone = vals[idx_zone] if idx_zone is not None and idx_zone < len(vals) else ""
            ref = vals[idx_ref] if idx_ref is not None and idx_ref < len(vals) else ""
            name = vals[idx_name] if idx_name is not None and idx_name < len(vals) else ""
            qty_raw = vals[idx_qty] if idx_qty is not None and idx_qty < len(vals) else ""
            note = vals[idx_note] if idx_note is not None and idx_note < len(vals) else ""
            
            # Проверить, является ли это строкой-заголовком группы
            section_headers = ['конденсаторы', 'резисторы', 'микросхемы', 'дроссели', 'индуктивности',
                             'разъемы', 'диоды', 'транзисторы', 'кабели', 'модули', 
                             'набор резисторов', 'набор конденсаторов', 'набор микросхем']
            name_lower = name.strip().lower()
            
            is_group_header = False
            has_tu_code = re.search(r'([А-ЯЁ]{2,}[\.\d]+ТУ)', name) or re.search(r'ТУ\s+[\d\-]+', name)
            
            if not ref.strip() and not qty_raw.strip():
                if any(name_lower.startswith(section) for section in section_headers) or has_tu_code:
                    is_group_header = True
            
            if is_group_header:
                # Извлекаем ТУ
                tu_pattern1 = r'([А-ЯЁ]{2,}\.\d+[\d\.\-]*\s*ТУ)'
                tu_match1 = re.search(tu_pattern1, name)
                if tu_match1:
                    current_group_tu = tu_match1.group(1)
                else:
                    tu_pattern2 = r'([А-ЯЁ]{2,}[\d\.]+\s*ТУ)'
                    tu_match2 = re.search(tu_pattern2, name)
                    if tu_match2:
                        current_group_tu = tu_match2.group(1)
                    else:
                        tu_pattern3 = r'ТУ\s+([\d\-]+)'
                        tu_match3 = re.search(tu_pattern3, name)
                        if tu_match3:
                            current_group_tu = 'ТУ ' + tu_match3.group(1)
                        else:
                            current_group_tu = ""
                
                # Извлекаем тип компонента
                type_text = name
                if current_group_tu:
                    type_text = type_text.replace(current_group_tu, '')
                type_text = re.sub(r'\s+[А-ЯЁ]+\d+[\dА-ЯЁ\-]*', '', type_text)
                type_text = re.sub(r'\s+[А-ЯЁ]+\.\d+[\d\.]*', '', type_text)
                current_group_type = type_text.strip()
                
                continue  # Пропускаем строку-заголовок

            # If header wasn't detected, try fallback mapping
            if not any([ref, name, qty_raw]) and len(vals) >= 2:
                name = " ".join(vals[:-1])
                qty_raw = vals[-1]

            # parse qty
            qty = None
            m = re.search(r"(\d+)", str(qty_raw))
            if m:
                try:
                    qty = int(m.group(1))
                except Exception:
                    qty = 1

            # Добавить ТУ и тип из заголовка группы в note
            if current_group_tu or current_group_type:
                parts = []
                if current_group_type:
                    parts.append(current_group_type)
                if current_group_tu:
                    parts.append(current_group_tu)
                note = ' | '.join(parts) if parts else note

            # Не добавлять строку без данных
            if not ref.strip() and not name.strip():
                continue

            row = {
                "zone": zone,
                "reference": ref,
                "description": name if name else note,
                "qty": qty if qty is not None else 1,
                "note": note,
            }
            extracted.append(row)

    # Additionally parse free text paragraphs (fallback)
    for p in doc.paragraphs:
        t = normalize_cell(p.text)
        if not t:
            continue
        parts = [s.strip() for s in LINE_SPLIT_RE.split(t) if s.strip()]
        if parts:
            pos = parts[0] if POS_PREFIX_RE.match(parts[0]) else ""
            m = re.search(r"(\d+)\s*(шт\.?|pcs|pieces)?", t, flags=re.IGNORECASE)
            qty = int(m.group(1)) if m else 1
            desc = " ".join(parts[1:]) if pos else " ".join(parts)
            extracted.append({"reference": pos, "description": desc, "qty": qty})

    if not extracted:
        extracted = [{"description": " ".join(normalize_cell(p.text) for p in doc.paragraphs), "qty": 1}]
    
    return pd.DataFrame(extracted)
